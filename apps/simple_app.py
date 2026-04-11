"""
简版工具 Flask 应用

单模型（智谱）快速生成简历定制工具。
独立启动入口，端口 5001。
"""

import os
import io
import json
import base64
import uuid
import logging
import time
from datetime import datetime
from pathlib import Path
from typing import Dict, Any, Tuple

from flask import Flask, request, jsonify, render_template, send_file, make_response
from flask_cors import CORS
from flask_limiter import Limiter
from flask_limiter.util import get_remote_address
from werkzeug.utils import secure_filename

# 添加父目录到路径
import sys
sys.path.insert(0, str(Path(__file__).parent.parent))

from core.config import config
from core.providers.zhipu_provider import ZhipuProvider
from core.model_manager import ModelManager
from core.resume_parser import ResumeParser, ParsedResume
from core.resume_builder import ResumeBuilder
from core.expert_team import ExpertTeam, ExpertTeamV2, AnalysisResult, GenerationResult, TailorResultV2
from core.evidence_tracker import EvidenceTracker
from core.resume_generator import ResumeGenerator
from core.cache_manager import CacheManager
from core.template_processor import TemplateProcessor
from core.template_manager import TemplateManager
from core.database import db
from core.auth import send_code, verify_code, login_or_register, login_required, get_current_user, set_login_duration
from core.quota import check_quota, use_quota, get_quota_display
from core.payment import create_payment, handle_payment_notify, query_payment, simulate_payment, get_available_providers

# 是否使用新版五阶段流程
USE_V2_PIPELINE = True

# 配置日志
logging.basicConfig(
    level=logging.INFO,
    format='%(asctime)s - %(name)s - %(levelname)s - %(message)s'
)
logger = logging.getLogger(__name__)

# 文件日志：记录完整流水线过程（特别是 Writer-Reviewer 闭环）
_pipeline_log = logging.FileHandler('storage/pipeline.log', encoding='utf-8', mode='a')
_pipeline_log.setFormatter(logging.Formatter('%(asctime)s - %(name)s - %(levelname)s - %(message)s'))
logging.getLogger().addHandler(_pipeline_log)


def create_app() -> Flask:
    """创建简版工具 Flask 应用"""
    app = Flask(__name__,
                template_folder='../web/templates/simple',
                static_folder='../web/static')

    # SECRET_KEY 安全处理：为空时自动生成随机 key
    secret_key = config.SECRET_KEY
    if not secret_key:
        import secrets
        secret_key = secrets.token_hex(32)
        print("[WARNING] SECRET_KEY 未设置，已自动生成临时密钥（重启后 session 失效）")
        print("  生产环境请设置环境变量 SECRET_KEY")
    app.config['SECRET_KEY'] = secret_key
    app.config['MAX_CONTENT_LENGTH'] = config.MAX_CONTENT_LENGTH
    CORS(app)

    # 全局禁止浏览器缓存（解决开发时反复遇到旧页面缓存的问题）
    @app.after_request
    def add_no_cache_headers(response):
        if 'text/html' in response.content_type:
            response.headers['Cache-Control'] = 'no-cache, no-store, must-revalidate, max-age=0'
            response.headers['Pragma'] = 'no-cache'
            response.headers['Expires'] = '0'
        return response

    # 限流初始化（不设全局默认限流，仅在敏感端点单独限流）
    limiter = Limiter(
        app=app,
        key_func=get_remote_address,
        storage_uri='memory://'
    )

    # 429 限流错误处理器（必须在全局 Exception 处理器之前）
    @app.errorhandler(429)
    def handle_rate_limit(e):
        return jsonify({'error': '请求过于频繁，请稍后再试', 'success': False, 'rate_limited': True}), 429

    # 全局异常处理器 - 防止未捕获的异常导致进程崩溃
    @app.errorhandler(Exception)
    def handle_exception(e):
        import uuid
        error_id = uuid.uuid4().hex[:8]
        logger.error(f"[{error_id}] 未捕获的全局异常: {e}", exc_info=True)

        # 已知异常类型 → 友好文案
        friendly_errors = {
            ValueError: '请求参数无效',
            TypeError: '请求数据类型错误',
            KeyError: '缺少必要的请求数据',
            RuntimeError: '服务处理异常，请稍后重试',
            FileNotFoundError: '请求的资源不存在',
            PermissionError: '权限不足',
        }
        user_msg = friendly_errors.get(type(e), '服务器内部错误，请稍后重试')

        return jsonify({
            'error': user_msg,
            'success': False,
            'error_id': error_id,
        }), 500

    # 初始化组件
    provider = ZhipuProvider()
    model_manager = ModelManager(provider)
    parser = ResumeParser()
    builder = ResumeBuilder()

    # 根据配置选择专家团队版本
    if USE_V2_PIPELINE:
        expert_team = ExpertTeamV2(model_manager)
        logger.info("使用五阶段定制流程 (ExpertTeamV2)")
    else:
        expert_team = ExpertTeam(model_manager)
        logger.info("使用两阶段定制流程 (ExpertTeam)")

    generator = ResumeGenerator()
    cache_manager = CacheManager()
    template_processor = TemplateProcessor()
    template_manager = TemplateManager()

    # 任务状态存储
    task_status: Dict[str, Dict[str, Any]] = {}
    _TASK_STATUS_TTL = 1800  # 30 minutes TTL for stale task entries

    def _cleanup_task_status():
        """Remove stale task_status entries older than TTL."""
        import time as _time
        now = _time.time()
        stale_keys = [
            k for k, v in task_status.items()
            if now - v.get('_created_at', 0) > _TASK_STATUS_TTL
        ]
        for k in stale_keys:
            del task_status[k]
        if stale_keys:
            logger.debug(f"Cleaned {len(stale_keys)} stale task_status entries")


    # ==================== 辅助函数 ====================

    def save_uploaded_file(file, session_id: str, user_id: str = None) -> str:
        upload_dir = Path('storage/uploads') / (user_id or 'anonymous') / session_id
        upload_dir.mkdir(parents=True, exist_ok=True)
        ext = Path(file.filename).suffix.lower()
        file_path = upload_dir / f'original{ext}'
        file.save(str(file_path))
        return str(file_path)

    def save_uploaded_bytes(content: bytes, filename: str, session_id: str, user_id: str = None) -> str:
        upload_dir = Path('storage/uploads') / (user_id or 'anonymous') / session_id
        upload_dir.mkdir(parents=True, exist_ok=True)
        ext = Path(filename).suffix.lower()
        file_path = upload_dir / f'original{ext}'
        with open(file_path, 'wb') as f:
            f.write(content)
        return str(file_path)

    def save_tailored_file(content: bytes, session_id: str, user_id: str = None) -> str:
        output_dir = Path('storage/tailored') / (user_id or 'anonymous') / session_id
        output_dir.mkdir(parents=True, exist_ok=True)
        file_path = output_dir / 'tailored.docx'
        with open(file_path, 'wb') as f:
            f.write(content)
        return str(file_path)

    def parse_jd_info(jd_content: str) -> Tuple[str, str]:
        job_title = ''
        company = ''
        lines = jd_content.strip().split('\n')
        for line in lines[:10]:
            line = line.strip()
            if not line:
                continue
            if any(kw in line for kw in ['职位', '岗位', '招聘', 'Job Title']):
                if '：' in line:
                    job_title = line.split('：')[-1].strip()
                elif ':' in line:
                    job_title = line.split(':')[-1].strip()
            if any(kw in line for kw in ['公司', '企业', 'Company']):
                if '：' in line:
                    company = line.split('：')[-1].strip()
        if not job_title and lines:
            first_line = lines[0].strip()
            if len(first_line) < 100:
                job_title = first_line
        return job_title[:100] if job_title else '', company[:50] if company else ''

    def allowed_file(filename: str) -> bool:
        allowed_extensions = {'.pdf', '.docx', '.doc', '.txt', '.md'}
        return Path(filename).suffix.lower() in allowed_extensions

    def convert_tailored_format(resume: Dict[str, Any]) -> Dict[str, Any]:
        """
        转换 AI 返回的嵌套格式为扁平格式（双重保障）

        AI 可能返回:
        - work_experience[].tailored_bullets: [{content: "...", evidence: {...}}]
        - projects[].tailored_description: "..."

        代码期望:
        - work_experience[].tailored: "..."
        - projects[].tailored: "..."

        改进:
        - 保留【】标记的 JD 关键词
        - 删除原始字段避免模板渲染混乱
        """
        if not resume:
            return resume

        logger.info(f"📝 格式转换: 开始处理 tailored_resume")

        # 处理 work_experience: tailored_bullets -> tailored
        work_exp = resume.get('work_experience', [])
        if isinstance(work_exp, list):
            for exp in work_exp:
                if not isinstance(exp, dict):
                    continue

                # 优先使用已有的 tailored 字符串（增强降级可能已经生成）
                if 'tailored' in exp and exp['tailored']:
                    # 删除原始字段避免模板渲染混乱
                    exp.pop('responsibilities', None)
                    exp.pop('description', None)
                    continue

                # 从 tailored_bullets 合并（保留【】标记）
                if 'tailored_bullets' in exp:
                    bullets = exp.get('tailored_bullets', [])
                    if isinstance(bullets, list) and bullets:
                        contents = []
                        for b in bullets:
                            if isinstance(b, dict):
                                # 保留完整的 content，包含【】标记
                                content = b.get('content', '')
                                if content:
                                    contents.append(content)
                            elif isinstance(b, str):
                                contents.append(b)
                        merged = '\n'.join(filter(None, contents))
                        if merged:
                            exp['tailored'] = merged
                            # 删除原始字段避免模板渲染混乱
                            exp.pop('responsibilities', None)
                            exp.pop('description', None)
                            logger.info(f"📊 格式转换: work_experience {len(bullets)} bullets -> tailored ({len(merged)} 字符)")

        # 处理 projects: tailored_description -> tailored
        projects = resume.get('projects', [])
        if isinstance(projects, list):
            for proj in projects:
                if not isinstance(proj, dict):
                    continue

                # 优先使用已有的 tailored 字符串
                if 'tailored' in proj and proj['tailored']:
                    proj.pop('description', None)
                    continue

                if 'tailored_description' in proj:
                    desc = proj.get('tailored_description', '')
                    if desc:
                        proj['tailored'] = desc
                        proj.pop('description', None)
                        logger.info(f"📊 格式转换: projects tailored_description -> tailored ({len(desc)} 字符)")

        # 处理 education: tailored_highlights -> tailored
        education = resume.get('education', [])
        if isinstance(education, list):
            for edu in education:
                if not isinstance(edu, dict):
                    continue

                # 优先使用已有的 tailored 字符串
                if 'tailored' in edu and edu['tailored']:
                    continue

                if 'tailored_highlights' in edu:
                    highlights = edu.get('tailored_highlights', [])
                    if isinstance(highlights, list) and highlights:
                        merged = '\n'.join(filter(None, highlights))
                        if merged:
                            edu['tailored'] = merged
                            logger.info(f"📊 格式转换: education {len(highlights)} highlights -> tailored ({len(merged)} 字符)")

        return resume

    def get_sample_resume_data() -> Dict[str, Any]:
        """获取示例简历数据用于模板预览"""
        return {
            "basic_info": {
                "name": "张明",
                "phone": "138-0000-0001",
                "email": "zhangming@example.com",
                "location": "北京",
                "age": 28
            },
            "summary": "资深软件工程师，拥有5年互联网行业经验，精通Python、Java开发，具备良好的团队协作能力和项目管理经验。",
            "education": [
                {
                    "school": "北京大学",
                    "degree": "硕士",
                    "major": "计算机科学与技术",
                    "time": "2016-2019",
                    "tailored": "GPA 3.8/4.0，专业排名前5%\n主修课程：数据结构与算法、机器学习、分布式系统\n获国家奖学金、优秀毕业生称号"
                }
            ],
            "work_experience": [
                {
                    "company": "科技有限公司",
                    "position": "高级软件工程师",
                    "time": "2019-至今",
                    "tailored": "主导核心系统架构设计，提升系统性能30%\n带领5人团队完成多个重点项目\n负责技术选型和代码审查",
                    "content": "主导核心系统架构设计，带领团队完成项目开发"
                }
            ],
            "projects": [
                {
                    "name": "智能推荐系统",
                    "role": "技术负责人",
                    "time": "2021-2022",
                    "tailored": "设计并实现基于机器学习的推荐算法\n日均处理百万级用户请求",
                    "content": "智能推荐系统的设计与实现"
                },
                {
                    "name": "行业热点聚合与智能简报系统",
                    "role": "原型设计",
                    "time": "2022-2023",
                    "tailored": "构建自动化信息采集pipeline\n实现结构化简报自动生成",
                },
                {
                    "name": "时令养生食谱生成器",
                    "role": "全栈开发",
                    "time": "2023-2024",
                    "tailored": "搭建垂直领域知识库\n集成LLM实现智能问答",
                }
            ],
            "skills": [
                {"name": "Python", "tailored_description": "精通Python开发，有丰富的Web开发经验"},
                {"name": "Java", "tailored_description": "熟练使用Java进行企业级应用开发"},
                {"name": "MySQL", "tailored_description": "熟悉数据库设计与优化"},
                {"name": "Docker", "tailored_description": "容器化部署与运维"}
            ],
            "awards": [{"name": "年度优秀员工"}, {"name": "技术创新奖"}],
            "certificates": [{"name": "PMP项目管理认证"}, {"name": "AWS架构师认证"}],
            "self_evaluation": "具备扎实的编程基础和丰富的项目经验，善于解决复杂技术问题，有良好的沟通能力和团队协作精神。"
        }

    def run_tailor_pipeline(resume_text: str, jd_content: str, session_id: str = None):
        """
        运行定制流程，统一处理 V1/V2 版本

        Args:
            resume_text: 简历文本
            jd_content: JD内容
            session_id: 会话ID（用于进度更新）

        Returns:
            dict: 包含 tailored_resume, evidence_report, optimization_summary, analysis
        """

        def progress_callback(stage: int, message: str, progress: int):
            """进度回调函数"""
            if session_id and session_id in task_status:
                task_status[session_id]['progress'] = progress
                task_status[session_id]['message'] = message
                task_status[session_id]['stage'] = stage
                logger.info(f"阶段{stage}: {message} ({progress}%)")

        if USE_V2_PIPELINE:
            # 五阶段流程
            result: TailorResultV2 = expert_team.tailor(
                resume_text, jd_content,
                progress_callback=progress_callback
            )

            # 检查是否有错误
            if isinstance(result.tailored_resume, dict) and 'error' in result.tailored_resume:
                error_msg = result.tailored_resume['error']
                logger.error(f"五阶段流程失败: {error_msg}")
                raise RuntimeError(error_msg)

            # 提取 Writer-Reviewer 闭环数据（从 optimization_summary 派生，避免双重构建）
            review_data = {}
            review_loop = result.optimization_summary.get('review_loop') if isinstance(result.optimization_summary, dict) else {}
            if review_loop:
                review_data = {
                    'iterations': review_loop['iterations'],
                    'stop_reason': review_loop['stop_reason'],
                    'feedback_summary': review_loop.get('feedback_summary', ''),
                    'scores': review_loop.get('iteration_details', []),
                }

            return {
                'tailored_resume': result.tailored_resume,
                'evidence_report': result.evidence_report,
                'optimization_summary': result.optimization_summary,
                'analysis': result.analysis,
                'quality_score': result.quality_result.overall_score if result.quality_result else 0,
                'is_v2': True,
                'review_data': review_data,
            }
        else:
            # 两阶段流程（兼容旧版）
            analysis, generation = expert_team.tailor(resume_text, jd_content)
            return {
                'tailored_resume': generation.tailored_resume,
                'evidence_report': generation.evidence_report,
                'optimization_summary': generation.optimization_summary,
                'analysis': {
                    'match_score': analysis.matching_strategy.get('match_score', 0),
                    'match_level': analysis.matching_strategy.get('match_level', ''),
                    'strengths': analysis.matching_strategy.get('strengths', []),
                    'gaps': analysis.matching_strategy.get('gaps', [])
                },
                'quality_score': 0,
                'is_v2': False
            }

    # ==================== 路由 ====================

    @app.route('/')
    def index():
        # 添加防缓存头，确保浏览器始终加载最新的 HTML 文件
        response = make_response(render_template('index.html'))
        response.headers['Cache-Control'] = 'no-cache, no-store, must-revalidate'
        response.headers['Pragma'] = 'no-cache'
        response.headers['Expires'] = '0'
        return response

    @app.route('/favicon.ico')
    def favicon():
        return '', 204  # No Content

    @app.route('/guided')
    def guided_input():
        return render_template('guided_input.html')

    @app.route('/api/health', methods=['GET'])
    def health_check():
        return jsonify({
            'status': 'healthy',
            'version': '1.0.0',
            'mode': 'simple',
            'provider': 'zhipu',
            'timestamp': datetime.now().isoformat()
        })

    @app.route('/api/user_params', methods=['GET'])
    def get_user_params():
        """读取保存的用户参数（服务端持久化，重启不丢失）"""
        params_file = Path('data/user_params.json')
        if params_file.exists():
            try:
                return jsonify(json.loads(params_file.read_text(encoding='utf-8')))
            except Exception:
                return jsonify({})
        return jsonify({})

    @app.route('/api/user_params', methods=['POST'])
    def save_user_params():
        """保存用户参数到服务端"""
        params = request.json
        params_file = Path('data/user_params.json')
        params_file.parent.mkdir(exist_ok=True)
        params_file.write_text(json.dumps(params, ensure_ascii=False), encoding='utf-8')
        return jsonify({'success': True})

    @app.route('/api/shutdown', methods=['POST'])
    def shutdown():
        """关闭服务（由主服务器调用）"""
        import threading
        import os

        def do_shutdown():
            time.sleep(1)
            os._exit(0)

        logger.info("Received shutdown request")
        threading.Thread(target=do_shutdown, daemon=True).start()
        return jsonify({'status': 'shutting_down'})

    @app.route('/api/cache/clear', methods=['POST'])
    def api_cache_clear():
        """手动清除所有缓存"""
        try:
            count = cache_manager.clear_all()
            # 同时清除 preprocessed 模板缓存
            preprocessed_dir = config.BASE_DIR / 'templates' / 'preprocessed'
            if preprocessed_dir.exists():
                for f in preprocessed_dir.glob('*.docx'):
                    f.unlink()
                logger.info(f"  templates/preprocessed/: 已清理所有预处理模板")
            logger.info(f"手动清缓存: {count} 个文件")
            return jsonify({'success': True, 'cleared': count})
        except Exception as e:
            logger.error(f"清缓存失败: {e}")
            return jsonify({'success': False, 'error': str(e)}), 500

    # ==================== 认证 API ====================

    @app.route('/api/auth/send-code', methods=['POST'])
    @limiter.limit("1 per minute")
    def api_send_code():
        """发送邮箱验证码"""
        data = request.get_json() or {}
        email = data.get('email', '').strip()
        if not email:
            return jsonify({'success': False, 'error': '请输入邮箱'}), 400
        if not send_code(email):
            return jsonify({'success': False, 'error': '验证码发送失败，请稍后再试'}), 400
        return jsonify({'success': True, 'message': '验证码已发送到您的邮箱'})

    @app.route('/api/auth/login', methods=['POST'])
    @limiter.limit("5 per minute")
    def api_login():
        """邮箱验证码登录/注册"""
        data = request.get_json() or {}
        email = data.get('email', '').strip()
        code = data.get('code', '').strip()
        duration = data.get('duration', 'session')  # session / 7d / 30d / forever
        if not email or not code:
            return jsonify({'success': False, 'error': '邮箱和验证码不能为空'}), 400
        if not verify_code(email, code):
            return jsonify({'success': False, 'error': '验证码错误或已过期'}), 401
        try:
            result = login_or_register(email)
            from flask import session
            session.clear()  # 防止 session fixation
            session['user_id'] = result['user_id']
            session['email'] = result['email']
            # 设置登录有效期
            set_login_duration(duration)
            quota_info = get_quota_display(result['user_id'])
            return jsonify({
                'success': True,
                'user': {
                    'user_id': result['user_id'],
                    'email': result['email'],
                    'is_new_user': result['is_new_user'],
                },
                'quota': quota_info
            })
        except Exception as e:
            logger.error(f"登录失败: {e}")
            return jsonify({'success': False, 'error': '登录失败'}), 500

    @app.route('/api/auth/logout', methods=['POST'])
    def api_logout():
        """退出登录"""
        from flask import session
        session.clear()
        return jsonify({'success': True})

    @app.route('/api/auth/me', methods=['GET'])
    @login_required
    def api_auth_me():
        """获取当前用户信息"""
        user = get_current_user()
        quota_info = get_quota_display(request.user_id)
        return jsonify({
            'success': True,
            'user': {
                'user_id': user['id'],
                'email': user['email'],
                'phone': user.get('phone', ''),
                'nickname': user.get('nickname', ''),
                'plan_type': user['plan_type'],
                'created_at': user['created_at'],
            },
            'quota': quota_info
        })

    # ==================== 配额 API ====================

    @app.route('/api/quota', methods=['GET'])
    @login_required
    def api_get_quota():
        """获取当前用户配额"""
        quota_info = get_quota_display(request.user_id)
        return jsonify({'success': True, 'quota': quota_info})

    # ==================== 支付 API ====================

    @app.route('/api/payment/plans', methods=['GET'])
    def api_get_plans():
        """获取所有套餐信息"""
        plans = []
        for key, val in config.PLANS.items():
            plans.append({
                'type': key,
                'name': val['name'],
                'price': val['price'],
                'quota': val['quota'],
                'daily_limit': val['daily_limit'],
            })
        return jsonify({'success': True, 'plans': plans})

    @app.route('/api/payment/providers', methods=['GET'])
    def api_get_providers():
        """获取可用支付方式列表"""
        providers = get_available_providers()
        return jsonify({'success': True, 'providers': providers})

    @app.route('/api/payment/create', methods=['POST'])
    @login_required
    def api_create_payment():
        """创建支付订单"""
        data = request.get_json() or {}
        plan_type = data.get('plan_type', '')
        provider_id = data.get('provider', '')
        if plan_type not in config.PLANS or config.PLANS[plan_type]['price'] <= 0:
            return jsonify({'success': False, 'error': '无效的套餐'}), 400
        try:
            result = create_payment(request.user_id, plan_type, provider_id=provider_id)
            return jsonify({'success': True, **result})
        except Exception as e:
            logger.error(f"创建支付失败: {e}")
            return jsonify({'success': False, 'error': str(e)}), 500

    @app.route('/api/payment/query/<order_no>', methods=['GET'])
    @login_required
    def api_query_payment(order_no):
        """查询支付状态"""
        result = query_payment(order_no)
        return jsonify({'success': True, **result})

    @app.route('/api/payment/notify/alipay', methods=['POST'])
    def api_alipay_notify():
        """支付宝异步回调通知"""
        success = handle_payment_notify(request, 'alipay')
        if success:
            return 'success'
        return 'fail', 400

    @app.route('/api/payment/notify/wechat', methods=['POST'])
    def api_wechat_notify():
        """微信支付回调通知"""
        success = handle_payment_notify(request, 'wechat')
        if success:
            return jsonify({'code': 'SUCCESS', 'message': 'OK'})
        return jsonify({'code': 'FAIL', 'message': '处理失败'}), 500

    @app.route('/api/payment/simulate', methods=['POST'])
    @login_required
    def api_simulate_payment():
        """模拟支付成功（仅沙箱环境）"""
        data = request.get_json() or {}
        order_no = data.get('order_no', '')
        if not order_no:
            return jsonify({'success': False, 'error': '缺少订单号'}), 400
        if simulate_payment(order_no):
            quota_info = get_quota_display(request.user_id)
            return jsonify({'success': True, 'quota': quota_info})
        return jsonify({'success': False, 'error': '模拟支付失败'}), 400

    # ==================== 用户中心 API ====================

    @app.route('/api/user/history', methods=['GET'])
    @login_required
    def api_user_history():
        """获取用户使用历史"""
        records = db.get_user_usage_history(request.user_id, limit=20)
        return jsonify({'success': True, 'records': records})

    @app.route('/api/user/orders', methods=['GET'])
    @login_required
    def api_user_orders():
        """获取用户订单列表"""
        orders = db.get_user_orders(request.user_id)
        return jsonify({'success': True, 'orders': orders})

    def _check_quota_or_anonymous():
        """配额检查（复用逻辑，避免 tailor_file/text/form 重复代码）。
        Returns: (error_response, current_user) — error_response 为 None 表示通过。
        """
        current_user = get_current_user()
        if current_user:
            can_use, quota_info = check_quota(current_user['id'])
            if not can_use:
                return (jsonify({'error': quota_info.get('reason', '配额已用完'), 'quota_exceeded': True}), 403), current_user
        else:
            from flask import session as flask_session
            if not flask_session.get('anonymous_used'):
                flask_session['anonymous_used'] = True
            else:
                return (jsonify({'error': '免费体验次数已用完，请登录获取更多次数', 'need_login': True}), 403), current_user
        return None, current_user

    @app.route('/api/tailor/file', methods=['POST'])
    @limiter.limit("5 per hour")
    def tailor_file():
        try:
            # 配额检查（允许未登录用户使用免费额度）
            err_resp, current_user = _check_quota_or_anonymous()
            if err_resp:
                return err_resp

            start_time = time.time()

            # 记录请求参数
            logger.info(f"📝 收到定制请求 - template_mode: {request.form.get('template_mode')}, template_id: {request.form.get('template_id')}")

            if 'resume' not in request.files:
                return jsonify({'error': '未上传简历文件'}), 400

            resume_file = request.files['resume']
            if resume_file.filename == '':
                return jsonify({'error': '未选择简历文件'}), 400

            if not allowed_file(resume_file.filename):
                return jsonify({'error': '不支持的简历格式'}), 400

            jd_content = ''
            if 'jd' in request.files and request.files['jd'].filename != '':
                jd_file = request.files['jd']
                if allowed_file(jd_file.filename):
                    jd_content = jd_file.read().decode('utf-8', errors='ignore')
            elif request.form.get('jd_text'):
                jd_content = request.form.get('jd_text', '')

            if not jd_content:
                return jsonify({'error': '未提供职位JD'}), 400

            session_id = str(uuid.uuid4())
            _cleanup_task_status(); task_status[session_id] = {'status': 'processing', 'progress': 0, 'message': '正在解析简历...', '_created_at': time.time()}

            resume_content = resume_file.read()
            user_id_str = str(current_user['id']) if current_user else None
            save_uploaded_bytes(resume_content, resume_file.filename, session_id, user_id_str)

            parsed_resume = parser.parse(
                file_content=resume_content,
                filename=resume_file.filename
            )

            template_result = None
            original_doc = None
            if parsed_resume.source_format == 'word':
                task_status[session_id]['progress'] = 10
                task_status[session_id]['message'] = '正在提取模板样式...'
                try:
                    from docx import Document
                    original_doc = Document(io.BytesIO(resume_content))
                    template_result = template_processor.preprocess(
                        original_doc,
                        resume_file.filename,
                        original_content=resume_content
                    )
                except Exception as e:
                    logger.warning(f"模板预处理异常: {e}")

            task_status[session_id]['progress'] = 20
            task_status[session_id]['message'] = '正在分析JD需求...'

            # 提前提取模板参数（缓存快速路径需要）
            template_id = request.form.get('template_id', '') or request.form.get('templateId', '')
            template_mode = request.form.get('template_mode', '') or request.form.get('templateMode', 'auto')

            # 检查是否强制跳过缓存（用户点击"开始定制简历"时默认跳过）
            no_cache = request.form.get('no_cache', 'true').lower() == 'true'

            cached_result = cache_manager.get(parsed_resume.raw_text, jd_content)
            cached_tailored = None
            cached_analysis = None
            cached_evidence = None
            cached_quality_score = 0
            cached_optimization = {}
            if cached_result and cached_result.get('tailored_resume'):
                cached_tailored = cached_result['tailored_resume']
                cached_analysis = cached_result.get('analysis', {})
                cached_evidence = cached_result.get('evidence_report', {})
                cached_quality_score = cached_result.get('quality_score', 0)
                cached_optimization = cached_result.get('optimization_summary', {})

            # 快速路径：JD+简历不变，用户切换了模板 → 复用缓存 AI 内容，只重新渲染
            # 注意：必须检查 no_cache，否则代码/prompt 更新后仍返回旧缓存
            if not no_cache and cached_tailored and template_mode == 'selected' and template_id:
                logger.info(f"快速路径: 复用缓存 tailored_resume + 新模板 {template_id}")
                task_status[session_id]['progress'] = 80
                task_status[session_id]['message'] = '正在渲染模板...'

                tailored_resume = convert_tailored_format(cached_tailored)
                analysis_data = cached_analysis
                is_v2 = True
                evidence_report = cached_evidence

                # 直接跳到模板渲染（与正常流程共享渲染逻辑）
                style_preserved = False
                used_template_id = None
                selected_template = template_manager.get_template(template_id)
                if not selected_template:
                    task_status[session_id]['status'] = 'failed'
                    task_status[session_id]['message'] = f'选定的模板不存在: {template_id}'
                    return jsonify({'error': f'选定的模板不存在: {template_id}', 'error_code': 'TEMPLATE_NOT_FOUND'}), 400

                try:
                    word_bytes = template_processor.render(template_id, tailored_resume, parsed_resume.style_metadata)
                    used_template_id = template_id
                    style_preserved = True
                    template_manager.increment_use_count(template_id)
                    logger.info(f"快速路径渲染成功: {selected_template.get('name', template_id)}")
                except Exception as e:
                    logger.error(f"快速路径渲染失败: {e}", exc_info=True)
                    task_status[session_id]['status'] = 'failed'
                    task_status[session_id]['message'] = f'模板渲染失败: {str(e)}'
                    return jsonify({'error': f'模板渲染失败: {str(e)}', 'error_code': 'TEMPLATE_RENDER_FAILED'}), 500

                # 构建返回结果（与正常流程一致）
                result = {
                    'tailored_resume': tailored_resume,
                    'analysis': analysis_data,
                    'quality_score': cached_quality_score,
                    'optimization_summary': cached_optimization,
                    'style_preserved': style_preserved,
                    'used_template_id': used_template_id,
                    'pipeline_version': 'v2',
                    'cache_hit': True,
                }
                if cached_result.get('review_data'):
                    result['review_data'] = cached_result['review_data']

                task_status[session_id]['progress'] = 100
                task_status[session_id]['status'] = 'completed'
                save_tailored_file(word_bytes, session_id)
                return jsonify(result)

            # 完整缓存命中（非模板切换场景）
            if not no_cache and cached_result and cached_result.get('tailored_resume'):
                logger.info(f"命中缓存且包含 tailored_resume: session={session_id}")
                task_status[session_id]['progress'] = 100
                task_status[session_id]['status'] = 'completed'
                return jsonify(cached_result)

            task_status[session_id]['message'] = '正在AI分析...'

            # 使用统一的定制流程（带进度回调）
            pipeline_result = run_tailor_pipeline(parsed_resume.raw_text, jd_content, session_id)
            tailored_resume = pipeline_result['tailored_resume']

            # 格式转换：将 AI 返回的嵌套格式转为扁平格式（双重保障）
            tailored_resume = convert_tailored_format(tailored_resume)

            analysis_data = pipeline_result['analysis']
            is_v2 = pipeline_result['is_v2']

            task_status[session_id]['progress'] = 60
            task_status[session_id]['message'] = '正在验证依据...'

            # V2 版本已有依据报告，V1 版本需要单独验证
            if is_v2:
                evidence_report = pipeline_result['evidence_report']
            else:
                evidence_tracker = EvidenceTracker(model_manager)
                evidence_report = evidence_tracker.validate_resume(
                    parsed_resume.raw_text,
                    tailored_resume
                ).to_dict()

            task_status[session_id]['progress'] = 80
            task_status[session_id]['message'] = '正在生成文档...'

            output_format = request.form.get('format', 'word')  # word / ats_pdf
            logger.info(f"📝 output_format = '{output_format}' (from form)")
            style = request.form.get('style', 'original')
            style_preserved = False
            used_template_id = None

            # ATS PDF 模式：使用 career-ops 风格的 ATS 优化输出
            if output_format == 'ats_pdf':
                task_status[session_id]['message'] = '正在生成 ATS 优化简历...'
                try:
                    jd_keywords = analysis_data.get('matching_strategy', {}).get('must_have_skills', [])
                    if not jd_keywords:
                        jd_keywords = analysis_data.get('matching_strategy', {}).get('strengths', [])

                    ats_html_path = generator.generate_ats_html(
                        tailored_resume,
                        jd_keywords=jd_keywords
                    )
                    ats_pdf_path = ats_html_path.replace('.html', '.pdf')

                    import subprocess as _subprocess
                    pdf_tool = Path('tools/generate-pdf.mjs')
                    if pdf_tool.exists():
                        proc = _subprocess.run(
                            ['node', str(pdf_tool), ats_html_path, ats_pdf_path, '--format=letter'],
                            capture_output=True, text=True, timeout=30
                        )
                        if proc.returncode == 0:
                            with open(ats_pdf_path, 'rb') as f:
                                ats_pdf_bytes = f.read()
                            task_status[session_id]['progress'] = 100
                            task_status[session_id]['status'] = 'completed'
                            task_status[session_id]['message'] = 'ATS 简历生成完成'

                            candidate_name = parsed_resume.basic_info.get('name', '')
                            job_title, company = parse_jd_info(jd_content)

                            result = {
                                'session_id': session_id,
                                'status': 'completed',
                                'processing_time': int((time.time() - start_time) * 1000),
                                'tailored_word': '',
                                'tailored_ats_pdf': base64.b64encode(ats_pdf_bytes).decode('utf-8'),
                                'evidence_report': evidence_report if isinstance(evidence_report, dict) else evidence_report.to_dict(),
                                'validation_result': 'pass' if (evidence_report.coverage if hasattr(evidence_report, 'coverage') else evidence_report.get('coverage', 0)) >= 0.9 else 'pass_with_review',
                                'output_format': 'ats_pdf',
                                'style_preserved': False,
                                'template_used': False,
                                'template_mode': 'ats',
                                'analysis': analysis_data,
                                'pipeline_version': 'v2' if is_v2 else 'v1',
                                'candidate_name': candidate_name,
                                'job_title': job_title
                            }
                            if is_v2:
                                result['quality_score'] = pipeline_result.get('quality_score', 0)
                                result['optimization_summary'] = pipeline_result.get('optimization_summary', {})
                            if pipeline_result.get('review_data'):
                                result['review_data'] = pipeline_result['review_data']
                            cache_manager.set(parsed_resume.raw_text, jd_content, result)
                            logger.info(f"ATS PDF 生成成功: {ats_pdf_path}")
                            return jsonify(result)
                        else:
                            logger.error(f"ATS PDF 生成失败: {proc.stderr}")
                    logger.error("ATS PDF 工具不存在")
                except Exception as e:
                    logger.error(f"ATS PDF 异常: {e}", exc_info=True)

            # 根据模板模式选择渲染方式 - 增强日志用于诊断
            logger.debug(f"模板参数详情:")
            logger.debug(f"   template_mode = '{template_mode}'")
            logger.debug(f"   template_id = '{template_id}'")
            logger.debug(f"   条件判断: selected={template_mode == 'selected' and bool(template_id)}")
            logger.debug(f"原始表单数据: {list(request.form.keys())}")
            logger.debug(f"AI生成数据 keys: {list(tailored_resume.keys())}")

            if template_mode == 'selected' and template_id:
                # 用户指定模板 - 强制使用选定模板，不降级
                selected_template = template_manager.get_template(template_id)
                logger.debug(f"获取模板信息: {selected_template.get('name') if selected_template else 'None'}")
                if not selected_template:
                    # 模板不存在，返回错误
                    task_status[session_id]['status'] = 'failed'
                    task_status[session_id]['message'] = f'选定的模板不存在: {template_id}'
                    return jsonify({
                        'error': f'选定的模板不存在: {template_id}',
                        'error_code': 'TEMPLATE_NOT_FOUND'
                    }), 400

                try:
                    logger.info(f"开始渲染模板: {template_id}, 文件路径: {selected_template.get('file_path')}")
                    word_bytes = template_processor.render(
                        template_id, tailored_resume, parsed_resume.style_metadata
                    )
                    used_template_id = template_id
                    style_preserved = True
                    template_manager.increment_use_count(template_id)
                    logger.info(f"✅ 模板渲染成功: {selected_template.get('name', template_id)} ({template_id})")
                except Exception as e:
                    logger.error(f"❌ 选定模板渲染失败: {e}", exc_info=True)
                    task_status[session_id]['status'] = 'failed'
                    task_status[session_id]['message'] = f'模板渲染失败: {str(e)}'
                    return jsonify({
                        'error': f'模板渲染失败: {str(e)}',
                        'error_code': 'TEMPLATE_RENDER_FAILED'
                    }), 500
            elif template_mode == 'original' or (template_mode == 'auto' and template_result and template_result.success and original_doc):
                # 使用原简历样式
                try:
                    word_bytes, used_template = template_processor.render_with_fallback(
                        original_doc,
                        tailored_resume,
                        parsed_resume.style_metadata,
                        resume_file.filename,
                        preprocess_result=template_result,
                        original_content=resume_content
                    )
                    style_preserved = used_template
                    if used_template:
                        used_template_id = template_result.template_id
                except Exception as e:
                    logger.warning(f"模板渲染失败: {e}")
                    try:
                        word_bytes = generator.generate_bytes(
                            tailored_resume,
                            style_metadata=parsed_resume.style_metadata
                        )
                    except Exception as e2:
                        logger.error(f"降级生成器也失败: {e2}", exc_info=True)
                        raise
            else:
                # 使用默认模板或生成器
                default_template = template_manager.get_default_template()
                if default_template and template_mode == 'auto':
                    try:
                        word_bytes = template_processor.render(
                            default_template['template_id'],
                            tailored_resume,
                            parsed_resume.style_metadata
                        )
                        used_template_id = default_template['template_id']
                        style_preserved = True
                        template_manager.increment_use_count(default_template['template_id'])
                        logger.info(f"使用默认模板: {default_template.get('name', '默认模板')}")
                    except Exception as e:
                        logger.warning(f"默认模板渲染失败: {e}")
                        try:
                            word_bytes = generator.generate_bytes(
                                tailored_resume,
                                style_metadata=parsed_resume.style_metadata
                            )
                        except Exception as e2:
                            logger.error(f"降级生成器也失败: {e2}", exc_info=True)
                            raise
                else:
                    try:
                        word_bytes = generator.generate_bytes(
                            tailored_resume,
                            style_metadata=parsed_resume.style_metadata
                        )
                    except Exception as e2:
                        logger.error(f"无模板生成失败: {e2}", exc_info=True)
                        raise

            task_status[session_id]['progress'] = 100
            task_status[session_id]['status'] = 'completed'

            processing_time_ms = int((time.time() - start_time) * 1000)

            # 计算覆盖率
            if isinstance(evidence_report, dict):
                coverage = evidence_report.get('coverage', 0)
            else:
                coverage = evidence_report.coverage if hasattr(evidence_report, 'coverage') else 0

            # 获取模板名称（用于前端显示）
            template_name = None
            if used_template_id:
                tmpl = template_manager.get_template(used_template_id)
                if tmpl:
                    template_name = tmpl.get('name')

            # 获取候选人姓名和职位名称（用于文件名生成）
            candidate_name = parsed_resume.basic_info.get('name', '')
            job_title, company = parse_jd_info(jd_content)

            result = {
                'session_id': session_id,
                'status': 'completed',
                'processing_time': processing_time_ms,
                'tailored_resume': tailored_resume,
                'tailored_word': base64.b64encode(word_bytes).decode('utf-8'),
                'evidence_report': evidence_report if isinstance(evidence_report, dict) else evidence_report.to_dict(),
                'validation_result': 'pass' if coverage >= 0.9 else 'pass_with_review',
                'style_preserved': style_preserved,
                'template_id': used_template_id,
                'template_used': used_template_id is not None,
                'template_name': template_name,
                'template_mode': template_mode,
                'style_info': {
                    'font': parsed_resume.style_metadata.primary_font,
                    'font_size': parsed_resume.style_metadata.body_font_size,
                    'source': parsed_resume.style_metadata.source
                },
                'analysis': analysis_data,
                'pipeline_version': 'v2' if is_v2 else 'v1',
                'candidate_name': candidate_name,
                'job_title': job_title
            }

            # 添加匹配度日志
            logger.info(f"📊 匹配度数据: score={analysis_data.get('match_score')}, level={analysis_data.get('match_level')}")
            logger.info(f"📊 完整 analysis: {analysis_data}")

            # V2 版本添加额外信息
            if is_v2:
                result['quality_score'] = pipeline_result.get('quality_score', 0)
                result['optimization_summary'] = pipeline_result.get('optimization_summary', {})
            if pipeline_result.get('review_data'):
                result['review_data'] = pipeline_result['review_data']

            cache_manager.set(parsed_resume.raw_text, jd_content, result)
            save_tailored_file(word_bytes, session_id)

            has_experience = bool(parsed_resume.work_experience)
            db.save_history(session_id, {
                'candidate_name': parsed_resume.basic_info.get('name', ''),
                'candidate_type': 'experienced' if has_experience else 'entry_level',
                'job_title': job_title,
                'company': company,
                'match_score': analysis_data.get('match_score', 0),
                'match_level': analysis_data.get('match_level', ''),
                'original_resume': parsed_resume.raw_text,
                'tailored_resume': tailored_resume,
                'jd_content': jd_content,
                'evidence_report': evidence_report if isinstance(evidence_report, dict) else evidence_report.to_dict(),
                'optimization_summary': {'style_preserved': style_preserved, 'pipeline_version': 'v2' if is_v2 else 'v1', **pipeline_result.get('optimization_summary', {})},
                'model_used': model_manager.current_model,
                'tokens_used': 0,
                'processing_time_ms': processing_time_ms
            })

            # 扣减用户配额
            if current_user:
                use_quota(current_user['id'], session_id=session_id)

            return jsonify(result)

        except Exception as e:
            import traceback as tb_module
            full_tb = tb_module.format_exc()
            logger.error(f"文件定制失败: {e}\n{full_tb}")
            error_str = str(e)
            error_type = type(e).__name__
            # 更详细的错误分类
            if "429" in error_str or "rate" in error_str.lower() or "限制" in error_str or "并发" in error_str:
                return jsonify({'error': 'API调用频率超限，请等待1-2分钟后重试'}), 429
            elif "resume_analysis" in error_str or "{" in error_str:
                return jsonify({'error': 'AI响应解析失败，请重试'}), 500
            elif "timeout" in error_str.lower():
                return jsonify({'error': 'AI处理超时，请稍后重试'}), 504
            elif "connection" in error_str.lower():
                return jsonify({'error': 'AI服务连接失败，请检查网络'}), 503
            logger.error(f"文件定制详细错误类型: {error_type}, 消息: {error_str}")
            return jsonify({
                'error': '处理失败，请稍后重试',
                'success': False,
            }), 500

    @app.route('/api/tailor/text', methods=['POST'])
    def tailor_text():
        """处理纯文本简历定制请求"""
        try:
            # 配额检查
            err_resp, current_user = _check_quota_or_anonymous()
            if err_resp:
                return err_resp

            start_time = time.time()

            data = request.get_json()
            if not data:
                return jsonify({'error': '未提供请求数据'}), 400

            resume_text = data.get('resume_text', '')
            jd_content = data.get('jd_text', '')
            template_mode = data.get('template_mode', 'auto')
            template_id = data.get('template_id', '')

            if not resume_text:
                return jsonify({'error': '未提供简历内容'}), 400
            if not jd_content:
                return jsonify({'error': '未提供职位JD'}), 400

            session_id = str(uuid.uuid4())
            _cleanup_task_status(); task_status[session_id] = {'status': 'processing', 'progress': 0, 'message': '正在分析...', '_created_at': time.time()}

            task_status[session_id]['progress'] = 20
            task_status[session_id]['message'] = '正在分析JD需求...'

            # 检查是否强制跳过缓存（用户点击"开始定制简历"时默认跳过）
            no_cache = data.get('no_cache', True)

            cached_result = cache_manager.get(resume_text, jd_content)
            # 检查缓存是否包含必要字段（旧缓存可能缺少 tailored_resume）
            # 只有 no_cache=False 时才使用缓存
            if not no_cache and cached_result and cached_result.get('tailored_resume'):
                logger.info(f"命中缓存且包含 tailored_resume: session={session_id}")
                task_status[session_id]['progress'] = 100
                task_status[session_id]['status'] = 'completed'
                return jsonify(cached_result)

            task_status[session_id]['message'] = '正在AI分析...'

            # 使用统一的定制流程（带进度回调）
            pipeline_result = run_tailor_pipeline(resume_text, jd_content, session_id)
            tailored_resume = pipeline_result['tailored_resume']

            # 格式转换：将 AI 返回的嵌套格式转为扁平格式（双重保障）
            tailored_resume = convert_tailored_format(tailored_resume)

            analysis_data = pipeline_result['analysis']
            is_v2 = pipeline_result['is_v2']

            task_status[session_id]['progress'] = 96
            task_status[session_id]['message'] = '正在验证依据...'

            # V2 版本已有依据报告，V1 版本需要单独验证
            if is_v2:
                evidence_report = pipeline_result['evidence_report']
            else:
                evidence_tracker = EvidenceTracker(model_manager)
                evidence_report = evidence_tracker.validate_resume(
                    resume_text,
                    tailored_resume
                ).to_dict()

            task_status[session_id]['progress'] = 98
            task_status[session_id]['message'] = '正在生成文档...'

            # ATS PDF 模式
            output_format = data.get('format', 'word')
            logger.info(f"📝 [引导模式] output_format = '{output_format}' (from json)")
            if output_format == 'ats_pdf':
                task_status[session_id]['message'] = '正在生成 ATS 优化简历...'
                try:
                    jd_keywords = analysis_data.get('matching_strategy', {}).get('must_have_skills', [])
                    if not jd_keywords:
                        jd_keywords = analysis_data.get('matching_strategy', {}).get('strengths', [])

                    ats_html_path = generator.generate_ats_html(
                        tailored_resume,
                        jd_keywords=jd_keywords
                    )
                    ats_pdf_path = ats_html_path.replace('.html', '.pdf')

                    import subprocess as _subprocess
                    pdf_tool = Path('tools/generate-pdf.mjs')
                    if pdf_tool.exists():
                        proc = _subprocess.run(
                            ['node', str(pdf_tool), ats_html_path, ats_pdf_path, '--format=letter'],
                            capture_output=True, text=True, timeout=30
                        )
                        if proc.returncode == 0:
                            with open(ats_pdf_path, 'rb') as f:
                                ats_pdf_bytes = f.read()
                            task_status[session_id]['progress'] = 100
                            task_status[session_id]['status'] = 'completed'
                            task_status[session_id]['message'] = 'ATS 简历生成完成'

                            result = {
                                'session_id': session_id,
                                'status': 'completed',
                                'processing_time': int((time.time() - start_time) * 1000),
                                'tailored_word': '',
                                'tailored_ats_pdf': base64.b64encode(ats_pdf_bytes).decode('utf-8'),
                                'evidence_report': evidence_report if isinstance(evidence_report, dict) else evidence_report.to_dict(),
                                'validation_result': 'pass' if (evidence_report.coverage if hasattr(evidence_report, 'coverage') else evidence_report.get('coverage', 0)) >= 0.9 else 'pass_with_review',
                                'output_format': 'ats_pdf',
                                'style_preserved': False,
                                'template_used': False,
                                'template_mode': 'ats',
                                'analysis': analysis_data,
                                'pipeline_version': 'v2' if is_v2 else 'v1',
                                'candidate_name': tailored_resume.get('basic_info', {}).get('name', ''),
                                'job_title': ''
                            }
                            if is_v2:
                                result['quality_score'] = pipeline_result.get('quality_score', 0)
                                result['optimization_summary'] = pipeline_result.get('optimization_summary', {})
                            if pipeline_result.get('review_data'):
                                result['review_data'] = pipeline_result['review_data']
                            cache_manager.set(resume_text, jd_content, result)
                            logger.info(f"ATS PDF 生成成功(引导模式): {ats_pdf_path}")
                            return jsonify(result)
                        else:
                            logger.error(f"ATS PDF 生成失败(引导模式): {proc.stderr}")
                except Exception as e:
                    logger.error(f"ATS PDF 异常(引导模式): {e}", exc_info=True)

            # 根据模板模式选择渲染方式
            logger.info(f"引导模式 - 模板模式: {template_mode}, template_id: {template_id}")
            used_template_id = None
            style_preserved = False

            if template_mode == 'selected' and template_id:
                # 用户指定模板 - 强制使用选定模板，不降级
                selected_template = template_manager.get_template(template_id)
                if not selected_template:
                    # 模板不存在，返回错误
                    task_status[session_id]['status'] = 'failed'
                    task_status[session_id]['message'] = f'选定的模板不存在: {template_id}'
                    return jsonify({
                        'error': f'选定的模板不存在: {template_id}',
                        'error_code': 'TEMPLATE_NOT_FOUND'
                    }), 400

                try:
                    word_bytes = template_processor.render(
                        template_id, tailored_resume, None
                    )
                    used_template_id = template_id
                    style_preserved = True
                    template_manager.increment_use_count(template_id)
                    logger.info(f"引导模式使用指定模板: {selected_template.get('name')}")
                except Exception as e:
                    logger.error(f"模板渲染失败: {e}", exc_info=True)
                    task_status[session_id]['status'] = 'failed'
                    task_status[session_id]['message'] = f'模板渲染失败: {str(e)}'
                    return jsonify({
                        'error': f'模板渲染失败: {str(e)}',
                        'error_code': 'TEMPLATE_RENDER_FAILED'
                    }), 500
            else:
                # auto 模式或 original 模式：使用默认模板或生成器
                default_template = template_manager.get_default_template()
                if default_template:
                    try:
                        word_bytes = template_processor.render(
                            default_template['template_id'],
                            tailored_resume,
                            None
                        )
                        used_template_id = default_template['template_id']
                        style_preserved = True
                        template_manager.increment_use_count(default_template['template_id'])
                        logger.info(f"引导模式使用默认模板: {default_template.get('name', '默认模板')}")
                    except Exception as e:
                        logger.warning(f"默认模板渲染失败: {e}")
                        try:
                            word_bytes = generator.generate_bytes(tailored_resume)
                        except Exception as e2:
                            logger.error(f"降级生成器也失败: {e2}", exc_info=True)
                            raise
                else:
                    try:
                        word_bytes = generator.generate_bytes(tailored_resume)
                    except Exception as e2:
                        logger.error(f"无模板生成失败: {e2}", exc_info=True)
                        raise

            task_status[session_id]['progress'] = 100
            task_status[session_id]['status'] = 'completed'

            processing_time_ms = int((time.time() - start_time) * 1000)

            # 计算覆盖率
            if isinstance(evidence_report, dict):
                coverage = evidence_report.get('coverage', 0)
            else:
                coverage = evidence_report.coverage if hasattr(evidence_report, 'coverage') else 0

            # 获取模板名称
            template_name = None
            if used_template_id:
                tmpl = template_manager.get_template(used_template_id)
                if tmpl:
                    template_name = tmpl.get('name')

            result = {
                'session_id': session_id,
                'status': 'completed',
                'processing_time': processing_time_ms,
                'tailored_resume': tailored_resume,
                'tailored_word': base64.b64encode(word_bytes).decode('utf-8'),
                'evidence_report': evidence_report if isinstance(evidence_report, dict) else evidence_report.to_dict(),
                'validation_result': 'pass' if coverage >= 0.9 else 'pass_with_review',
                'style_preserved': style_preserved,
                'template_id': used_template_id,
                'template_used': used_template_id is not None,
                'template_name': template_name,
                'template_mode': template_mode,
                'analysis': analysis_data,
                'pipeline_version': 'v2' if is_v2 else 'v1'
            }

            # 添加匹配度日志
            logger.info(f"📊 匹配度数据: score={analysis_data.get('match_score')}, level={analysis_data.get('match_level')}")
            logger.info(f"📊 完整 analysis: {analysis_data}")

            # V2 版本添加额外信息
            if is_v2:
                result['quality_score'] = pipeline_result.get('quality_score', 0)
                result['optimization_summary'] = pipeline_result.get('optimization_summary', {})
            if pipeline_result.get('review_data'):
                result['review_data'] = pipeline_result['review_data']

            cache_manager.set(resume_text, jd_content, result)
            save_tailored_file(word_bytes, session_id)

            job_title, company = parse_jd_info(jd_content)
            # 尝试从简历文本中提取姓名
            candidate_name = ''
            for line in resume_text.split('\n')[:10]:
                if '姓名' in line or '名字' in line:
                    parts = line.replace('姓名', '').replace('名字', '').replace('：', ':').split(':')
                    if len(parts) > 1:
                        candidate_name = parts[-1].strip()
                        break

            db.save_history(session_id, {
                'candidate_name': candidate_name,
                'candidate_type': 'experienced',
                'job_title': job_title,
                'company': company,
                'match_score': analysis_data.get('match_score', 0),
                'match_level': analysis_data.get('match_level', ''),
                'original_resume': resume_text,
                'tailored_resume': tailored_resume,
                'jd_content': jd_content,
                'evidence_report': evidence_report if isinstance(evidence_report, dict) else evidence_report.to_dict(),
                'optimization_summary': {'input_mode': 'text', 'pipeline_version': 'v2' if is_v2 else 'v1', **pipeline_result.get('optimization_summary', {})},
                'model_used': model_manager.current_model,
                'tokens_used': 0,
                'processing_time_ms': processing_time_ms
            })

            # 扣减用户配额
            if current_user:
                use_quota(current_user['id'], session_id=session_id)

            return jsonify(result)

        except Exception as e:
            import traceback as tb_module
            full_tb = tb_module.format_exc()
            logger.error(f"文本定制失败: {e}\n{full_tb}")
            error_str = str(e)
            error_type = type(e).__name__
            # 更详细的错误分类
            if "429" in error_str or "rate" in error_str.lower() or "限制" in error_str or "并发" in error_str:
                return jsonify({'error': 'API调用频率超限，请等待1-2分钟后重试'}), 429
            elif "resume_analysis" in error_str or "{" in error_str:
                return jsonify({'error': 'AI响应解析失败，请重试'}), 500
            elif "timeout" in error_str.lower():
                return jsonify({'error': 'AI处理超时，请稍后重试'}), 504
            elif "connection" in error_str.lower():
                return jsonify({'error': 'AI服务连接失败，请检查网络'}), 503
            logger.error(f"文本定制详细错误类型: {error_type}, 消息: {error_str}")
            return jsonify({
                'error': '处理失败，请稍后重试',
                'success': False,
            }), 500

    @app.route('/api/tailor/form', methods=['POST'])
    def tailor_form():
        try:
            # 配额检查
            err_resp, current_user = _check_quota_or_anonymous()
            if err_resp:
                return err_resp

            start_time = time.time()

            data = request.get_json()
            if not data:
                return jsonify({'error': '未提供表单数据'}), 400

            jd_content = data.get('jd', '')
            if not jd_content:
                return jsonify({'error': '未提供职位JD'}), 400

            session_id = str(uuid.uuid4())
            _cleanup_task_status(); task_status[session_id] = {'status': 'processing', 'progress': 0, 'message': '正在构建简历...', '_created_at': time.time()}

            resume_text = builder.build_from_form(data)

            task_status[session_id]['progress'] = 20
            task_status[session_id]['message'] = '正在分析JD需求...'

            # 检查是否强制跳过缓存（与 tailor_file/tailor_text 保持一致）
            no_cache = data.get('no_cache', True)

            cached_result = cache_manager.get(resume_text, jd_content)
            # 检查缓存是否包含必要字段（旧缓存可能缺少 tailored_resume）
            # 只有 no_cache=False 时才使用缓存
            if not no_cache and cached_result and cached_result.get('tailored_resume'):
                logger.info(f"命中缓存且包含 tailored_resume: session={session_id}")
                task_status[session_id]['progress'] = 100
                task_status[session_id]['status'] = 'completed'
                return jsonify(cached_result)

            task_status[session_id]['message'] = '正在AI分析...'

            # 使用统一的定制流程（带进度回调）
            pipeline_result = run_tailor_pipeline(resume_text, jd_content, session_id)
            tailored_resume = pipeline_result['tailored_resume']

            # 格式转换：将 AI 返回的嵌套格式转为扁平格式（双重保障）
            tailored_resume = convert_tailored_format(tailored_resume)

            analysis_data = pipeline_result['analysis']
            is_v2 = pipeline_result['is_v2']

            task_status[session_id]['progress'] = 96
            task_status[session_id]['message'] = '正在验证依据...'

            # V2 版本已有依据报告，V1 版本需要单独验证
            if is_v2:
                evidence_report = pipeline_result['evidence_report']
            else:
                evidence_tracker = EvidenceTracker(model_manager)
                evidence_report = evidence_tracker.validate_resume(
                    resume_text,
                    tailored_resume
                ).to_dict()

            task_status[session_id]['progress'] = 98
            task_status[session_id]['message'] = '正在生成文档...'

            word_bytes = generator.generate_bytes(tailored_resume)

            task_status[session_id]['progress'] = 100
            task_status[session_id]['status'] = 'completed'

            processing_time_ms = int((time.time() - start_time) * 1000)

            # 计算覆盖率
            if isinstance(evidence_report, dict):
                coverage = evidence_report.get('coverage', 0)
            else:
                coverage = evidence_report.coverage if hasattr(evidence_report, 'coverage') else 0

            result = {
                'session_id': session_id,
                'status': 'completed',
                'processing_time': processing_time_ms,
                'tailored_resume': tailored_resume,
                'tailored_word': base64.b64encode(word_bytes).decode('utf-8'),
                'evidence_report': evidence_report if isinstance(evidence_report, dict) else evidence_report.to_dict(),
                'validation_result': 'pass' if coverage >= 0.9 else 'pass_with_review',
                'analysis': analysis_data,
                'pipeline_version': 'v2' if is_v2 else 'v1'
            }

            # 添加匹配度日志
            logger.info(f"📊 匹配度数据: score={analysis_data.get('match_score')}, level={analysis_data.get('match_level')}")
            logger.info(f"📊 完整 analysis: {analysis_data}")

            # V2 版本添加额外信息
            if is_v2:
                result['quality_score'] = pipeline_result.get('quality_score', 0)
                result['optimization_summary'] = pipeline_result.get('optimization_summary', {})
            if pipeline_result.get('review_data'):
                result['review_data'] = pipeline_result['review_data']

            cache_manager.set(resume_text, jd_content, result)
            save_tailored_file(word_bytes, session_id)

            job_title, company = parse_jd_info(jd_content)
            basic_info = builder.build_structured(data).get('basic_info', {})
            has_experience = bool(data.get('work_experience') or data.get('work_count', 0) > 0)
            db.save_history(session_id, {
                'candidate_name': basic_info.get('name', ''),
                'candidate_type': 'experienced' if has_experience else 'entry_level',
                'job_title': job_title,
                'company': company,
                'match_score': analysis_data.get('match_score', 0),
                'match_level': analysis_data.get('match_level', ''),
                'original_resume': resume_text,
                'tailored_resume': tailored_resume,
                'jd_content': jd_content,
                'evidence_report': evidence_report if isinstance(evidence_report, dict) else evidence_report.to_dict(),
                'optimization_summary': {'input_mode': 'guided', 'pipeline_version': 'v2' if is_v2 else 'v1', **pipeline_result.get('optimization_summary', {})},
                'model_used': model_manager.current_model,
                'tokens_used': 0,
                'processing_time_ms': processing_time_ms
            })

            # 扣减用户配额
            if current_user:
                use_quota(current_user['id'], session_id=session_id)

            return jsonify(result)

        except Exception as e:
            import traceback as tb_module
            full_tb = tb_module.format_exc()
            logger.error(f"表单定制失败: {e}\n{full_tb}")
            error_str = str(e)
            error_type = type(e).__name__
            if "resume_analysis" in error_str or "{" in error_str:
                return jsonify({'error': 'AI响应解析失败，请重试'}), 500
            return jsonify({
                'error': '处理失败，请稍后重试',
                'success': False,
            }), 500

    @app.route('/api/status/<task_id>', methods=['GET'])
    def get_status(task_id: str):
        if task_id not in task_status:
            return jsonify({'error': '任务不存在'}), 404
        return jsonify(task_status[task_id])

    @app.route('/api/preview', methods=['POST'])
    def preview():
        try:
            data = request.get_json()
            if not data:
                return jsonify({'error': '未提供数据'}), 400
            resume_text = builder.build_from_form(data)
            return jsonify({'preview_text': resume_text})
        except Exception as e:
            logger.error(f"预览失败: {e}")
            return jsonify({'error': str(e)}), 500

    @app.route('/api/stats', methods=['GET'])
    def get_stats():
        return jsonify({
            'model_stats': expert_team.get_stats(),
            'cache_stats': cache_manager.get_stats(),
            'parser_stats': parser.get_stats(),
            'template_stats': template_manager.get_stats()
        })

    # ==================== 模板管理 API ====================

    @app.route('/api/templates', methods=['GET'])
    def get_templates():
        """获取模板列表"""
        source = request.args.get('source')
        templates = template_manager.get_templates(source=source)
        return jsonify({
            'templates': templates,
            'stats': template_manager.get_stats()
        })

    @app.route('/api/templates', methods=['POST'])
    def upload_template():
        """上传新模板"""
        try:
            if 'template' not in request.files:
                return jsonify({'error': '未上传模板文件'}), 400

            template_file = request.files['template']
            if template_file.filename == '':
                return jsonify({'error': '未选择模板文件'}), 400

            ext = Path(template_file.filename).suffix.lower()
            if ext not in ['.docx']:
                return jsonify({'error': '仅支持 .docx 格式'}), 400

            name = request.form.get('name', '')
            description = request.form.get('description', '')
            tags_str = request.form.get('tags', '')
            tags = [t.strip() for t in tags_str.split(',') if t.strip()] if tags_str else []

            file_content = template_file.read()
            template_id, error = template_manager.upload_template(
                file_content, template_file.filename, name, description, tags
            )

            if error and not template_id:
                return jsonify({'error': error}), 400

            return jsonify({
                'template_id': template_id,
                'message': error if error else '模板上传成功'
            })

        except Exception as e:
            logger.error(f"上传模板失败: {e}", exc_info=True)
            return jsonify({'error': str(e)}), 500

    @app.route('/api/templates/<template_id>', methods=['GET'])
    def get_template(template_id: str):
        """获取模板详情"""
        template = template_manager.get_template(template_id)
        if not template:
            return jsonify({'error': '模板不存在'}), 404
        return jsonify(template)

    @app.route('/api/templates/<template_id>', methods=['DELETE'])
    def delete_template(template_id: str):
        """删除模板"""
        success, error = template_manager.delete_template(template_id)
        if not success:
            return jsonify({'error': error}), 400
        return jsonify({'message': '模板已删除'})

    @app.route('/api/templates/<template_id>/set_default', methods=['POST'])
    def set_default_template(template_id: str):
        """设置默认模板"""
        success = template_manager.set_default_template(template_id)
        if not success:
            return jsonify({'error': '设置失败'}), 400
        return jsonify({'message': '已设为默认模板'})

    @app.route('/api/templates/extract', methods=['POST'])
    def extract_template():
        """从简历提取模板"""
        try:
            if 'resume' not in request.files:
                return jsonify({'error': '未上传简历文件'}), 400

            resume_file = request.files['resume']
            if resume_file.filename == '':
                return jsonify({'error': '未选择简历文件'}), 400

            ext = Path(resume_file.filename).suffix.lower()
            if ext not in ['.docx']:
                return jsonify({'error': '仅支持 .docx 格式的简历'}), 400

            name = request.form.get('name', '')
            logger.info(f"[提取模板] 收到 name={repr(name)}, file={resume_file.filename}")

            file_content = resume_file.read()
            template_id, error = template_manager.extract_template_from_resume(
                file_content, resume_file.filename, name
            )

            if error and not template_id:
                return jsonify({'error': error}), 400

            return jsonify({
                'template_id': template_id,
                'message': error if error else '模板提取成功'
            })

        except Exception as e:
            logger.error(f"提取模板失败: {e}", exc_info=True)
            return jsonify({'error': str(e)}), 500

    @app.route('/api/templates/<template_id>/preview', methods=['GET'])
    @limiter.exempt
    def get_template_preview(template_id: str):
        """获取模板预览图"""
        import os

        template = template_manager.get_template(template_id)
        if not template:
            return jsonify({'error': '模板不存在'}), 404

        preview_path = template.get('preview_image', '')
        if preview_path:
            # 转换为绝对路径
            if not os.path.isabs(preview_path):
                preview_path = os.path.join(str(config.BASE_DIR), preview_path)

            # 标准化路径
            preview_path = os.path.normpath(preview_path)

            if os.path.exists(preview_path):
                return send_file(preview_path, mimetype='image/png')

        return jsonify({'error': '预览图不存在'}), 404

    @app.route('/api/templates/<template_id>/download', methods=['GET'])
    def download_template(template_id: str):
        """下载模板文件"""
        template = template_manager.get_template(template_id)
        if not template:
            return jsonify({'error': '模板不存在'}), 404

        file_content = template_manager.get_template_file(template_id)
        if not file_content:
            return jsonify({'error': '模板文件不存在'}), 404

        return send_file(
            io.BytesIO(file_content),
            mimetype='application/vnd.openxmlformats-officedocument.wordprocessingml.document',
            as_attachment=True,
            download_name=f"{template['name']}.docx"
        )

    @app.route('/api/templates/<template_id>/editable', methods=['GET'])
    def get_template_editable(template_id: str):
        """获取模板的可编辑文本内容（非 Jinja2 变量部分）"""
        template = template_manager.get_template(template_id)
        if not template:
            return jsonify({'error': '模板不存在'}), 404

        file_path = template.get('file_path', '')
        if not file_path or not Path(file_path).exists():
            return jsonify({'error': '模板文件不存在'}), 404

        try:
            from docx import Document as DocxDocument
            doc = DocxDocument(file_path)
            paragraphs = []
            for i, para in enumerate(doc.paragraphs):
                text = para.text.strip()
                if not text:
                    paragraphs.append({'index': i, 'text': '', 'editable': False})
                    continue
                # 判断是否包含 Jinja2 变量
                has_jinja = '{%' in text or '{{' in text
                paragraphs.append({
                    'index': i,
                    'text': text,
                    'editable': not has_jinja,  # 纯文本可编辑，Jinja2 标签不可编辑
                    'is_jinja': has_jinja
                })

            return jsonify({
                'template_id': template_id,
                'template_name': template.get('name', ''),
                'paragraphs': paragraphs
            })
        except Exception as e:
            logger.error(f"获取模板可编辑内容失败: {e}", exc_info=True)
            return jsonify({'error': str(e)}), 500

    @app.route('/api/templates/<template_id>/save-edited', methods=['POST'])
    def save_edited_template(template_id: str):
        """保存用户编辑后的模板为新模板"""
        try:
            data = request.get_json()
            if not data:
                return jsonify({'error': '未提供数据'}), 400

            edited_paragraphs = data.get('paragraphs', [])  # [{index, text}]
            new_name = data.get('name', '')

            if not edited_paragraphs:
                return jsonify({'error': '未提供编辑内容'}), 400

            template = template_manager.get_template(template_id)
            if not template:
                return jsonify({'error': '源模板不存在'}), 404

            file_path = template.get('file_path', '')
            if not file_path or not Path(file_path).exists():
                return jsonify({'error': '源模板文件不存在'}), 404

            from docx import Document as DocxDocument
            doc = DocxDocument(file_path)

            # 应用编辑
            edits_map = {p['index']: p['text'] for p in edited_paragraphs if 'index' in p}
            for i, para in enumerate(doc.paragraphs):
                if i in edits_map:
                    # 保留原有格式的 run 结构，只替换文本
                    if para.runs:
                        # 将文本分配给第一个 run，清空其余
                        para.runs[0].text = edits_map[i]
                        for run in para.runs[1:]:
                            run.text = ''
                    else:
                        para.text = edits_map[i]

            # 保存为新模板
            new_id = str(uuid.uuid4())[:12]
            new_path = str(Path(file_path).parent / f"{new_id}.docx")
            doc.save(new_path)

            # 注册为用户自定义模板（直接写入数据库）
            db.save_template({
                'template_id': new_id,
                'name': new_name or f"{template.get('name', '模板')}_编辑版",
                'source': 'user_edited',
                'file_path': new_path,
                'content_hash': '',
                'structure_confidence': template.get('structure_confidence', 0.5),
                'sections': template.get('sections', []),
                'variables': template.get('variables', []),
                'description': f'基于 {template.get("name", "")} 编辑',
                'tags': [],
                'preview_image': '',
                'is_default': False,
                'use_count': 0
            })

            logger.info(f"模板编辑保存成功: {template_id} -> {new_id} ({new_name})")

            return jsonify({
                'success': True,
                'template_id': new_id,
                'name': new_name or f"{template.get('name', '模板')}_编辑版"
            })

        except Exception as e:
            logger.error(f"保存编辑模板失败: {e}", exc_info=True)
            return jsonify({'error': str(e)}), 500

    @app.route('/api/templates/<template_id>/preview/html', methods=['GET'])
    @limiter.exempt
    def get_template_html_preview(template_id: str):
        """获取模板的 HTML 预览（使用示例数据渲染）"""
        import mammoth

        logger.debug(f"预览请求 - template_id: {template_id}")

        template = template_manager.get_template(template_id)
        logger.debug(f"模板信息: name={template.get('name') if template else 'N/A'}, "
                   f"source={template.get('source') if template else 'N/A'}, "
                   f"file_path={template.get('file_path') if template else 'N/A'}")

        if not template:
            logger.error(f"❌ 模板不存在: {template_id}")
            return jsonify({'error': '模板不存在', 'success': False}), 404

        file_path = template.get('file_path', '')
        if not file_path:
            logger.error(f"❌ 模板文件路径不存在: {template_id}")
            return jsonify({'error': '模板文件路径不存在', 'success': False}), 404

        try:
            # 所有模板统一用示例数据渲染，避免显示原始 Jinja2 标签
            sample_data = get_sample_resume_data()
            logger.info(f"📝 使用示例数据渲染模板预览: {template_id}")

            try:
                logger.info(f"🎨 尝试渲染模板: {template_id}")
                word_bytes = template_processor.render(template_id, sample_data)
                logger.info(f"✅ 模板渲染成功，字节数: {len(word_bytes)}")
            except Exception as render_error:
                logger.warning(f"❌ 模板渲染失败: {render_error}，降级到原始文件")
                with open(file_path, 'rb') as f:
                    word_bytes = f.read()

            # 3. 转换为 HTML（带错误处理）
            try:
                result = mammoth.convert_to_html(io.BytesIO(word_bytes))
                html = result.value
                if result.messages:
                    logger.warning(f"mammoth 转换警告: {result.messages}")
            except Exception as mammoth_error:
                logger.error(f"❌ mammoth 转换失败: {mammoth_error}", exc_info=True)
                # 降级：返回简单的纯文本预览
                html = f"<div style='padding:20px;color:#666;'><h3>预览生成失败</h3><p>错误信息: {str(mammoth_error)}</p><p>模板 ID: {template_id}</p><p>建议：请尝试下载模板后使用 Word 查看。</p></div>"

            # 4. 返回带样式的 HTML
            styled_html = f"""
            <style>
                body {{ font-family: 'Microsoft YaHei', Arial, sans-serif; padding: 20px; line-height: 1.6; color: #333; }}
                h1, h2, h3 {{ color: #2c5282; margin-top: 16px; margin-bottom: 8px; }}
                h1 {{ font-size: 24px; border-bottom: 2px solid #2c5282; padding-bottom: 8px; }}
                h2 {{ font-size: 18px; }}
                p {{ margin: 8px 0; }}
                table {{ border-collapse: collapse; width: 100%; margin: 12px 0; }}
                td, th {{ border: 1px solid #ddd; padding: 8px; text-align: left; }}
                ul, ol {{ margin: 8px 0; padding-left: 24px; }}
                li {{ margin: 4px 0; }}
                strong {{ color: #1a365d; }}
            </style>
            {html}
            """
            return jsonify({'html': styled_html, 'success': True})

        except FileNotFoundError:
            return jsonify({'error': '模板文件未找到', 'success': False}), 404
        except Exception as e:
            logger.error(f"生成 HTML 预览失败: {e}", exc_info=True)
            return jsonify({'error': str(e), 'success': False}), 500

    @app.route('/api/preview/tailored', methods=['POST'])
    def preview_tailored_resume():
        """将定制后的简历（base64 docx）转换为 HTML 预览"""
        logger.info("🔄 /api/preview/tailored endpoint called")
        import mammoth
        import binascii

        try:
            data = request.get_json()
            if not data:
                logger.error("❌ 请求体为空")
                return jsonify({'error': '请求数据为空', 'success': False}), 400

            if 'tailored_word' not in data:
                logger.error("❌ 缺少 tailored_word 参数")
                return jsonify({'error': '未提供简历数据', 'success': False}), 400

            b64_data = data['tailored_word']
            if not b64_data or not b64_data.strip():
                logger.error("❌ tailored_word 为空")
                return jsonify({'error': '简历数据为空', 'success': False}), 400

            logger.info(f"📊 base64 数据长度: {len(b64_data)}")

            # 解码 base64
            try:
                word_bytes = base64.b64decode(b64_data)
                logger.info(f"✅ 解码成功，字节数: {len(word_bytes)}")
            except binascii.Error as e:
                logger.error(f"❌ base64 解码失败: {e}")
                return jsonify({'error': f'简历数据格式错误: {str(e)}', 'success': False}), 400

            # 验证是否为有效的 docx 文件（ZIP 格式）
            if len(word_bytes) < 8 or word_bytes[:4] != b'PK\x03\x04':
                logger.error(f"❌ 无效的 Word 文档格式，文件头: {word_bytes[:8].hex() if len(word_bytes) >= 8 else 'too short'}")
                return jsonify({'error': '无效的 Word 文档格式', 'success': False}), 400

            # mammoth 转换 - 添加 style_map 选项处理 Word 样式
            try:
                result = mammoth.convert_to_html(
                    io.BytesIO(word_bytes),
                    style_map="""
                        p[style-name='Heading 1'] => h1:fresh
                        p[style-name='Heading 2'] => h2:fresh
                        p[style-name='Heading 3'] => h3:fresh
                        p[style-name='标题 1'] => h1:fresh
                        p[style-name='标题 2'] => h2:fresh
                        p[style-name='标题 3'] => h3:fresh
                    """
                )
                html = result.value

                if result.messages:
                    for msg in result.messages:
                        logger.warning(f"⚠️ mammoth 警告: {msg}")

                logger.info(f"✅ HTML 生成成功，长度: {len(html)}")

                # 检查转换结果是否有效
                if not html or len(html.strip()) < 10:
                    logger.warning("⚠️ HTML 内容过短，可能转换失败")
                    raise ValueError("转换结果为空或内容过短")

            except Exception as mammoth_error:
                logger.error(f"❌ mammoth 转换失败: {type(mammoth_error).__name__}: {mammoth_error}", exc_info=True)

                # 降级方案：返回友好的提示（返回 success: True 让前端显示降级内容）
                fallback_html = """
                <div style="padding: 40px; text-align: center; background: #f8f9fa; border-radius: 8px; margin: 20px;">
                    <div style="font-size: 48px; margin-bottom: 16px;">📄</div>
                    <h3 style="color: #495057; margin-bottom: 8px;">预览暂时不可用</h3>
                    <p style="color: #6c757d; margin-bottom: 16px;">简历文件已生成，请点击上方"下载简历"按钮查看完整内容</p>
                    <p style="color: #adb5bd; font-size: 12px;">技术提示: 文档格式转换遇到问题</p>
                </div>
                """
                return jsonify({
                    'html': fallback_html,
                    'success': True,  # 返回 True，让前端显示降级内容
                    'preview_failed': True
                })

            # 返回带样式的 HTML
            styled_html = f"""
            <style>
                body {{ font-family: 'Microsoft YaHei', Arial, sans-serif; padding: 20px; line-height: 1.6; color: #333; }}
                h1, h2, h3 {{ color: #2c5282; margin-top: 16px; margin-bottom: 8px; }}
                h1 {{ font-size: 24px; border-bottom: 2px solid #2c5282; padding-bottom: 8px; }}
                h2 {{ font-size: 18px; }}
                p {{ margin: 8px 0; }}
                table {{ border-collapse: collapse; width: 100%; margin: 12px 0; }}
                td, th {{ border: 1px solid #ddd; padding: 8px; text-align: left; }}
                ul, ol {{ margin: 8px 0; padding-left: 24px; }}
                li {{ margin: 4px 0; }}
                strong {{ color: #1a365d; }}
            </style>
            {html}
            """

            return jsonify({'html': styled_html, 'success': True})

        except Exception as e:
            logger.error(f"❌ 预览 API 异常: {e}", exc_info=True)
            # 异常时也返回降级内容，而不是 500 错误
            return jsonify({
                'html': '<div style="padding:20px;text-align:center;color:#666;">预览加载失败，请下载查看</div>',
                'success': True,
                'preview_failed': True
            })

    @app.route('/api/templates/<template_id>/compatibility', methods=['POST'])
    def check_template_compatibility(template_id: str):
        """检查模板兼容性"""
        try:
            data = request.get_json()
            if not data:
                return jsonify({'error': '未提供简历数据'}), 400

            resume_data = data.get('resume_data', {})
            is_compatible, missing_sections = template_manager.check_compatibility(
                template_id, resume_data
            )

            return jsonify({
                'is_compatible': is_compatible,
                'missing_sections': missing_sections
            })

        except Exception as e:
            logger.error(f"检查兼容性失败: {e}")
            return jsonify({'error': str(e)}), 500

    @app.route('/api/resume/update', methods=['POST'])
    def resume_update():
        """用户编辑简历后，重新渲染 DOCX 并更新历史记录"""
        try:
            data = request.get_json()
            if not data:
                return jsonify({'error': '未提供数据'}), 400

            session_id = data.get('session_id', '')
            tailored_resume = data.get('tailored_resume')
            template_id = data.get('template_id')

            if not session_id:
                return jsonify({'error': '缺少 session_id'}), 400
            if not tailored_resume:
                return jsonify({'error': '缺少 tailored_resume'}), 400

            # 查找原始历史记录（获取 original_resume 和 jd_content）
            history = db.get_history(session_id)
            if not history:
                return jsonify({'error': '未找到该会话记录'}), 404

            # 重新渲染 DOCX
            if template_id:
                try:
                    word_bytes = template_processor.render(template_id, tailored_resume, None)
                except Exception as e:
                    logger.error(f"简历编辑后模板渲染失败: {e}", exc_info=True)
                    return jsonify({'error': f'模板渲染失败: {str(e)}'}), 500
            else:
                # 无模板时使用默认模板或生成器
                default_template = template_manager.get_default_template()
                if default_template:
                    try:
                        word_bytes = template_processor.render(
                            default_template['template_id'], tailored_resume, None
                        )
                    except Exception as e:
                        logger.warning(f"简历编辑后默认模板失败，降级到生成器: {e}")
                        word_bytes = generator.generate_bytes(tailored_resume)
                else:
                    word_bytes = generator.generate_bytes(tailored_resume)

            # 更新历史记录
            db.update_history(session_id, {
                'tailored_resume': tailored_resume,
                'optimization_summary': {
                    **(history.get('optimization_summary') or {}),
                    'user_edited': True,
                    'edited_at': datetime.now().isoformat()
                }
            })

            # 更新磁盘上的 DOCX 文件
            save_tailored_file(word_bytes, session_id)

            logger.info(f"简历编辑保存成功: session={session_id}")

            return jsonify({
                'success': True,
                'session_id': session_id,
                'tailored_word': base64.b64encode(word_bytes).decode('utf-8'),
            })

        except Exception as e:
            logger.error(f"简历编辑保存失败: {e}", exc_info=True)
            return jsonify({'error': str(e)}), 500

    @app.route('/api/templates/recommend', methods=['POST'])
    def recommend_template():
        """根据 JD 推荐模板"""
        try:
            data = request.get_json() or {}
            jd_content = data.get('jd_content', '')
            industry = data.get('industry', '')
            position_level = data.get('position_level', '')

            recommendations = template_manager.recommend_template(
                jd_content=jd_content,
                industry=industry,
                position_level=position_level
            )

            return jsonify({
                'recommendations': recommendations[:5]  # 返回前5个推荐
            })

        except Exception as e:
            logger.error(f"模板推荐失败: {e}")
            return jsonify({'error': str(e)}), 500

    return app


if __name__ == '__main__':
    app = create_app()
    app.run(host='0.0.0.0', port=config.SIMPLE_APP_PORT, debug=True, use_reloader=False)
