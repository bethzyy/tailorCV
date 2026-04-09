"""
Jinja2 标记插入器模块

根据结构检测的结果，在 Word 文档中插入 Jinja2 模板标记，
生成可用于 docxtpl 渲染的模板文档。
"""

import re
import logging
from typing import Dict, List, Optional, Tuple, Any
from dataclasses import dataclass, field
from copy import deepcopy
import io

# Word 文档处理
try:
    from docx import Document
    from docx.text.paragraph import Paragraph
    from docx.text.run import Run
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement
    HAS_PYTHON_DOCX = True
except ImportError:
    HAS_PYTHON_DOCX = False

from .structure_detector import (
    StructureMap, SectionInfo, EntryInfo, SectionType
)

logger = logging.getLogger(__name__)


@dataclass
class TemplateMetadata:
    """模板元数据"""
    template_id: str                           # 模板ID（用于存储和检索）
    original_filename: str                     # 原始文件名
    structure_confidence: float                # 结构检测置信度
    variables: List[str] = field(default_factory=list)  # 模板变量列表
    has_dynamic_content: bool = False          # 是否包含动态内容（循环）
    sections_detected: int = 0                 # 检测到的章节数
    entries_detected: int = 0                  # 检测到的条目数


class JinjaTagInserter:
    """
    Jinja2 标记插入器

    将普通 Word 文档转换为 Jinja2 模板，支持：
    1. 静态变量：{{ basic_info.name }}
    2. 动态循环：{%tr for exp in work_experience %} ... {%tr endfor %}
    3. 条件判断：{%p if summary %} ... {%p endif %}
    """

    # 变量名映射
    VARIABLE_MAPPING = {
        SectionType.BASIC_INFO: {
            'name': 'basic_info.name',
            'phone': 'basic_info.phone',
            'email': 'basic_info.email',
            'location': 'basic_info.location',
            'age': 'basic_info.age',
            'gender': 'basic_info.gender',
        },
        SectionType.SUMMARY: 'summary',
        SectionType.SELF_EVALUATION: 'self_evaluation',
    }

    # 循环变量映射
    LOOP_VARIABLES = {
        SectionType.EDUCATION: {
            'loop_var': 'edu',
            'list_var': 'education',
            'fields': ['time', 'school', 'major', 'degree', 'tailored', 'highlights']
        },
        SectionType.WORK: {
            'loop_var': 'exp',
            'list_var': 'work_experience',
            'fields': ['time', 'company', 'position', 'tailored', 'content']
        },
        SectionType.PROJECT: {
            'loop_var': 'proj',
            'list_var': 'projects',
            'fields': ['time', 'name', 'role', 'tailored', 'content']
        },
        SectionType.SKILLS: {
            'loop_var': 'skill',
            'list_var': 'skills',
            'fields': ['name', 'tailored_description']
        },
        SectionType.AWARDS: {
            'loop_var': 'award',
            'list_var': 'awards',
            'fields': ['name']
        },
        SectionType.CERTIFICATES: {
            'loop_var': 'cert',
            'list_var': 'certificates',
            'fields': ['name']
        },
    }

    def __init__(self):
        self.insertion_stats = {
            'static_variables': 0,
            'dynamic_loops': 0,
            'conditional_blocks': 0,
            'failed': 0
        }

    def insert_tags(self, doc: 'Document', structure: StructureMap,
                   template_id: str) -> Tuple['Document', TemplateMetadata]:
        """
        在文档中插入 Jinja2 标记

        Args:
            doc: python-docx Document 对象
            structure: 结构映射
            template_id: 模板ID

        Returns:
            Tuple[Document, TemplateMetadata]: (标记后的文档, 模板元数据)
        """
        if not HAS_PYTHON_DOCX:
            raise ImportError("未安装 python-docx")

        # 创建文档副本（通过保存和重新加载，确保深拷贝）
        # deepcopy 对 python-docx 对象不完全有效
        bio = io.BytesIO()
        doc.save(bio)
        bio.seek(0)
        tagged_doc = Document(bio)

        variables = []
        has_dynamic = False

        # 1. 处理姓名
        if structure.name_paragraph_index is not None:
            self._insert_name_tag(tagged_doc, structure.name_paragraph_index)
            variables.append('basic_info.name')
            self.insertion_stats['static_variables'] += 1

        # 2. 处理联系方式
        if structure.contact_paragraph_index is not None:
            self._insert_contact_tag(tagged_doc, structure.contact_paragraph_index)
            variables.extend(['basic_info.phone', 'basic_info.email', 'basic_info.location'])
            self.insertion_stats['static_variables'] += 1

        # 3. 处理各章节
        for section in structure.sections:
            if section.is_dynamic:
                # 动态章节（循环）
                section_vars = self._insert_dynamic_section(
                    tagged_doc, section, structure.entries
                )
                variables.extend(section_vars)
                has_dynamic = True
                self.insertion_stats['dynamic_loops'] += 1
            else:
                # 静态章节
                section_var = self._insert_static_section(tagged_doc, section)
                if section_var:
                    variables.append(section_var)
                    self.insertion_stats['static_variables'] += 1

        # 创建元数据
        metadata = TemplateMetadata(
            template_id=template_id,
            original_filename='',  # 由调用方设置
            structure_confidence=structure.confidence,
            variables=list(set(variables)),
            has_dynamic_content=has_dynamic,
            sections_detected=len(structure.sections),
            entries_detected=len(structure.entries)
        )

        logger.info(f"Jinja2 标记插入完成: {len(variables)} 变量, "
                   f"动态内容: {has_dynamic}")

        return tagged_doc, metadata

    def _insert_name_tag(self, doc: 'Document', para_index: int):
        """
        在姓名段落插入变量标记 - 保留原始格式

        不清除段落，保留原始 run 的格式（字体、字号、加粗等），
        只在原位置插入 Jinja2 变量标记。
        """
        para = doc.paragraphs[para_index]

        # 保留原始文本作为默认值
        original_text = para.text.strip()

        # 保存第一个 run 的格式（通常姓名只有一两个 run）
        if para.runs:
            # 保留第一个 run 的格式，修改其文本
            first_run = para.runs[0]
            # 保存格式属性
            font_name = first_run.font.name
            font_size = first_run.font.size
            font_bold = first_run.font.bold
            font_color = first_run.font.color.rgb if first_run.font.color.rgb else None

            # 清除所有 runs
            para.clear()

            # 添加新的 run，恢复原始格式
            new_run = para.add_run(f"{{{{ basic_info.name | default('{original_text}') }}}}")
            if font_name:
                new_run.font.name = font_name
            if font_size:
                new_run.font.size = font_size
            if font_bold is not None:
                new_run.font.bold = font_bold
            if font_color:
                new_run.font.color.rgb = font_color
        else:
            # 没有 run 的情况下才添加新 run
            run = para.add_run('{{ basic_info.name }}')

        logger.debug(f"插入姓名变量 (段落 {para_index})")

    def _insert_contact_tag(self, doc: 'Document', para_index: int):
        """
        在联系方式段落插入变量标记 - 保留原始格式

        分析原始联系方式的格式，保留分隔符和布局，
        只将电话、邮箱、地址替换为 Jinja2 变量。
        """
        para = doc.paragraphs[para_index]
        original_text = para.text.strip()

        # 保存原始格式
        original_runs = []
        for run in para.runs:
            run_info = {
                'text': run.text,
                'font_name': run.font.name,
                'font_size': run.font.size,
                'font_bold': run.font.bold,
                'font_color': run.font.color.rgb if run.font.color.rgb else None
            }
            original_runs.append(run_info)

        # 分析原始文本，提取分隔符
        separator = self._detect_contact_separator(original_text)

        # 使用简单的变量替换语法（不使用条件判断）
        # 因为 docxtpl 在同一段落中的条件语法支持有限
        if separator:
            contact_template = (
                f'电话: {{{{ basic_info.phone }}}} {separator} '
                f'邮箱: {{{{ basic_info.email }}}} {separator} '
                f'现居: {{{{ basic_info.location }}}}'
            )
        else:
            contact_template = (
                '电话: {{ basic_info.phone }} '
                '邮箱: {{ basic_info.email }} '
                '现居: {{ basic_info.location }}'
            )

        # 保留第一个 run 的格式
        if original_runs:
            first_run_info = original_runs[0]
            para.clear()
            new_run = para.add_run(contact_template)
            if first_run_info['font_name']:
                new_run.font.name = first_run_info['font_name']
            if first_run_info['font_size']:
                new_run.font.size = first_run_info['font_size']
            if first_run_info['font_bold'] is not None:
                new_run.font.bold = first_run_info['font_bold']
            if first_run_info['font_color']:
                new_run.font.color.rgb = first_run_info['font_color']
        else:
            para.clear()
            para.add_run(contact_template)

        logger.debug(f"插入联系方式变量 (段落 {para_index})")

    def _detect_contact_separator(self, text: str) -> Optional[str]:
        """
        检测联系方式中的分隔符

        Args:
            text: 原始联系方式文本

        Returns:
            Optional[str]: 分隔符，如果没有则返回 None
        """
        # 常见分隔符优先级（包括全角和半角字符）
        separators = ['|', '｜', '/', '／', '·', '•', '-', '—']

        for sep in separators:
            if sep in text:
                return sep

        return None

    def _insert_static_section(self, doc: 'Document', section: SectionInfo) -> Optional[str]:
        """
        处理静态章节（个人简介、自我评价）

        Args:
            doc: 文档对象
            section: 章节信息

        Returns:
            Optional[str]: 变量名，如果没有处理则返回 None
        """
        var_name = self.VARIABLE_MAPPING.get(section.section_type)
        if not var_name:
            return None

        # 获取章节内容段落（跳过标题）
        content_paras = []
        for i in range(section.content_start, section.content_end + 1):
            if i < len(doc.paragraphs):
                para = doc.paragraphs[i]
                if para.text.strip():
                    content_paras.append(i)

        if not content_paras:
            return None

        # 将第一个内容段落替换为变量
        first_content_idx = content_paras[0]
        para = doc.paragraphs[first_content_idx]

        # 使用简单变量替换（不使用 {%p %} 条件标签，避免 XML runs 拆分导致 docxtpl 解析失败）
        template_text = f'{{{{ {var_name} }}}}'
        self._replace_paragraph_text(para, template_text)

        # 清除其他内容段落
        for idx in content_paras[1:]:
            doc.paragraphs[idx].clear()

        logger.debug(f"插入静态章节变量: {var_name}")
        return var_name

    def _insert_dynamic_section(self, doc: 'Document', section: SectionInfo,
                               entries: List[EntryInfo]) -> List[str]:
        """
        处理动态章节（工作经历、项目经历、教育背景）

        使用简单的变量替换语法，不使用循环（因为 docxtpl 的循环语法在段落中有限制）

        Args:
            doc: 文档对象
            section: 章节信息
            entries: 该章节的条目列表

        Returns:
            List[str]: 使用的变量列表
        """
        loop_config = self.LOOP_VARIABLES.get(section.section_type)
        if not loop_config:
            return []

        loop_var = loop_config['loop_var']
        list_var = loop_config['list_var']
        fields = loop_config['fields']

        # 过滤出该章节的条目
        section_entries = [e for e in entries if e.entry_type == section.section_type]

        if not section_entries:
            # 没有条目，添加一个简单的模板条目
            self._add_simple_template_entry(doc, section, loop_var, list_var)
            return [list_var]

        # 处理每个条目 - 使用简单的变量替换
        for i, entry in enumerate(section_entries):
            self._insert_entry_simple(doc, entry, loop_var, list_var, i)

        return [list_var]

    def _add_simple_template_entry(self, doc: 'Document', section: SectionInfo,
                                   loop_var: str, list_var: str):
        """
        添加简单的模板条目（当没有检测到条目时）
        """
        insert_idx = section.content_start

        # 构建简单的模板文本
        if section.section_type == SectionType.WORK:
            template_text = f'{{{{ {list_var}_0_time }}}}  {{{{ {list_var}_0_company }}}}  |  {{{{ {list_var}_0_position }}}}'
        elif section.section_type == SectionType.PROJECT:
            template_text = f'1. {{{{ {list_var}_0_time }}}}  {{{{ {list_var}_0_name }}}}  |  {{{{ {list_var}_0_role }}}}'
        elif section.section_type == SectionType.EDUCATION:
            template_text = f'{{{{ {list_var}_0_time }}}}  {{{{ {list_var}_0_school }}}}  {{{{ {list_var}_0_major }}}}'
        else:
            template_text = f'{{{{ {list_var}_0_name }}}}'

        if insert_idx < len(doc.paragraphs):
            para = doc.paragraphs[insert_idx]
            self._replace_paragraph_text(para, template_text)

            # 替换下一个内容段落为 tailored 变量
            next_idx = insert_idx + 1
            while next_idx < len(doc.paragraphs):
                next_para = doc.paragraphs[next_idx]
                if next_para.text.strip():
                    self._replace_paragraph_text(
                        next_para,
                        f'{{{{ {list_var}_0_tailored }}}}'
                    )
                    # 清除后续内容段落直到空行或下一个条目
                    clear_paras = []
                    clear_idx = next_idx + 1
                    while clear_idx < len(doc.paragraphs):
                        cp = doc.paragraphs[clear_idx]
                        if not cp.text.strip():
                            break
                        if re.match(r'^\d+\.', cp.text.strip()):
                            break
                        clear_paras.append(cp)
                        clear_idx += 1
                    for cp in clear_paras:
                        parent = cp._element.getparent()
                        if parent is not None:
                            parent.remove(cp._element)
                    logger.debug(f"简单条目插入内容变量 (段落 {next_idx})")
                    break
                next_idx += 1

        logger.debug(f"添加简单模板条目: {list_var}")

    def _insert_entry_simple(self, doc: 'Document', entry: EntryInfo,
                            loop_var: str, list_var: str, index: int):
        """
        插入简单变量替换的条目
        """
        para = doc.paragraphs[entry.paragraph_index]

        # 使用带索引的变量名
        if entry.entry_type == SectionType.WORK:
            template = (
                f'{{{{ {list_var}_{index}_time }}}}  '
                f'{{{{ {list_var}_{index}_company }}}}  |  '
                f'{{{{ {list_var}_{index}_position }}}}'
            )
        elif entry.entry_type == SectionType.PROJECT:
            if entry.time:
                template = (
                    f'{index + 1}. {{{{ {list_var}_{index}_time }}}}  '
                    f'{{{{ {list_var}_{index}_name }}}}  |  '
                    f'{{{{ {list_var}_{index}_role }}}}'
                )
            else:
                # 无时间字段，保留编号 + 名称 + 分隔符 + 角色
                # 从原始段落文本检测分隔符（_parse_entry_header 已拆分，需看原文）
                original_text = para.text
                if '——' in original_text:
                    sep = ' —— '
                elif '—' in original_text:
                    sep = ' — '
                else:
                    sep = ' | '
                template = (
                    f'{index + 1}. {{{{ {list_var}_{index}_name }}}}'
                    f'{sep}'
                    f'{{{{ {list_var}_{index}_role }}}}'
                )
        elif entry.entry_type == SectionType.EDUCATION:
            template = (
                f'{{{{ {list_var}_{index}_time }}}}  '
                f'{{{{ {list_var}_{index}_school }}}}  '
                f'{{{{ {list_var}_{index}_major }}}}'
            )
        else:
            template = f'{{{{ {list_var}_{index}_name }}}}'

        self._replace_paragraph_text(para, template)

        # 处理内容段落 → 合并到第一个，替换为 tailored 变量
        if entry.content_paragraphs:
            first_content_idx = entry.content_paragraphs[0]
            if first_content_idx < len(doc.paragraphs):
                content_para = doc.paragraphs[first_content_idx]
                self._replace_paragraph_text(
                    content_para,
                    f'{{{{ {list_var}_{index}_tailored }}}}'
                )
                logger.debug(f"插入内容段落变量 (段落 {first_content_idx})")
            # 清除剩余内容段落
            # 清除剩余内容段落（从XML中彻底删除，避免空行残留）
            paragraphs_to_remove = []
            for idx in entry.content_paragraphs[1:]:
                if idx < len(doc.paragraphs):
                    paragraphs_to_remove.append(doc.paragraphs[idx])
            for para in paragraphs_to_remove:
                parent = para._element.getparent()
                if parent is not None:
                    parent.remove(para._element)
        else:
            # 兜底：content_paragraphs 为空时，向前扫描找到内容段落
            next_idx = entry.paragraph_index + 1
            while next_idx < len(doc.paragraphs):
                next_para = doc.paragraphs[next_idx]
                if next_para.text.strip():
                    self._replace_paragraph_text(
                        next_para,
                        f'{{{{ {list_var}_{index}_tailored }}}}'
                    )
                    # 清除后续内容段落直到空行或下一个条目
                    clear_paras = []
                    clear_idx = next_idx + 1
                    while clear_idx < len(doc.paragraphs):
                        cp = doc.paragraphs[clear_idx]
                        if not cp.text.strip():
                            break
                        if re.match(r'^\d+\.', cp.text.strip()):
                            break
                        clear_paras.append(cp)
                        clear_idx += 1
                    for cp in clear_paras:
                        parent = cp._element.getparent()
                        if parent is not None:
                            parent.remove(cp._element)
                    logger.debug(f"兜底插入内容段落变量 (段落 {next_idx})")
                    break
                next_idx += 1

        logger.debug(f"插入简单条目变量 (段落 {entry.paragraph_index}, 索引 {index})")

    def _replace_paragraph_text(self, para: 'Paragraph', new_text: str):
        """
        替换段落文本，保留段落格式和第一个 run 的字体格式

        Args:
            para: 段落对象
            new_text: 新文本
        """
        # 保存第一个 run 的格式（如果存在）
        first_run_info = None
        if para.runs:
            first_run = para.runs[0]
            first_run_info = {
                'font_name': first_run.font.name,
                'font_size': first_run.font.size,
                'font_bold': first_run.font.bold,
                'font_color': first_run.font.color.rgb if first_run.font.color.rgb else None,
                'font_italic': first_run.font.italic,
                'font_underline': first_run.font.underline
            }

        # 保存段落格式
        para_alignment = para.paragraph_format.alignment
        para_space_before = para.paragraph_format.space_before
        para_space_after = para.paragraph_format.space_after
        para_line_spacing = para.paragraph_format.line_spacing

        # 清除所有 runs
        para.clear()

        # 添加新的 run
        run = para.add_run(new_text)

        # 恢复第一个 run 的格式
        if first_run_info:
            if first_run_info['font_name']:
                run.font.name = first_run_info['font_name']
            if first_run_info['font_size']:
                run.font.size = first_run_info['font_size']
            if first_run_info['font_bold'] is not None:
                run.font.bold = first_run_info['font_bold']
            if first_run_info['font_color']:
                run.font.color.rgb = first_run_info['font_color']
            if first_run_info['font_italic'] is not None:
                run.font.italic = first_run_info['font_italic']
            if first_run_info['font_underline'] is not None:
                run.font.underline = first_run_info['font_underline']

        # 恢复段落格式
        if para_alignment is not None:
            para.paragraph_format.alignment = para_alignment
        if para_space_before is not None:
            para.paragraph_format.space_before = para_space_before
        if para_space_after is not None:
            para.paragraph_format.space_after = para_space_after
        if para_line_spacing is not None:
            para.paragraph_format.line_spacing = para_line_spacing

        # 如果没有 run 格式，尝试使用段落样式
        if not first_run_info and para.style and para.style.font:
            run.font.name = para.style.font.name
            run.font.size = para.style.font.size
            run.font.bold = para.style.font.bold

    def get_stats(self) -> Dict[str, int]:
        """获取插入统计信息"""
        return self.insertion_stats.copy()
