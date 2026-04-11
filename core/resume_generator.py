"""
简历生成器模块

负责生成定制后的简历文档，支持 Word 格式输出。
"""

import io
import os
import logging
import subprocess
import tempfile
from typing import Dict, Any, Optional, List
from pathlib import Path
from datetime import datetime

from docx import Document
from docx.shared import Pt, Inches, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.enum.style import WD_STYLE_TYPE
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from jinja2 import Environment, FileSystemLoader

from .config import config
from .resume_parser import StyleMetadata

logger = logging.getLogger(__name__)


class WordDocumentBuilder:
    """负责构建和格式化 Word 文档的细节"""

    def __init__(self):
        pass

    def create_document(self, tailored_resume: Dict[str, Any],
                        style_metadata: StyleMetadata,
                        add_content_fn) -> Document:
        """
        创建文档对象并设置基础样式
        Args:
            tailored_resume: 简历数据
            style_metadata: 样式元数据
            add_content_fn: 用于添加内容的回调函数
        """
        doc = Document()

        # 使用传入的样式元数据或默认值
        if style_metadata is None:
            style_metadata = StyleMetadata()

        # 设置页面边距（使用提取的边距）
        for section in doc.sections:
            section.top_margin = Cm(style_metadata.margin_top)
            section.bottom_margin = Cm(style_metadata.margin_bottom)
            section.left_margin = Cm(style_metadata.margin_left)
            section.right_margin = Cm(style_metadata.margin_right)

        self._set_document_font(doc, style_metadata)

        # 调用回调函数添加内容
        add_content_fn(doc, style_metadata)

        return doc

    def _set_document_font(self, doc: Document, style_metadata: StyleMetadata):
        """设置文档默认字体（使用提取的样式）"""
        # 设置正文样式
        style = doc.styles['Normal']
        font = style.font
        font.name = style_metadata.primary_font
        font.size = Pt(style_metadata.body_font_size)

        # 设置中文字体
        style._element.rPr.rFonts.set(qn('w:eastAsia'), style_metadata.primary_font)

    def add_basic_info(self, doc: Document, basic_info: Dict[str, Any],
                       style_metadata: StyleMetadata):
        """添加基本信息（使用动态字号）"""
        name = basic_info.get('name', '')
        if not name:
            return

        # 姓名（标题）- 使用动态字号
        title = doc.add_paragraph()
        title_run = title.add_run(name)
        title_run.bold = True
        title_run.font.size = Pt(style_metadata.get_name_font_size())
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER

        # 联系方式（一行）
        contact_parts = []
        if basic_info.get('phone'):
            contact_parts.append(f"电话: {basic_info['phone']}")
        if basic_info.get('email'):
            contact_parts.append(f"邮箱: {basic_info['email']}")
        if basic_info.get('location'):
            contact_parts.append(f"现居: {basic_info['location']}")
        if basic_info.get('age'):
            age = str(basic_info['age'])
            if not age.endswith('岁'):
                age = f"{age}岁"
            contact_parts.append(f"年龄: {age}")
        if basic_info.get('gender'):
            contact_parts.append(f"性别: {basic_info['gender']}")

        if contact_parts:
            contact = doc.add_paragraph()
            contact_run = contact.add_run(' | '.join(contact_parts))
            contact_run.font.size = Pt(style_metadata.get_time_font_size())
            contact.alignment = WD_ALIGN_PARAGRAPH.CENTER

        # 分隔线
        doc.add_paragraph('─' * 40)

    def add_summary(self, doc: Document, summary: str, style_metadata: StyleMetadata):
        """添加个人简介"""
        if not summary:
            return

        self._add_section_title(doc, '个人简介', style_metadata)
        p = doc.add_paragraph(summary)
        p.paragraph_format.first_line_indent = Cm(0.74)  # 两字符缩进

    def add_education(self, doc: Document, education: List[Dict[str, Any]],
                      style_metadata: StyleMetadata):
        """添加教育背景 - 支持复杂对象格式"""
        if not education:
            return

        self._add_section_title(doc, '教育背景', style_metadata)

        for edu in education:
            p = doc.add_paragraph()

            # 如果有 tailored 字段（优化后内容），使用它
            tailored = edu.get('tailored', '')
            if tailored:
                p.add_run(tailored)
                continue

            # 否则从各字段组装
            # 时间
            time = edu.get('time', '')
            if time:
                run = p.add_run(time + '  ')
                run.font.size = Pt(style_metadata.get_time_font_size())

            # 学校和专业
            school = edu.get('school', '')
            major = edu.get('major', '')
            degree = edu.get('degree', '')

            if school:
                run = p.add_run(school)
                run.bold = True
            if major:
                p.add_run(f'  {major}')
            if degree:
                run = p.add_run(f'  [{degree}]')
                run.font.size = Pt(style_metadata.get_degree_font_size())

            # 如果有 highlights，添加核心课程等信息
            highlights = edu.get('highlights', [])
            if highlights:
                p2 = doc.add_paragraph()
                p2.add_run('核心课程/亮点: ' + ' | '.join(highlights))
                p2.paragraph_format.left_indent = Cm(0.5)

    def add_work_experience(self, doc: Document, work_experience: List[Dict[str, Any]],
                            style_metadata: StyleMetadata):
        """添加工作经历"""
        if not work_experience:
            return

        self._add_section_title(doc, '工作经历', style_metadata)

        for exp in work_experience:
            # 标题行：时间 公司 职位
            p = doc.add_paragraph()

            time = exp.get('time', '')
            if time:
                run = p.add_run(time + '  ')
                run.font.size = Pt(style_metadata.get_time_font_size())

            company = exp.get('company', '')
            if company:
                run = p.add_run(company)
                run.bold = True

            position = exp.get('position', '')
            if position:
                p.add_run(f'  |  {position}')

            # 工作内容 - 优先使用 tailored，支持多格式降级
            content = exp.get('tailored', '')

            # 降级：如果 tailored 为空但有 tailored_bullets
            if not content and 'tailored_bullets' in exp:
                bullets = exp.get('tailored_bullets', [])
                if isinstance(bullets, list):
                    contents = [b.get('content', '') if isinstance(b, dict) else b for b in bullets]
                    content = '\n'.join(filter(None, contents))

            # 最终降级：使用原始 content
            if not content:
                content = exp.get('content', '')

            if content:
                # 按行分割并添加
                lines = content.split('\n')
                for line in lines:
                    line = line.strip()
                    if line:
                        p = doc.add_paragraph(line, style='List Bullet')
                        p.paragraph_format.left_indent = Cm(0.5)

    def add_projects(self, doc: Document, projects: List[Dict[str, Any]],
                     style_metadata: StyleMetadata):
        """添加项目经历"""
        if not projects:
            return

        self._add_section_title(doc, '项目经历', style_metadata)

        for idx, proj in enumerate(projects, 1):
            # 标题行
            p = doc.add_paragraph()

            time = proj.get('time', '')
            if time:
                run = p.add_run(time + '  ')
                run.font.size = Pt(style_metadata.get_time_font_size())

            name = proj.get('name', '')
            if name:
                run = p.add_run(f'{idx}. {name}')
                run.bold = True

            role = proj.get('role', '')
            if role:
                p.add_run(f'  |  {role}')

            # 项目内容 - 优先使用 tailored，支持多格式降级
            content = proj.get('tailored', '')

            # 降级：如果 tailored 为空但有 tailored_description
            if not content and 'tailored_description' in proj:
                content = proj['tailored_description']

            # 最终降级：使用原始 content
            if not content:
                content = proj.get('content', '')

            if content:
                lines = content.split('\n')
                for line in lines:
                    line = line.strip()
                    if line:
                        p = doc.add_paragraph(line, style='List Bullet')
                        p.paragraph_format.left_indent = Cm(0.5)

    def add_skills(self, doc: Document, skills: List[Any],
                   style_metadata: StyleMetadata):
        """添加专业技能"""
        if not skills:
            return

        self._add_section_title(doc, '专业技能', style_metadata)

        # 处理不同的技能格式
        if isinstance(skills, list):
            skill_names = []
            for skill in skills:
                if isinstance(skill, dict):
                    name = skill.get('name', '')
                    desc = skill.get('tailored_description', '')
                    if name:
                        if desc:
                            skill_names.append(f"{name}: {desc}")
                        else:
                            skill_names.append(name)
                elif isinstance(skill, str):
                    skill_names.append(skill)

            if skill_names:
                p = doc.add_paragraph(' | '.join(skill_names))

    def add_awards(self, doc: Document, awards: List[Any],
                   style_metadata: StyleMetadata):
        """添加奖项荣誉 - 支持字符串和对象格式"""
        if not awards:
            return

        self._add_section_title(doc, '奖项荣誉', style_metadata)

        for award in awards:
            if isinstance(award, dict):
                # 对象格式：提取 name 字段
                name = award.get('name', '')
                if name:
                    p = doc.add_paragraph(name, style='List Bullet')
            elif isinstance(award, str) and award:
                # 字符串格式：直接使用
                p = doc.add_paragraph(award, style='List Bullet')

    def add_certificates(self, doc: Document, certificates: List[Any],
                         style_metadata: StyleMetadata):
        """添加证书资质 - 支持字符串和对象格式"""
        if not certificates:
            return

        self._add_section_title(doc, '证书资质', style_metadata)

        for cert in certificates:
            if isinstance(cert, dict):
                # 对象格式：提取 name 字段
                name = cert.get('name', '')
                if name:
                    p = doc.add_paragraph(name, style='List Bullet')
            elif isinstance(cert, str) and cert:
                # 字符串格式：直接使用
                p = doc.add_paragraph(cert, style='List Bullet')

    def add_self_evaluation(self, doc: Document, self_evaluation: str,
                            style_metadata: StyleMetadata):
        """添加自我评价"""
        if not self_evaluation:
            return

        self._add_section_title(doc, '自我评价', style_metadata)
        p = doc.add_paragraph(self_evaluation)
        p.paragraph_format.first_line_indent = Cm(0.74)

    def _add_section_title(self, doc: Document, title: str,
                           style_metadata: StyleMetadata):
        """添加章节标题（使用动态字号）"""
        p = doc.add_paragraph()
        run = p.add_run(title)
        run.bold = True
        run.font.size = Pt(style_metadata.get_section_title_font_size())
        p.paragraph_format.space_before = Pt(12)
        p.paragraph_format.space_after = Pt(6)

        # 添加下划线
        p.add_run('\n' + '─' * 20)


class AtsHtmlBuilder:
    """负责构建 ATS 优化的 HTML 内容"""

    def __init__(self):
        pass

    def build_context(self, tailored_resume: Dict[str, Any],
                      jd_keywords: Optional[List[str]] = None) -> Dict[str, Any]:
        """Build template context from tailored_resume for ATS HTML template."""

        basic = tailored_resume.get('basic_info', {})
        name = basic.get('name', 'Candidate')
        email = basic.get('email', '')
        phone = basic.get('phone', '')
        location = basic.get('location', '')

        # Build contact display (linkedin/github if available)
        contact_parts = []
        if email:
            contact_parts.append(email)
        if phone:
            contact_parts.append(phone)
        contact_display = ' | '.join(contact_parts)

        # Professional Summary
        summary_text = tailored_resume.get('summary', '')

        # Core Competencies from JD keywords
        if jd_keywords:
            competency_tags = '\n              '.join(
                f'<span class="competency-tag">{kw}</span>' for kw in jd_keywords[:8]
            )
        else:
            competency_tags = ''

        # Work Experience → HTML
        experience_html = self._build_experience(tailored_resume.get('work_experience', []))

        # Projects → HTML (top 4)
        projects_html = self._build_projects(tailored_resume.get('projects', [])[:4])

        # Education → HTML
        education_html = self._build_education(tailored_resume.get('education', []))

        # Certifications → HTML
        certs_html = self._build_certificates(tailored_resume.get('certificates', []))

        # Skills → HTML
        skills_html = self._build_skills(tailored_resume.get('skills', []))

        return {
            'LANG': 'en',
            'PAGE_WIDTH': '8.5in',
            'NAME': name,
            'EMAIL': email,
            'LINKEDIN_URL': '#',
            'LINKEDIN_DISPLAY': 'LinkedIn',
            'PORTFOLIO_URL': '#',
            'PORTFOLIO_DISPLAY': 'Portfolio',
            'LOCATION': location,
            'SECTION_SUMMARY': 'Professional Summary',
            'SUMMARY_TEXT': summary_text,
            'SECTION_COMPETENCIES': 'Core Competencies',
            'COMPETENCIES': competency_tags,
            'SECTION_EXPERIENCE': 'Work Experience',
            'EXPERIENCE': experience_html,
            'SECTION_PROJECTS': 'AI Projects',
            'PROJECTS': projects_html,
            'SECTION_EDUCATION': 'Education',
            'EDUCATION': education_html,
            'SECTION_CERTIFICATIONS': 'Certifications',
            'CERTIFICATIONS': certs_html,
            'SECTION_SKILLS': 'Technical Skills',
            'SKILLS': skills_html,
        }

    def _build_experience(self, work_experience: List[Dict]) -> str:
        """Build HTML for work experience section."""
        if not work_experience:
            return '<p style="color:#888;">No work experience listed.</p>'

        html_parts = []
        for exp in work_experience:
            company = exp.get('company', '')
            time = exp.get('time', '')
            position = exp.get('position', '')

            # Get tailored content
            content = exp.get('tailored', '')
            if not content and 'tailored_bullets' in exp:
                bullets = exp.get('tailored_bullets', [])
                if isinstance(bullets, list):
                    contents = [b.get('content', '') if isinstance(b, dict) else b for b in bullets]
                    content = '\n'.join(filter(None, contents))
            if not content:
                content = exp.get('content', '')

            # Build bullet items
            bullets_html = ''
            if content:
                lines = content.split('\n')
                for line in lines:
                    line = line.strip()
                    if line:
                        # Strip common bullet prefixes
                        line = line.lstrip('- •·').strip()
                        if line:
                            bullets_html += f'<li>{line}</li>\n'

            html_parts.append(f'''
            <div class="job">
              <div class="job-header">
                <span class="job-company">{self._esc_html(company)}</span>
                <span class="job-period">{self._esc_html(time)}</span>
              </div>
              <div class="job-role">{self._esc_html(position)}</div>
              <ul>{bullets_html}</ul>
            </div>''')

        return '\n'.join(html_parts)

    def _build_projects(self, projects: List[Dict]) -> str:
        """Build HTML for projects section."""
        if not projects:
            return ''

        html_parts = []
        for proj in projects:
            name = proj.get('name', 'Project')
            content = proj.get('tailored', '') or proj.get('tailored_description', '') or proj.get('content', '')

            # Build tech from project if available
            tech = proj.get('tech', proj.get('technologies', ''))
            tech_html = f'<div class="project-tech">{self._esc_html(tech)}</div>' if tech else ''

            # Clean content for description
            desc = content.strip().split('\n')[0] if content else ''

            html_parts.append(f'''
            <div class="project">
              <div class="project-title">{self._esc_html(name)}</div>
              <div class="project-desc">{self._esc_html(desc)}</div>
              {tech_html}
            </div>''')

        return '\n'.join(html_parts)

    def _build_education(self, education: List[Dict]) -> str:
        """Build HTML for education section."""
        if not education:
            return ''

        html_parts = []
        for edu in education:
            school = edu.get('school', '')
            major = edu.get('major', '')
            degree = edu.get('degree', '')

            title_parts = []
            if major:
                title_parts.append(major)
            if degree:
                title_parts.append(degree)
            title_text = ' &mdash; '.join(title_parts)

            html_parts.append(f'''
            <div class="edu-item">
              <div class="edu-header">
                <span><span class="edu-title">{title_text}</span> &mdash; <span class="edu-org">{self._esc_html(school)}</span></span>
              </div>
            </div>''')

        return '\n'.join(html_parts)

    def _build_certificates(self, certificates: List) -> str:
        """Build HTML for certifications section."""
        if not certificates:
            return ''

        html_parts = []
        for cert in certificates:
            if isinstance(cert, dict):
                name = cert.get('name', '')
                org = cert.get('org', cert.get('issuer', ''))
                year = cert.get('year', cert.get('date', ''))
            else:
                name = str(cert)
                org = ''
                year = ''

            if name:
                org_html = f'<span class="cert-org">{self._esc_html(org)}</span>' if org else ''
                year_html = f'<span class="cert-year">{self._esc_html(year)}</span>' if year else ''
                html_parts.append(f'''
            <div class="cert-item">
              <span class="cert-title">{self._esc_html(name)}</span>
              {org_html}
              {year_html}
            </div>''')

        return '\n'.join(html_parts)

    def _build_skills(self, skills: List) -> str:
        """Build HTML for skills section."""
        if not skills:
            return ''

        html_parts = []
        for skill in skills:
            if isinstance(skill, dict):
                name = skill.get('name', '')
                desc = skill.get('tailored_description', skill.get('description', ''))
                text = f'{name}: {desc}' if desc else name
            else:
                text = str(skill)
            if text:
                html_parts.append(f'<span class="skill-item">{self._esc_html(text)}</span>')

        return '\n'.join(html_parts)

    @staticmethod
    def _esc_html(text: str) -> str:
        """Escape HTML special characters."""
        if not text:
            return ''
        return (str(text)
                .replace('&', '&amp;')
                .replace('<', '&lt;')
                .replace('>', '&gt;')
                .replace('"', '&quot;'))


class ResumeGenerator:
    """简历生成器 - Word 文档输出"""

    def __init__(self):
        self.template_dir = config.BASE_DIR / 'templates'
        self.output_dir = config.BASE_DIR / 'output'
        self.output_dir.mkdir(exist_ok=True)
        
        # 组合分离出的构建器
        self._word_builder = WordDocumentBuilder()
        self._ats_builder = AtsHtmlBuilder()

    def generate_word(self, tailored_resume: Dict[str, Any],
                      output_path: Optional[str] = None,
                      style: str = 'original') -> str:
        """
        生成 Word 文档

        Args:
            tailored_resume: 定制后的简历数据
            output_path: 输出路径（可选）
            style: 样式（MVP仅支持 'original'）

        Returns:
            str: 生成的文件路径
        """
        logger.info(f"开始生成 Word 文档: style={style}")

        # 创建文档
        doc = self._create_document(tailored_resume)

        # 生成输出路径
        if not output_path:
            timestamp = datetime.now().strftime('%Y%m%d_%H%M%S')
            name = tailored_resume.get('basic_info', {}).get('name', '简历')
            output_path = str(self.output_dir / f'{name}_定制简历_{timestamp}.docx')

        # 保存文档
        doc.save(output_path)
        logger.info(f"Word 文档生成完成: {output_path}")

        return output_path

    def generate_pdf(self, tailored_resume: Dict[str, Any],
                     output_path: Optional[str] = None) -> str:
        """
        生成 PDF 文档（先 Word 后转换）

        Args:
            tailored_resume: 定制后的简历数据
            output_path: 输出路径（可选）

        Returns:
            str: 生成的 PDF 文件路径
        """
        # 先生成 Word
        word_path = self.generate_word(tailored_resume)

        # 转换为 PDF
        pdf_path = word_path.replace('.docx', '.pdf')

        try:
            from docx2pdf import convert
            convert(word_path, pdf_path)
            logger.info(f"PDF 文档生成完成: {pdf_path}")
            return pdf_path
        except ImportError:
            logger.warning("docx2pdf 未安装，跳过 PDF 生成")
            return word_path
        except Exception as e:
            logger.error(f"PDF 转换失败: {e}")
            return word_path

    def generate_bytes(self, tailored_resume: Dict[str, Any],
                       format: str = 'word',
                       style_metadata: StyleMetadata = None) -> bytes:
        """
        生成文档字节流

        Args:
            tailored_resume: 定制后的简历数据
            format: 格式 ('word' 或 'pdf')
            style_metadata: 样式元数据（可选）

        Returns:
            bytes: 文档字节流
        """
        # 生成到内存
        doc = self._create_document(tailored_resume, style_metadata)

        # 保存到字节流
        bio = io.BytesIO()
        doc.save(bio)
        bio.seek(0)

        if format == 'pdf':
            # TODO: 实现 PDF 字节流转换
            pass

        return bio.read()

    def _create_document(self, tailored_resume: Dict[str, Any],
                        style_metadata: StyleMetadata = None) -> Document:
        """创建文档对象"""
        
        def add_content(doc: Document, style: StyleMetadata):
            """内部回调函数，用于添加内容"""
            self._word_builder.add_basic_info(doc, tailored_resume.get('basic_info', {}), style)
            self._word_builder.add_summary(doc, tailored_resume.get('summary', ''), style)
            self._word_builder.add_education(doc, tailored_resume.get('education', []), style)
            self._word_builder.add_work_experience(doc, tailored_resume.get('work_experience', []), style)
            self._word_builder.add_projects(doc, tailored_resume.get('projects', []), style)
            self._word_builder.add_skills(doc, tailored_resume.get('skills', []), style)
            self._word_builder.add_awards(doc, tailored_resume.get('awards', []), style)
            self._word_builder.add_certificates(doc, tailored_resume.get('certificates', []), style)
            self._word_builder.add_self_evaluation(doc, tailored_resume.get('self_evaluation', ''), style)

        return self._word_builder.create_document(tailored_resume, style_metadata, add_content)

    # ================================================================
    # ATS-optimized HTML/PDF output (career-ops integration)
    # ================================================================

    def generate_ats_html(self, tailored_resume: Dict[str, Any],
                          jd_keywords: Optional[List[str]] = None,
                          output_path: Optional[str] = None) -> str:
        """
        Generate ATS-optimized HTML resume.

        Args:
            tailored_resume: Tailored resume data from Stage 3
            jd_keywords: List of JD keywords for competency tags (optional)
            output_path: Output HTML file path (optional)

        Returns:
            str: Path to generated HTML file
        """
        ats_template_dir = config.BASE_DIR / 'templates' / 'ats'
        env = Environment(loader=FileSystemLoader(str(ats_template_dir)))
        template = env.get_template('ats_template.html')

        # Build template context from tailored_resume
        context = self._ats_builder.build_context(tailored_resume, jd_keywords)

        # Render HTML
        html_content = template.render(**context)

        # Save to file
        if not output_path:
            timestamp = datetime.now().strftime('%Y%m%d_%H%M%S')
            name = tailored_resume.get('basic_info', {}).get('name', 'resume')
            output_path = str(self.output_dir / f'{name}_ats_{timestamp}.html')

        Path(output_path).write_text(html_content, encoding='utf-8')
        logger.info(f"ATS HTML 生成完成: {output_path}")
        return output_path

    def generate_ats_pdf(self, tailored_resume: Dict[str, Any],
                         jd_keywords: Optional[List[str]] = None,
                         output_path: Optional[str] = None) -> str:
        """
        Generate ATS-optimized PDF resume via Puppeteer.

        Args:
            tailored_resume: Tailored resume data from Stage 3
            jd_keywords: List of JD keywords (optional)
            output_path: Output PDF file path (optional)

        Returns:
            str: Path to generated PDF file
        """
        # First generate HTML
        html_path = self.generate_ats_html(tailored_resume, jd_keywords)

        # Derive PDF path
        if not output_path:
            output_path = html_path.replace('.html', '.pdf')

        # Call generate-pdf.mjs via Node.js
        pdf_tool = config.BASE_DIR / 'tools' / 'generate-pdf.mjs'
        if not pdf_tool.exists():
            logger.error(f"PDF 生成工具不存在: {pdf_tool}")
            return html_path

        try:
            result = subprocess.run(
                ['node', str(pdf_tool), html_path, output_path, '--format=letter'],
                capture_output=True, text=True, timeout=30,
                cwd=str(config.BASE_DIR)
            )
            if result.returncode == 0:
                logger.info(f"ATS PDF 生成完成: {output_path}")
                return output_path
            else:
                logger.error(f"PDF 生成失败: {result.stderr}")
                return html_path
        except Exception as e:
            logger.error(f"PDF 生成异常: {e}")
            return html_path
