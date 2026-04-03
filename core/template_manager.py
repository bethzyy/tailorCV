"""
模板管理器模块

管理简历模板的增删改查、去重、默认模板等功能。
支持多种来源：内置模板、用户上传、从简历提取。
"""

import os
import io
import json
import hashlib
import logging
import shutil
from datetime import datetime
from pathlib import Path
from typing import Dict, Any, List, Optional, Tuple
from dataclasses import dataclass, field, asdict

try:
    from docx import Document
    HAS_PYTHON_DOCX = True
except ImportError:
    HAS_PYTHON_DOCX = False

from .config import config
from .database import db
from .structure_detector import StructureDetector

logger = logging.getLogger(__name__)


@dataclass
class TemplateInfo:
    """模板信息数据类"""
    template_id: str
    name: str
    source: str  # builtin/uploaded/extracted
    file_path: str
    content_hash: str = ""
    structure_confidence: float = 0.0
    sections: List[str] = field(default_factory=list)
    variables: List[str] = field(default_factory=list)
    description: str = ""
    tags: List[str] = field(default_factory=list)
    preview_image: str = ""
    use_count: int = 0
    is_default: bool = False
    created_at: str = ""

    def to_dict(self) -> Dict[str, Any]:
        return asdict(self)


class TemplateManager:
    """
    模板管理器

    主要功能：
    1. 模板 CRUD 操作
    2. 模板去重（基于 content_hash）
    3. 默认模板管理
    4. 内置模板注册
    """

    # 模板来源
    SOURCE_BUILTIN = 'builtin'
    SOURCE_UPLOADED = 'uploaded'
    SOURCE_EXTRACTED = 'extracted'

    # 内置模板注册表
    BUILTIN_TEMPLATES = [
        {
            'id': 'classic_professional',
            'name': '经典专业',
            'description': '传统正式风格，适合金融、政府、教育等行业',
            'tags': ['正式', '传统', '专业'],
            'style': 'classic'
        },
        {
            'id': 'modern_minimal',
            'name': '现代简约',
            'description': '简洁清爽风格，适合互联网、科技行业',
            'tags': ['简约', '现代', '清爽'],
            'style': 'modern'
        },
        {
            'id': 'creative_design',
            'name': '创意设计',
            'description': '个性化设计风格，适合设计、营销、创意行业',
            'tags': ['创意', '设计', '个性'],
            'style': 'creative'
        },
        {
            'id': 'executive_senior',
            'name': '高管资深',
            'description': '大气稳重风格，适合高管、资深职位',
            'tags': ['高管', '资深', '大气'],
            'style': 'executive'
        },
        {
            'id': 'academic_research',
            'name': '学术研究',
            'description': '规范学术风格，适合科研、学术、教育行业',
            'tags': ['学术', '科研', '教育'],
            'style': 'academic'
        },
        {
            'id': 'tech_engineer',
            'name': '技术工程师',
            'description': '结构化清晰风格，适合程序员、工程师',
            'tags': ['技术', '工程师', 'IT'],
            'style': 'tech'
        },
    ]

    def __init__(self):
        """初始化模板管理器"""
        self.base_dir = config.BASE_DIR / 'templates'
        self.builtin_dir = self.base_dir / 'builtin'
        self.uploaded_dir = self.base_dir / 'uploaded'
        self.extracted_dir = self.base_dir / 'extracted'
        self.previews_dir = self.base_dir / 'previews'

        # 确保目录存在
        for dir_path in [self.builtin_dir, self.uploaded_dir,
                         self.extracted_dir, self.previews_dir]:
            dir_path.mkdir(parents=True, exist_ok=True)

        # 结构检测器
        self.detector = StructureDetector()

        # 初始化内置模板
        self._init_builtin_templates()

        # 恢复数据库丢失的自建模板（从文件重新注册）
        self._recover_user_templates()

    def _init_builtin_templates(self):
        """初始化内置模板到数据库"""
        import io as io_module
        import re as re_module

        for template_info in self.BUILTIN_TEMPLATES:
            template_id = template_info['id']
            file_path = self.builtin_dir / f"{template_id}.docx"

            # 检查数据库中是否已存在
            existing = db.get_template(template_id)
            if existing:
                continue

            # 检查文件是否存在
            if not file_path.exists():
                logger.warning(f"内置模板文件不存在: {file_path}")
                continue

            # 读取文件内容
            with open(file_path, 'rb') as f:
                file_content = f.read()

            # 计算内容哈希
            content_hash = hashlib.md5(file_content).hexdigest()

            # 检测模板结构
            sections, confidence = self._detect_template_structure(file_path)

            # 检查模板是否已有 Jinja2 标签
            has_jinja = self._check_existing_jinja_tags(file_content)
            template_variables = []

            if has_jinja:
                # 模板已有 Jinja2 标签，直接提取变量并复制到 preprocessed 目录
                logger.info(f"内置模板已有 Jinja2 标签，跳过预处理: {template_id}")
                template_variables = self._get_template_variables(file_path)

                # 复制到 preprocessed 目录（使用 template_id 作为文件名）
                preprocessed_path = self.base_dir / 'preprocessed' / f"{template_id}.docx"
                shutil.copy(file_path, preprocessed_path)
                logger.info(f"复制内置模板到 preprocessed 目录: {template_id}")
            else:
                # 模板没有 Jinja2 标签，需要预处理
                try:
                    from .template_processor import TemplateProcessor
                    template_processor = TemplateProcessor()

                    doc = Document(io_module.BytesIO(file_content))
                    preprocess_result = template_processor.preprocess(
                        doc,
                        original_filename=f"{template_id}.docx",
                        original_content=file_content
                    )

                    if preprocess_result.success:
                        template_variables = list(preprocess_result.metadata.variables) if preprocess_result.metadata.variables else []
                        logger.info(f"内置模板预处理成功: {template_id}, 变量数: {len(template_variables)}")

                        # 复制预处理后的模板到 template_id 命名的文件
                        preprocessed_path = self.base_dir / 'preprocessed' / f"{template_id}.docx"
                        shutil.copy(preprocess_result.template_path, preprocessed_path)
                    else:
                        logger.warning(f"内置模板预处理失败: {preprocess_result.error_message}")
                        template_variables = self._get_template_variables(file_path)
                except Exception as e:
                    logger.warning(f"内置模板预处理异常: {e}")
                    template_variables = self._get_template_variables(file_path)

            # 保存到数据库
            db.save_template({
                'template_id': template_id,
                'name': template_info['name'],
                'source': self.SOURCE_BUILTIN,
                'file_path': str(file_path),
                'content_hash': content_hash,
                'structure_confidence': confidence,
                'sections': sections,
                'variables': template_variables,
                'description': template_info['description'],
                'tags': template_info['tags'],
                'preview_image': str(self.previews_dir / f"{template_id}.png"),
                'is_default': template_id == 'classic_professional',  # 默认模板
                'use_count': 0
            })
            logger.info(f"注册内置模板: {template_info['name']}")

    def get_templates(self, source: str = None) -> List[Dict[str, Any]]:
        """
        获取模板列表

        Args:
            source: 来源过滤 (builtin/uploaded/extracted)，None 表示全部

        Returns:
            List[Dict]: 模板列表
        """
        return db.get_templates(source=source)

    def _recover_user_templates(self):
        """从 uploaded/ 和 extracted/ 目录恢复数据库中丢失的自建模板"""
        import hashlib

        for source_dir, source_type in [(self.uploaded_dir, 'uploaded'),
                                       (self.extracted_dir, 'extracted')]:
            if not source_dir.exists():
                continue
            for file_path in source_dir.glob('*.docx'):
                template_id = file_path.stem
                # 检查数据库中是否已有此模板
                existing = db.get_template(template_id)
                if existing:
                    continue
                try:
                    content = file_path.read_bytes()
                    content_hash = hashlib.md5(content).hexdigest()
                    # 检查文件名是否与哈希匹配（旧格式文件名就是哈希）
                    if template_id == content_hash[:16]:
                        logger.info(f"恢复自建模板 [{source_type}]: {file_path.name}")
                        from .template_processor import TemplateProcessor
                        template_processor = TemplateProcessor()
                        from docx import Document
                        import io as io_module
                        doc = Document(io_module.BytesIO(content))
                        sections, confidence = self._detect_template_structure_from_doc(doc)
                        preprocessed_path = Path('templates/preprocessed') / f"{template_id}.docx"
                        if preprocessed_path.exists():
                            shutil.copy(preprocessed_path, file_path)
                        else:
                            # 无预处理文件，直接用原始文件
                            logger.warning(f"无预处理文件，使用原始文件恢复: {template_id}")
                        db.save_template({
                            'template_id': template_id,
                            'name': f"恢复模板 {template_id[:8]}",
                            'source': source_type,
                            'file_path': str(file_path),
                            'content_hash': content_hash,
                            'structure_confidence': confidence,
                            'sections': sections,
                            'variables': [],
                            'description': '从文件恢复',
                            'tags': [],
                            'preview_image': '',
                            'is_default': False,
                            'use_count': 0
                        })
                except Exception as e:
                    logger.error(f"恢复模板失败 {file_path.name}: {e}")

    def get_template(self, template_id: str) -> Optional[Dict[str, Any]]:
        """获取指定模板"""
        return db.get_template(template_id)

    def get_default_template(self) -> Optional[Dict[str, Any]]:
        """获取默认模板"""
        return db.get_default_template()

    def set_default_template(self, template_id: str) -> bool:
        """设置默认模板"""
        template = db.get_template(template_id)
        if not template:
            logger.warning(f"模板不存在: {template_id}")
            return False
        return db.set_default_template(template_id)

    def upload_template(self, file_content: bytes, filename: str,
                        name: str = None, description: str = "",
                        tags: List[str] = None) -> Tuple[Optional[str], str]:
        """
        上传新模板

        Args:
            file_content: 文件内容（字节）
            filename: 原始文件名
            name: 模板名称（可选，默认使用文件名）
            description: 模板描述
            tags: 标签列表

        Returns:
            Tuple[template_id, error_message]: 模板ID和错误信息
        """
        if not HAS_PYTHON_DOCX:
            return None, "未安装 python-docx"

        # 检查文件格式
        ext = Path(filename).suffix.lower()
        if ext not in ['.docx', '.doc']:
            return None, "仅支持 .docx 格式"

        # 检查文件大小
        if len(file_content) > 5 * 1024 * 1024:  # 5MB
            return None, "文件大小不能超过 5MB"

        # 计算内容哈希
        content_hash = hashlib.md5(file_content).hexdigest()

        # 检查是否已存在相同模板
        existing = db.get_template_by_hash(content_hash)
        if existing:
            return existing['template_id'], f"已存在相同模板: {existing['name']}"

        try:
            # 解析文档
            doc = Document(io=io.BytesIO(file_content)) if 'io' in dir() else None
            import io as io_module
            doc = Document(io_module.BytesIO(file_content))

            # 检测结构
            sections, confidence = self._detect_template_structure_from_doc(doc)

            # 生成模板ID
            template_id = content_hash[:16]

            # 确定模板名称
            if not name:
                name = Path(filename).stem

            # 使用 TemplateProcessor 预处理，生成带 Jinja2 标记的模板
            from .template_processor import TemplateProcessor
            template_processor = TemplateProcessor()

            # 预处理文档，插入 Jinja2 标记
            preprocess_result = template_processor.preprocess(
                doc,
                original_filename=filename,
                original_content=file_content
            )

            if preprocess_result.success:
                # 使用预处理后的模板（已保存到 templates/preprocessed/ 目录）
                preprocessed_path = Path(preprocess_result.template_path)
                # 复制到 uploaded 目录
                file_path = self.uploaded_dir / f"{template_id}.docx"
                shutil.copy(preprocessed_path, file_path)
                logger.info(f"上传模板预处理成功，已插入 Jinja2 标记: {template_id}")

                # 更新变量列表（从预处理结果获取）
                template_variables = list(preprocess_result.metadata.variables) if preprocess_result.metadata.variables else []
            else:
                # 降级：保存原始文件，但记录警告
                logger.warning(f"上传模板预处理失败: {preprocess_result.error_message}，使用原始文件")
                file_path = self.uploaded_dir / f"{template_id}.docx"
                with open(file_path, 'wb') as f:
                    f.write(file_content)
                template_variables = self._get_template_variables_from_doc(doc)

            # 保存到数据库
            db.save_template({
                'template_id': template_id,
                'name': name,
                'source': self.SOURCE_UPLOADED,
                'file_path': str(file_path),
                'content_hash': content_hash,
                'structure_confidence': confidence,
                'sections': sections,
                'variables': template_variables,
                'description': description,
                'tags': tags or [],
                'preview_image': '',
                'is_default': False,
                'use_count': 0
            })

            logger.info(f"上传模板成功: {name} ({template_id})")
            return template_id, ""

        except Exception as e:
            logger.error(f"上传模板失败: {e}", exc_info=True)
            return None, str(e)

    def extract_template_from_resume(self, file_content: bytes, filename: str,
                                     name: str = None) -> Tuple[Optional[str], str]:
        """
        从简历提取模板

        Args:
            file_content: 简历文件内容（字节）
            filename: 原始文件名
            name: 模板名称（可选）

        Returns:
            Tuple[template_id, error_message]: 模板ID和错误信息
        """
        if not HAS_PYTHON_DOCX:
            return None, "未安装 python-docx"

        # 检查文件格式
        ext = Path(filename).suffix.lower()
        if ext not in ['.docx', '.doc']:
            return None, "仅支持 .docx 格式的简历"

        # 计算内容哈希
        content_hash = hashlib.md5(file_content).hexdigest()

        # 检查是否已存在相同模板
        existing = db.get_template_by_hash(content_hash)
        if existing:
            return existing['template_id'], f"已存在相同模板: {existing['name']}"

        try:
            import io as io_module
            doc = Document(io_module.BytesIO(file_content))

            # 检测结构
            sections, confidence = self._detect_template_structure_from_doc(doc)

            # 结构置信度检查
            if confidence < 0.3:
                return None, f"简历结构不清晰（置信度 {confidence:.0%}），无法提取模板"

            # 生成模板ID
            template_id = content_hash[:16]

            # 确定模板名称
            if not name:
                name = f"从 {Path(filename).stem} 提取"

            # 使用 TemplateProcessor 预处理，生成带 Jinja2 标记的模板
            from .template_processor import TemplateProcessor
            template_processor = TemplateProcessor()

            # 预处理文档，插入 Jinja2 标记
            preprocess_result = template_processor.preprocess(
                doc,
                original_filename=filename,
                original_content=file_content
            )

            if preprocess_result.success:
                # 使用预处理后的模板（已保存到 templates/preprocessed/ 目录）
                preprocessed_path = Path(preprocess_result.template_path)
                # 复制到 extracted 目录
                file_path = self.extracted_dir / f"{template_id}.docx"
                shutil.copy(preprocessed_path, file_path)
                logger.info(f"模板预处理成功，已插入 Jinja2 标记: {template_id}")

                # 更新变量列表（从预处理结果获取）
                template_variables = list(preprocess_result.metadata.variables) if preprocess_result.metadata.variables else []
            else:
                # 降级：保存原始文件，但记录警告
                logger.warning(f"模板预处理失败: {preprocess_result.error_message}，使用原始文件")
                file_path = self.extracted_dir / f"{template_id}.docx"
                with open(file_path, 'wb') as f:
                    f.write(file_content)
                template_variables = self._get_template_variables_from_doc(doc)

            # 保存到数据库
            db.save_template({
                'template_id': template_id,
                'name': name,
                'source': self.SOURCE_EXTRACTED,
                'file_path': str(file_path),
                'content_hash': content_hash,
                'structure_confidence': confidence,
                'sections': sections,
                'variables': template_variables,
                'description': f"从简历 {filename} 自动提取",
                'tags': ['自动提取'],
                'preview_image': '',
                'is_default': False,
                'use_count': 0
            })

            logger.info(f"提取模板成功: {name} ({template_id})")
            return template_id, ""

        except Exception as e:
            logger.error(f"提取模板失败: {e}", exc_info=True)
            return None, str(e)

    def delete_template(self, template_id: str) -> Tuple[bool, str]:
        """
        删除模板

        Args:
            template_id: 模板ID

        Returns:
            Tuple[success, error_message]: 是否成功和错误信息
        """
        template = db.get_template(template_id)
        if not template:
            return False, "模板不存在"

        if template['source'] == self.SOURCE_BUILTIN:
            return False, "内置模板不能删除"

        # 删除文件
        file_path = Path(template['file_path'])
        if file_path.exists():
            try:
                file_path.unlink()
            except Exception as e:
                logger.warning(f"删除模板文件失败: {e}")

        # 删除预览图
        if template.get('preview_image'):
            preview_path = Path(template['preview_image'])
            if preview_path.exists():
                try:
                    preview_path.unlink()
                except Exception as e:
                    logger.warning(f"删除预览图失败: {e}")

        # 从数据库删除
        success = db.delete_template(template_id)
        return success, "" if success else "删除失败"

    def increment_use_count(self, template_id: str):
        """增加模板使用次数"""
        db.increment_template_use_count(template_id)

    def get_template_file(self, template_id: str) -> Optional[bytes]:
        """获取模板文件内容"""
        template = db.get_template(template_id)
        if not template:
            return None

        file_path = Path(template['file_path'])
        if not file_path.exists():
            logger.warning(f"模板文件不存在: {file_path}")
            return None

        with open(file_path, 'rb') as f:
            return f.read()

    def check_compatibility(self, template_id: str,
                            resume_data: Dict[str, Any]) -> Tuple[bool, List[str]]:
        """
        检查模板与简历数据的兼容性

        Args:
            template_id: 模板ID
            resume_data: 简历数据

        Returns:
            Tuple[is_compatible, missing_sections]: 是否兼容和缺失的章节列表
        """
        template = db.get_template(template_id)
        if not template:
            return False, ["模板不存在"]

        template_sections = set(template.get('sections', []))
        resume_sections = set()

        # 从简历数据中提取已有章节
        if resume_data.get('basic_info'):
            resume_sections.add('basic_info')
        if resume_data.get('education'):
            resume_sections.add('education')
        if resume_data.get('work_experience'):
            resume_sections.add('work_experience')
        if resume_data.get('projects'):
            resume_sections.add('projects')
        if resume_data.get('skills'):
            resume_sections.add('skills')
        if resume_data.get('awards'):
            resume_sections.add('awards')
        if resume_data.get('certificates'):
            resume_sections.add('certificates')
        if resume_data.get('self_evaluation'):
            resume_sections.add('self_evaluation')

        # 检查缺失的章节
        missing = template_sections - resume_sections

        # 如果缺失关键章节，返回不兼容
        critical_sections = {'basic_info', 'work_experience', 'education'}
        critical_missing = critical_sections & missing

        is_compatible = len(critical_missing) == 0
        return is_compatible, list(missing)

    def recommend_template(self, jd_content: str = "",
                          industry: str = "",
                          position_level: str = "") -> List[Dict[str, Any]]:
        """
        根据 JD 内容推荐合适模板

        Args:
            jd_content: JD 文本内容
            industry: 行业类型（可选）
            position_level: 职位级别（可选）

        Returns:
            List[Dict]: 推荐模板列表，按推荐度排序
        """
        # 行业关键词映射
        industry_keywords = {
            'tech': ['技术', '开发', '工程师', '程序', 'IT', '互联网', '软件', '代码', '算法',
                     'developer', 'engineer', 'programmer', 'tech', 'software'],
            'finance': ['金融', '银行', '投资', '证券', '基金', '财务', '会计', '审计',
                       'finance', 'bank', 'investment', 'accounting'],
            'creative': ['设计', '创意', '营销', '品牌', '广告', '市场', '运营', '产品',
                        'design', 'creative', 'marketing', 'brand'],
            'academic': ['研究', '学术', '科研', '高校', '大学', '教授', '博士',
                        'research', 'academic', 'university', 'phd'],
            'executive': ['总监', '经理', 'VP', 'CEO', 'CTO', '高管', '领导', '负责',
                         'director', 'manager', 'executive', 'vp', 'chief'],
        }

        # 模板推荐权重
        template_weights = {
            'tech_engineer': {'tech': 3, 'creative': 1},
            'classic_professional': {'finance': 3, 'default': 2},
            'modern_minimal': {'tech': 2, 'creative': 2, 'default': 2},
            'creative_design': {'creative': 3},
            'academic_research': {'academic': 3, 'default': 1},
            'executive_senior': {'executive': 3, 'finance': 2, 'default': 1},
        }

        # 检测行业
        detected_industries = set()
        if industry:
            detected_industries.add(industry.lower())
        if jd_content:
            jd_lower = jd_content.lower()
            for ind, keywords in industry_keywords.items():
                for kw in keywords:
                    if kw.lower() in jd_lower:
                        detected_industries.add(ind)
                        break

        # 计算每个模板的推荐分数
        templates = self.get_templates()
        recommendations = []

        for template in templates:
            template_id = template['template_id']
            weights = template_weights.get(template_id, {'default': 1})

            score = weights.get('default', 0)
            for ind in detected_industries:
                if ind in weights:
                    score = max(score, weights[ind] * 2)

            # 默认模板加分
            if template.get('is_default'):
                score += 1

            # 使用次数加分（热度）
            score += min(template.get('use_count', 0) / 10, 2)

            recommendations.append({
                'template': template,
                'score': score,
                'reason': self._get_recommendation_reason(template_id, detected_industries)
            })

        # 按分数排序
        recommendations.sort(key=lambda x: x['score'], reverse=True)
        return recommendations

    def _get_recommendation_reason(self, template_id: str,
                                   industries: set) -> str:
        """获取推荐原因"""
        reasons = {
            'tech_engineer': '适合技术岗位，结构化展示技术栈和项目经验',
            'classic_professional': '传统正式风格，适合大多数行业',
            'modern_minimal': '现代简约设计，突出核心内容',
            'creative_design': '创意风格，适合设计、营销类岗位',
            'academic_research': '学术规范格式，适合科研、教育岗位',
            'executive_senior': '大气稳重风格，适合高级管理岗位',
        }

        base_reason = reasons.get(template_id, '推荐使用')

        if 'tech' in industries:
            if template_id == 'tech_engineer':
                return '技术岗位首选'
        if 'executive' in industries:
            if template_id == 'executive_senior':
                return '高级职位首选'
        if 'academic' in industries:
            if template_id == 'academic_research':
                return '学术岗位首选'
        if 'creative' in industries:
            if template_id == 'creative_design':
                return '创意岗位首选'

        return base_reason

    def _calculate_file_hash(self, file_path: Path) -> str:
        """计算文件内容哈希"""
        with open(file_path, 'rb') as f:
            return hashlib.md5(f.read()).hexdigest()

    def _detect_template_structure(self, file_path: Path) -> Tuple[List[str], float]:
        """从文件检测模板结构"""
        if not HAS_PYTHON_DOCX:
            return [], 0.0

        try:
            doc = Document(str(file_path))
            return self._detect_template_structure_from_doc(doc)
        except Exception as e:
            logger.warning(f"检测模板结构失败: {e}")
            return [], 0.0

    def _detect_template_structure_from_doc(self, doc: 'Document') -> Tuple[List[str], float]:
        """从 Document 对象检测模板结构"""
        try:
            structure = self.detector.detect_structure(doc)
            sections = [s.section_type.value for s in structure.sections]
            return sections, structure.confidence
        except Exception as e:
            logger.warning(f"检测文档结构失败: {e}")
            return [], 0.0

    def _get_template_variables(self, file_path: Path) -> List[str]:
        """从文件提取模板变量"""
        if not HAS_PYTHON_DOCX:
            return []

        try:
            doc = Document(str(file_path))
            return self._get_template_variables_from_doc(doc)
        except Exception as e:
            logger.warning(f"提取模板变量失败: {e}")
            return []

    def _get_template_variables_from_doc(self, doc: 'Document') -> List[str]:
        """从 Document 对象提取模板变量"""
        import re
        variables = set()

        # Jinja2 变量模式
        patterns = [
            r'\{\{\s*(\w+)',  # {{ variable
            r'\{\%\s*for\s+(\w+)\s+in\s+(\w+)',  # {% for item in list
            r'\{\%\s*if\s+(\w+)',  # {% if variable
        ]

        for paragraph in doc.paragraphs:
            text = paragraph.text
            for pattern in patterns:
                matches = re.findall(pattern, text)
                for match in matches:
                    if isinstance(match, tuple):
                        variables.update(match)
                    else:
                        variables.add(match)

        # 也检查表格
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    text = cell.text
                    for pattern in patterns:
                        matches = re.findall(pattern, text)
                        for match in matches:
                            if isinstance(match, tuple):
                                variables.update(match)
                            else:
                                variables.add(match)

        # 过滤常见的 Jinja2 关键字
        keywords = {'if', 'else', 'endif', 'for', 'endfor', 'in', 'not', 'and', 'or'}
        return sorted(list(variables - keywords))

    def _check_existing_jinja_tags(self, file_content: bytes) -> bool:
        """
        检查文件是否已包含 Jinja2 标签

        Args:
            file_content: 文件内容（字节）

        Returns:
            bool: 是否已包含 Jinja2 标签
        """
        import zipfile
        import re

        try:
            # 解压 docx 并读取 document.xml
            with zipfile.ZipFile(io.BytesIO(file_content)) as zf:
                try:
                    doc_xml = zf.read('word/document.xml').decode('utf-8')
                except KeyError:
                    return False

            # 检查是否有 Jinja2 变量或控制标签
            # 使用更宽松的模式，匹配 {{ }} 和 {% %}
            jinja_patterns = [
                r'\{\{.*?\}\}',  # 变量 {{ ... }}
                r'\{%.*?%\}',    # 控制标签 {% ... %}
            ]

            for pattern in jinja_patterns:
                if re.search(pattern, doc_xml, re.DOTALL):
                    return True

            return False

        except Exception as e:
            logger.warning(f"检查 Jinja2 标签失败: {e}")
            return False

    def get_stats(self) -> Dict[str, Any]:
        """获取模板统计信息"""
        templates = db.get_templates()
        return {
            'total_count': len(templates),
            'builtin_count': len([t for t in templates if t['source'] == self.SOURCE_BUILTIN]),
            'uploaded_count': len([t for t in templates if t['source'] == self.SOURCE_UPLOADED]),
            'extracted_count': len([t for t in templates if t['source'] == self.SOURCE_EXTRACTED]),
            'total_use_count': sum(t.get('use_count', 0) for t in templates)
        }


# 创建全局实例
template_manager = TemplateManager()
