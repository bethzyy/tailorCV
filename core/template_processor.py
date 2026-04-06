"""
模板预处理器模块

整合结构检测和 Jinja2 标记插入，提供完整的模板预处理和渲染功能。
使用 docxtpl 进行模板渲染，确保输出文档保留原简历的样式。
"""

import os
import io
import re
import copy
import uuid
import hashlib
import logging
from typing import Dict, List, Optional, Tuple, Any
from dataclasses import dataclass, field
from pathlib import Path
from datetime import datetime

# Word 文档处理
try:
    from docx import Document
    HAS_PYTHON_DOCX = True
except ImportError:
    HAS_PYTHON_DOCX = False

# docxtpl 模板渲染
try:
    from docxtpl import DocxTemplate
    HAS_DOCXTPL = True
except ImportError:
    HAS_DOCXTPL = False

from .config import config
from .structure_detector import StructureDetector, StructureMap, SectionType
from .jinja_inserter import JinjaTagInserter, TemplateMetadata
from .resume_parser import StyleMetadata

logger = logging.getLogger(__name__)


@dataclass
class PreprocessResult:
    """预处理结果"""
    success: bool
    template_id: str
    template_path: str                       # 模板文件路径
    metadata: TemplateMetadata
    style_metadata: StyleMetadata
    error_message: str = ""


class TemplateProcessor:
    """
    模板预处理器

    主要功能：
    1. preprocess(): 解析原简历，生成带 Jinja2 标记的模板
    2. render(): 使用模板和上下文数据渲染简历
    3. 自动降级：模板提取失败时回退到样式元数据方案
    """

    # 支持的动态章节字段
    DYNAMIC_SECTION_FIELDS = {
        'education': ['time', 'school', 'major', 'degree', 'tailored', 'highlights'],
        'work_experience': ['time', 'company', 'position', 'tailored', 'content'],
        'projects': ['time', 'name', 'role', 'tailored', 'content'],
        'skills': ['name', 'tailored_description'],
        'awards': ['name'],
        'certificates': ['name'],
    }

    def __init__(self):
        self.detector = StructureDetector()
        self.inserter = JinjaTagInserter()

        # 模板存储目录
        self.template_dir = config.BASE_DIR / 'templates' / 'preprocessed'
        self.template_dir.mkdir(parents=True, exist_ok=True)

        # 通用兜底模板路径
        self.generic_template_path = config.BASE_DIR / 'templates' / 'generic.docx'

        # 统计信息
        self.stats = {
            'preprocessed': 0,
            'rendered': 0,
            'fallback_used': 0,
            'failed': 0
        }

    def preprocess(self, doc: 'Document', original_filename: str = "",
                   original_content: bytes = None) -> PreprocessResult:
        """
        预处理原简历，生成带 Jinja2 标记的模板

        Args:
            doc: python-docx Document 对象
            original_filename: 原始文件名
            original_content: 原始文件内容（字节），用于基于内容哈希的去重

        Returns:
            PreprocessResult: 预处理结果
        """
        if not HAS_PYTHON_DOCX:
            return PreprocessResult(
                success=False,
                template_id="",
                template_path="",
                metadata=TemplateMetadata(
                    template_id="",
                    original_filename=original_filename,
                    structure_confidence=0.0
                ),
                style_metadata=StyleMetadata(),
                error_message="未安装 python-docx"
            )

        if not HAS_DOCXTPL:
            return PreprocessResult(
                success=False,
                template_id="",
                template_path="",
                metadata=TemplateMetadata(
                    template_id="",
                    original_filename=original_filename,
                    structure_confidence=0.0
                ),
                style_metadata=StyleMetadata(),
                error_message="未安装 docxtpl"
            )

        try:
            # 1. 检测文档结构
            structure = self.detector.detect_structure(doc)

            # 2. 提取样式元数据
            style_metadata = self._extract_style_metadata(doc)

            # 3. 生成模板ID（基于内容哈希）
            template_id = self._generate_template_id(doc, original_content)

            # 4. 检查是否已存在相同模板（去重）
            template_path = self.template_dir / f"{template_id}.docx"
            if template_path.exists():
                logger.info(f"复用已存在模板: {template_id}")
                # 返回已存在的模板信息
                return PreprocessResult(
                    success=True,
                    template_id=template_id,
                    template_path=str(template_path),
                    metadata=TemplateMetadata(
                        template_id=template_id,
                        original_filename=original_filename,
                        structure_confidence=structure.confidence
                    ),
                    style_metadata=style_metadata
                )

            # 5. 插入 Jinja2 标记
            tagged_doc, metadata = self.inserter.insert_tags(doc, structure, template_id)
            metadata.original_filename = original_filename

            # 6. 保存模板
            tagged_doc.save(str(template_path))

            self.stats['preprocessed'] += 1

            logger.info(f"模板预处理成功: {template_id}, "
                       f"置信度 {structure.confidence:.2f}, "
                       f"变量 {len(metadata.variables)}")

            return PreprocessResult(
                success=True,
                template_id=template_id,
                template_path=str(template_path),
                metadata=metadata,
                style_metadata=style_metadata
            )

        except Exception as e:
            self.stats['failed'] += 1
            logger.error(f"模板预处理失败: {e}", exc_info=True)
            return PreprocessResult(
                success=False,
                template_id="",
                template_path="",
                metadata=TemplateMetadata(
                    template_id="",
                    original_filename=original_filename,
                    structure_confidence=0.0
                ),
                style_metadata=StyleMetadata(),
                error_message=str(e)
            )

    def render(self, template_id: str, context: Dict[str, Any],
               style_metadata: Optional[StyleMetadata] = None) -> bytes:
        """
        使用模板渲染简历

        支持两种调用方式：
        1. 仅提供 template_id：自动查找模板文件
        2. 已有 preprocessed 模板：使用原有逻辑

        Args:
            template_id: 模板ID
            context: 渲染上下文（AI 生成的数据）
            style_metadata: 样式元数据（备用）

        Returns:
            bytes: 渲染后的 Word 文档字节流
        """
        logger.info(f"🔍 render() - template_id: {template_id}")

        if not HAS_DOCXTPL:
            raise ImportError("未安装 docxtpl")

        # 首先检查 preprocessed 目录
        template_path = self.template_dir / f"{template_id}.docx"
        logger.info(f"📁 检查预处理模板: {template_path}, 存在: {template_path.exists()}")

        if template_path.exists():
            # 使用原有逻辑
            try:
                logger.info(f"✅ 使用预处理模板渲染")
                doc = DocxTemplate(str(template_path))
                render_context = self._build_context(context)
                logger.info(f"📝 构建渲染上下文完成，变量数: {len(render_context)}")
                doc.render(render_context)
                self._post_process_tailored_content(doc)
                bio = io.BytesIO()
                doc.save(bio)
                bio.seek(0)

                self.stats['rendered'] += 1
                logger.info(f"✅ 模板渲染成功: {template_id}")

                return bio.read()

            except Exception as e:
                self.stats['failed'] += 1
                logger.error(f"❌ 模板渲染失败: {e}", exc_info=True)
                raise

        logger.info(f"⚠️ 预处理模板不存在，尝试 render_by_id 查找其他位置")
        # 尝试使用 render_by_id 查找其他位置的模板
        return self.render_by_id(template_id, context, style_metadata)

    def render_with_fallback(self, doc: 'Document', context: Dict[str, Any],
                            style_metadata: StyleMetadata,
                            original_filename: str = "",
                            preprocess_result=None,
                            original_content: bytes = None) -> Tuple[bytes, bool]:
        """
        带降级的渲染

        如果模板提取失败，自动降级到样式元数据方案。

        Args:
            doc: 原始文档对象
            context: 渲染上下文
            style_metadata: 样式元数据
            original_filename: 原始文件名
            preprocess_result: 已有的预处理结果（避免重复预处理）
            original_content: 原始文件字节（用于生成一致的 template_id）

        Returns:
            Tuple[bytes, bool]: (文档字节流, 是否使用了模板)
        """
        # 使用已有的预处理结果，或重新预处理
        if preprocess_result is None:
            preprocess_result = self.preprocess(doc, original_filename, original_content=original_content)

        if preprocess_result.success and preprocess_result.metadata.structure_confidence >= 0.2:
            # 模板提取成功，使用模板渲染
            try:
                word_bytes = self.render(
                    preprocess_result.template_id,
                    context,
                    preprocess_result.style_metadata
                )
                logger.info(f"使用模板渲染成功 (置信度: {preprocess_result.metadata.structure_confidence:.2f})")
                return word_bytes, True
            except Exception as e:
                logger.warning(f"模板渲染失败，降级到样式方案: {e}")

        # 降级：使用样式元数据方案
        self.stats['fallback_used'] += 1
        logger.info("使用样式元数据方案渲染")

        # 使用原有的 ResumeGenerator 逻辑
        from .resume_generator import ResumeGenerator
        generator = ResumeGenerator()
        word_bytes = generator.generate_bytes(context, style_metadata=style_metadata)

        return word_bytes, False

    def _build_context(self, tailored_resume: Dict[str, Any]) -> Dict[str, Any]:
        """
        构建模板渲染上下文

        将 AI 生成的数据转换为模板可用的格式。
        支持两种格式：
        1. 列表格式（循环）：work_experience[0].time
        2. 扁平格式：work_experience_0_time

        Args:
            tailored_resume: AI 生成的定制简历数据

        Returns:
            Dict[str, Any]: 渲染上下文
        """
        context = {}

        # 基本信息
        context['basic_info'] = tailored_resume.get('basic_info', {})

        # 个人简介
        context['summary'] = tailored_resume.get('summary', '')

        # 动态章节 - 同时支持列表格式和扁平格式
        dynamic_sections = [
            ('education', tailored_resume.get('education', [])),
            ('work_experience', tailored_resume.get('work_experience', [])),
            ('projects', tailored_resume.get('projects', [])),
            ('skills', tailored_resume.get('skills', [])),
            ('awards', tailored_resume.get('awards', [])),
            ('certificates', tailored_resume.get('certificates', [])),
        ]

        for section_name, section_data in dynamic_sections:
            # 列表格式（用于循环）
            context[section_name] = self._process_section(section_data, section_name)

            # 扁平格式（用于简单变量替换）
            self._add_flat_variables(context, section_name, section_data)

        # 自我评价
        context['self_evaluation'] = tailored_resume.get('self_evaluation', '')

        return context

    def _add_flat_variables(self, context: Dict[str, Any], section_name: str,
                           section_data: List[Dict[str, Any]]):
        """
        添加扁平格式的变量（如 work_experience_0_time）

        Args:
            context: 上下文字典（会被修改）
            section_name: 章节名称
            section_data: 章节数据列表
        """
        expected_fields = self.DYNAMIC_SECTION_FIELDS.get(section_name, [])

        for index, item in enumerate(section_data):
            if isinstance(item, dict):
                for field in expected_fields:
                    key = f"{section_name}_{index}_{field}"
                    context[key] = item.get(field, '')
                # 保留其他字段
                for key, value in item.items():
                    if key not in expected_fields:
                        flat_key = f"{section_name}_{index}_{key}"
                        context[flat_key] = value
            elif isinstance(item, str):
                key = f"{section_name}_{index}_name"
                context[key] = item

    def _process_section(self, section_data: Any, section_type: str) -> List[Dict[str, Any]]:
        """
        处理章节数据，确保格式正确

        Args:
            section_data: 章节数据（可能是列表或字典）
            section_type: 章节类型

        Returns:
            List[Dict[str, Any]]: 处理后的数据列表
        """
        if not section_data:
            return []

        # 兜底：skills 可能是 dict 格式（AI 返回 ordered_by_jd_relevance/other_skills）
        if section_type == 'skills' and isinstance(section_data, dict):
            skill_list = []
            for key in ['ordered_by_jd_relevance', 'other_skills', 'items', 'technical']:
                items = section_data.get(key, [])
                if isinstance(items, list):
                    for item in items:
                        if isinstance(item, dict):
                            name = item.get('skill', item.get('name', ''))
                            desc = item.get('context', item.get('tailored_description', item.get('description', '')))
                            level = item.get('level', '')
                            full_desc = f"{level} — {desc}" if level and desc else (level or desc or '')
                            skill_list.append({'name': name, 'tailored_description': full_desc})
                        elif isinstance(item, str) and item.strip():
                            skill_list.append({'name': item.strip(), 'tailored_description': ''})
            if skill_list:
                section_data = skill_list

        # 确保是列表
        if not isinstance(section_data, list):
            return []

        result = []
        expected_fields = self.DYNAMIC_SECTION_FIELDS.get(section_type, [])

        for item in section_data:
            if isinstance(item, dict):
                # 确保所有必需字段都存在
                processed = {}
                for field in expected_fields:
                    processed[field] = item.get(field, '')
                # 保留原始数据中的其他字段
                for key, value in item.items():
                    if key not in processed:
                        processed[key] = value
                result.append(processed)
            elif isinstance(item, str):
                # 字符串转换为字典
                if section_type in ['awards', 'certificates']:
                    result.append({'name': item})
                elif section_type == 'skills':
                    result.append({'name': item, 'tailored_description': ''})
                else:
                    result.append({'name': item})

        return result

    def _post_process_tailored_content(self, doc: 'Document'):
        """
        渲染后处理：将 tailored 内容中的 \\n 拆分为独立段落。

        docxtpl 将 {{ var }} 中的 \\n 渲染为单个段落中的换行符，
        需要拆分为独立的 <w:p> 元素，复制原段落的格式。
        """
        from docx.oxml.ns import qn

        paragraphs = list(doc.paragraphs)
        for para in paragraphs:
            text = para.text
            if '\n' not in text:
                continue

            lines = text.split('\n')
            # 只有多行时才拆分
            if len(lines) <= 1:
                continue

            # 保留第一行的内容
            para.clear()
            para.add_run(lines[0])

            # 获取段落的 XML 元素
            para_element = para._element
            para_parent = para_element.getparent()

            if para_parent is None:
                continue

            # 找到当前段落在父元素中的位置
            para_index = list(para_parent).index(para_element)

            # 为后续每行创建新段落，复制格式
            for line in lines[1:]:
                if not line.strip():
                    continue
                new_para_element = copy.deepcopy(para_element)
                # 清除新段落中的 run 文本，设置为新行内容
                for r in new_para_element.findall(qn('w:r')):
                    for t in r.findall(qn('w:t')):
                        t.text = ''
                    # 设置第一个 run 的文本
                    r_t = r.find(qn('w:t'))
                    if r_t is not None:
                        r_t.text = line
                    break  # 只保留第一个 run

                # 移除多余的 run 元素（只保留第一个）
                runs = new_para_element.findall(qn('w:r'))
                for r in runs[1:]:
                    new_para_element.remove(r)

                # 在当前段落之后插入新段落
                para_parent.insert(para_index + 1, new_para_element)
                para_index += 1

            logger.debug(f"拆分 tailored 段落: {len(lines)} 行")

    def _generate_template_id(self, doc: 'Document', original_content: bytes = None) -> str:
        """
        生成模板ID - 基于内容哈希

        如果提供 original_content，使用 MD5 哈希（用于去重）
        否则回退到时间戳+UUID（兼容旧逻辑）

        Args:
            doc: Document 对象（未使用，保留用于未来扩展）
            original_content: 原始文件内容（字节）

        Returns:
            str: 模板ID（16位哈希或时间戳_UUID格式）
        """
        if original_content:
            # 基于内容哈希生成 ID（用于去重）
            content_hash = hashlib.md5(original_content).hexdigest()[:16]
            return content_hash
        else:
            # 回退方案：时间戳+UUID（兼容旧逻辑）
            timestamp = datetime.now().strftime('%Y%m%d_%H%M%S')
            unique_id = str(uuid.uuid4())[:8]
            return f"{timestamp}_{unique_id}"

    def _extract_style_metadata(self, doc: 'Document') -> StyleMetadata:
        """
        从 Word 文档提取样式元数据

        复用 ResumeParser 的逻辑
        """
        from .resume_parser import ResumeParser
        parser = ResumeParser()
        return parser._extract_style_metadata(doc)

    def get_stats(self) -> Dict[str, int]:
        """获取统计信息"""
        stats = self.stats.copy()
        stats.update({
            'detector_stats': self.detector.get_stats(),
            'inserter_stats': self.inserter.get_stats()
        })
        return stats

    def render_by_id(self, template_id: str, context: Dict[str, Any],
                     style_metadata: Optional[StyleMetadata] = None) -> bytes:
        """
        根据模板ID渲染简历

        支持两种模板来源：
        1. templates/preprocessed/ 目录（从原简历提取的模板）
        2. templates/builtin/ 目录（内置模板）

        Args:
            template_id: 模板ID
            context: 渲染上下文（AI 生成的数据）
            style_metadata: 样式元数据（备用）

        Returns:
            bytes: 渲染后的 Word 文档字节流
        """
        logger.info(f"🔍 render_by_id() - template_id: {template_id}")

        if not HAS_DOCXTPL:
            raise ImportError("未安装 docxtpl")

        # 1. 首先检查预处理后的模板
        preprocessed_path = self.template_dir / f"{template_id}.docx"
        logger.info(f"📁 检查预处理模板: {preprocessed_path}, 存在: {preprocessed_path.exists()}")

        if preprocessed_path.exists():
            try:
                logger.info(f"✅ 使用预处理模板渲染")
                doc = DocxTemplate(str(preprocessed_path))
                render_context = self._build_context(context)
                doc.render(render_context)
                self._post_process_tailored_content(doc)
                bio = io.BytesIO()
                doc.save(bio)
                bio.seek(0)

                self.stats['rendered'] += 1
                logger.info(f"✅ 使用预处理模板渲染成功: {template_id}")

                return bio.read()
            except Exception as e:
                logger.warning(f"❌ 预处理模板渲染失败: {e}，尝试动态预处理")

        # 2. 查找原始模板
        possible_paths = [
            config.BASE_DIR / 'templates' / 'builtin' / f"{template_id}.docx",  # builtin 目录
            config.BASE_DIR / 'templates' / 'uploaded' / f"{template_id}.docx",  # uploaded 目录
            config.BASE_DIR / 'templates' / 'extracted' / f"{template_id}.docx",  # extracted 目录
        ]

        logger.info(f"🔍 查找原始模板...")
        template_path = None
        for path in possible_paths:
            logger.info(f"   检查: {path}, 存在: {path.exists()}")
            if path.exists():
                template_path = path
                break

        if not template_path:
            logger.error(f"❌ 模板不存在: {template_id}")
            raise FileNotFoundError(f"模板不存在: {template_id}")

        logger.info(f"✅ 找到原始模板: {template_path}")

        # 3. 尝试动态预处理
        try:
            logger.info(f"🔧 尝试动态预处理: {template_path}")
            with open(template_path, 'rb') as f:
                file_content = f.read()
            doc = Document(io.BytesIO(file_content))
            preprocess_result = self.preprocess(
                doc,
                original_filename=f"{template_id}.docx",
                original_content=file_content
            )

            if preprocess_result.success:
                # 使用预处理后的模板渲染
                logger.info(f"✅ 动态预处理成功: {template_id}, 新 template_id: {preprocess_result.template_id}")
                return self.render(preprocess_result.template_id, context, style_metadata)
            else:
                logger.warning(f"❌ 动态预处理失败: {preprocess_result.error_message}")
        except Exception as e:
            logger.warning(f"❌ 动态预处理异常: {e}")

        # 4. 降级：直接使用原始模板（可能无法正确替换变量）
        try:
            logger.warning(f"⚠️ 降级：直接使用原始模板渲染（可能无法正确替换变量）: {template_id}")
            doc = DocxTemplate(str(template_path))
            render_context = self._build_context(context)
            doc.render(render_context)
            self._post_process_tailored_content(doc)
            bio = io.BytesIO()
            doc.save(bio)
            bio.seek(0)

            self.stats['rendered'] += 1
            logger.info(f"✅ 原始模板渲染完成: {template_id}")

            return bio.read()

        except Exception as e:
            self.stats['failed'] += 1
            logger.error(f"❌ 模板渲染失败: {e}", exc_info=True)
            raise

    def check_template_compatibility(self, template_id: str,
                                     context: Dict[str, Any]) -> Tuple[bool, List[str]]:
        """
        检查模板与数据的兼容性

        Args:
            template_id: 模板ID
            context: 渲染上下文

        Returns:
            Tuple[bool, List[str]]: (是否兼容, 缺失字段列表)
        """
        # 获取模板所需变量
        from .template_manager import template_manager
        template = template_manager.get_template(template_id)
        if not template:
            return False, ["模板不存在"]

        required_vars = set(template.get('variables', []))

        # 从上下文中提取已有变量
        available_vars = set()
        context_flat = self._flatten_context(context)
        available_vars.update(context_flat.keys())

        # 检查缺失的关键变量
        critical_vars = {'basic_info', 'work_experience', 'education'}
        missing_critical = critical_vars & required_vars - available_vars

        return len(missing_critical) == 0, list(required_vars - available_vars)

    def _flatten_context(self, context: Dict[str, Any], prefix: str = '') -> Dict[str, Any]:
        """扁平化上下文，用于变量检查"""
        result = {}
        for key, value in context.items():
            full_key = f"{prefix}_{key}" if prefix else key
            if isinstance(value, dict):
                result.update(self._flatten_context(value, full_key))
            elif isinstance(value, list):
                result[full_key] = value
                if value and isinstance(value[0], dict):
                    for i, item in enumerate(value):
                        result.update(self._flatten_context(item, f"{full_key}_{i}"))
            else:
                result[full_key] = value
        return result
        """
        清理旧模板文件

        Args:
            max_age_days: 最大保留天数
        """
        import time

        current_time = time.time()
        max_age_seconds = max_age_days * 24 * 60 * 60

        for template_file in self.template_dir.glob('*.docx'):
            file_age = current_time - template_file.stat().st_mtime
            if file_age > max_age_seconds:
                try:
                    template_file.unlink()
                    logger.info(f"清理旧模板: {template_file.name}")
                except Exception as e:
                    logger.warning(f"清理模板失败: {e}")
