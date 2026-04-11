"""
简历解析器模块

支持多种格式的简历解析，包括 PDF、Word、TXT、Markdown。
使用多级 fallback 机制确保解析成功率。
"""

import io
import re
import logging
from pathlib import Path
from typing import Optional, Dict, Any, List
from dataclasses import dataclass, field

# PDF 解析
try:
    import pdfplumber
    HAS_PDFPLUMBER = True
except ImportError:
    HAS_PDFPLUMBER = False

try:
    from PyPDF2 import PdfReader
    HAS_PYPDF2 = True
except ImportError:
    HAS_PYPDF2 = False

# Word 解析
try:
    from docx import Document
    from docx.oxml.ns import qn
    HAS_PYTHON_DOCX = True
except ImportError:
    HAS_PYTHON_DOCX = False

logger = logging.getLogger(__name__)


@dataclass
class StyleMetadata:
    """从原简历提取的样式元数据"""
    primary_font: str = "Microsoft YaHei"
    body_font_size: float = 10.5
    margin_top: float = 2.5
    margin_bottom: float = 2.5
    margin_left: float = 2.0
    margin_right: float = 2.0
    source: str = "default"  # "word", "pdf", "default"

    def get_name_font_size(self) -> float:
        """获取姓名字号 = 正文字号 × 1.7"""
        return round(self.body_font_size * 1.7, 1)

    def get_section_title_font_size(self) -> float:
        """获取章节标题字号 = 正文字号 × 1.14"""
        return round(self.body_font_size * 1.14, 1)

    def get_time_font_size(self) -> float:
        """获取时间字号 = 正文字号 × 0.95"""
        return round(self.body_font_size * 0.95, 1)

    def get_degree_font_size(self) -> float:
        """获取学位字号 = 正文字号 × 0.86"""
        return round(self.body_font_size * 0.86, 1)


@dataclass
class ParsedResume:
    """解析后的简历数据结构"""
    raw_text: str                                    # 原始文本
    basic_info: Dict[str, Any] = field(default_factory=dict)      # 基本信息
    education: List[Dict[str, Any]] = field(default_factory=list) # 教育背景
    work_experience: List[Dict[str, Any]] = field(default_factory=list)  # 工作经历
    projects: List[Dict[str, Any]] = field(default_factory=list)  # 项目经历
    skills: List[str] = field(default_factory=list)               # 技能
    awards: List[str] = field(default_factory=list)               # 奖项
    certificates: List[str] = field(default_factory=list)         # 证书
    self_evaluation: str = ""                         # 自我评价
    source_format: str = ""                           # 来源格式
    parse_confidence: float = 0.0                     # 解析置信度
    style_metadata: StyleMetadata = field(default_factory=StyleMetadata)  # 样式元数据


class FileParser:
    """文件解析器 - 负责将二进制流转换为文本"""

    def __init__(self):
        self.parse_stats = {
            'pdfplumber': 0,
            'pypdf2': 0,
            'docx': 0,
            'text': 0,
            'failed': 0
        }

    def parse_pdf(self, content: bytes) -> str:
        """PDF 解析 - 多级 fallback"""
        text = ""

        # Level 1: pdfplumber (推荐，效果最好)
        if HAS_PDFPLUMBER:
            try:
                with pdfplumber.open(io.BytesIO(content)) as pdf:
                    pages_text = []
                    for page in pdf.pages:
                        page_text = page.extract_text()
                        if page_text:
                            pages_text.append(page_text)
                    text = "\n\n".join(pages_text)

                if text and len(text.strip()) > 50:
                    self.parse_stats['pdfplumber'] += 1
                    logger.info(f"PDF 解析成功 (pdfplumber): {len(text)} 字符")
                    return text
            except Exception as e:
                logger.warning(f"pdfplumber 解析失败: {e}")

        # Level 2: PyPDF2 (备选)
        if HAS_PYPDF2:
            try:
                reader = PdfReader(io.BytesIO(content))
                pages_text = []
                for page in reader.pages:
                    page_text = page.extract_text()
                    if page_text:
                        pages_text.append(page_text)
                text = "\n\n".join(pages_text)

                if text and len(text.strip()) > 50:
                    self.parse_stats['pypdf2'] += 1
                    logger.info(f"PDF 解析成功 (PyPDF2): {len(text)} 字符")
                    return text
            except Exception as e:
                logger.warning(f"PyPDF2 解析失败: {e}")

        # Level 3: OCR (可选，需要额外安装)
        # TODO: 实现 OCR fallback

        self.parse_stats['failed'] += 1
        raise ValueError("PDF 解析失败，请尝试上传 Word 或文本格式")

    def parse_word(self, content: bytes) -> tuple:
        """
        Word 文档解析

        Returns:
            tuple: (Document对象, 提取的文本)
        """
        if not HAS_PYTHON_DOCX:
            raise ValueError("未安装 python-docx，无法解析 Word 文档")

        try:
            doc = Document(io.BytesIO(content))
            paragraphs = [p.text for p in doc.paragraphs if p.text.strip()]
            text = "\n".join(paragraphs)

            # 提取表格内容
            for table in doc.tables:
                for row in table.rows:
                    row_text = " | ".join(cell.text for cell in row.cells)
                    if row_text.strip():
                        text += "\n" + row_text

            self.parse_stats['docx'] += 1
            logger.info(f"Word 解析成功: {len(text)} 字符")
            return doc, text
        except Exception as e:
            self.parse_stats['failed'] += 1
            raise ValueError(f"Word 解析失败: {e}")

    def get_stats(self) -> Dict[str, int]:
        """获取解析统计信息"""
        return self.parse_stats.copy()


class StyleExtractor:
    """样式提取器 - 从文档对象中提取样式元数据"""

    @staticmethod
    def extract_from_word(doc: 'Document') -> StyleMetadata:
        """
        从 Word 文档提取样式元数据

        提取策略：
        1. 统计所有段落中最常见的字体和字号作为主样式
        2. 提取页面边距设置

        Args:
            doc: python-docx Document 对象

        Returns:
            StyleMetadata: 提取的样式元数据
        """
        from collections import Counter
        from docx.shared import Pt

        # 默认值
        primary_font = "Microsoft YaHei"
        body_font_size = 10.5
        margin_top = 2.5
        margin_bottom = 2.5
        margin_left = 2.0
        margin_right = 2.0

        try:
            # 1. 提取页面边距
            if doc.sections:
                section = doc.sections[0]
                # 转换为 cm
                margin_top = round(section.top_margin.cm, 1) if section.top_margin else 2.5
                margin_bottom = round(section.bottom_margin.cm, 1) if section.bottom_margin else 2.5
                margin_left = round(section.left_margin.cm, 1) if section.left_margin else 2.0
                margin_right = round(section.right_margin.cm, 1) if section.right_margin else 2.0

            # 2. 统计最常见的字体和字号
            font_counter = Counter()
            font_size_counter = Counter()

            for para in doc.paragraphs:
                if not para.text.strip():
                    continue

                # 遍历段落中的所有 run
                for run in para.runs:
                    # 提取字体
                    font_name = run.font.name
                    if font_name:
                        # 处理中文字体
                        east_asia_font = None
                        try:
                            rPr = run._element.rPr
                            if rPr is not None:
                                rFonts = rPr.find(qn('w:rFonts'))
                                if rFonts is not None:
                                    east_asia_font = rFonts.get(qn('w:eastAsia'))
                        except AttributeError:
                            pass

                        # 优先使用东亚字体（中文字体）
                        actual_font = east_asia_font or font_name
                        font_counter[actual_font] += len(run.text)

                    # 提取字号
                    font_size = run.font.size
                    if font_size:
                        size_pt = font_size.pt
                        if size_pt and 8 <= size_pt <= 24:  # 合理的字号范围
                            font_size_counter[size_pt] += len(run.text)

            # 3. 选择最常见的字体和字号
            if font_counter:
                primary_font = font_counter.most_common(1)[0][0]
                logger.info(f"提取到主字体: {primary_font}")

            if font_size_counter:
                body_font_size = font_size_counter.most_common(1)[0][0]
                logger.info(f"提取到正文字号: {body_font_size}pt")

            return StyleMetadata(
                primary_font=primary_font,
                body_font_size=body_font_size,
                margin_top=margin_top,
                margin_bottom=margin_bottom,
                margin_left=margin_left,
                margin_right=margin_right,
                source='word'
            )

        except Exception as e:
            logger.warning(f"样式提取失败，使用默认样式: {e}")
            return StyleMetadata(source='default')


class InfoExtractor:
    """信息提取器 - 从文本中提取结构化简历信息"""

    def extract_all(self, text: str) -> ParsedResume:
        """从文本中提取所有结构化信息"""
        parsed = ParsedResume(raw_text=text)

        # 提取基本信息
        parsed.basic_info = self._extract_basic_info(text)

        # 提取教育背景
        parsed.education = self._extract_education(text)

        # 提取工作经历
        parsed.work_experience = self._extract_work_experience(text)

        # 提取项目经历
        parsed.projects = self._extract_projects(text)

        # 提取技能
        parsed.skills = self._extract_skills(text)

        # 提取奖项
        parsed.awards = self._extract_awards(text)

        # 提取证书
        parsed.certificates = self._extract_certificates(text)

        # 提取自我评价
        parsed.self_evaluation = self._extract_self_evaluation(text)

        # 计算置信度
        parsed.parse_confidence = self._calculate_confidence(parsed)

        return parsed

    def _extract_basic_info(self, text: str) -> Dict[str, Any]:
        """提取基本信息"""
        info = {}

        # 姓名（通常在前几行）
        lines = text.split('\n')[:10]
        for line in lines:
            line = line.strip()
            if len(line) >= 2 and len(line) <= 10 and not any(c.isdigit() for c in line):
                # 可能是姓名
                if re.match(r'^[\u4e00-\u9fa5]+$', line):
                    info['name'] = line
                    break

        # 手机号
        phone_match = re.search(r'1[3-9]\d{9}', text)
        if phone_match:
            info['phone'] = phone_match.group()

        # 邮箱
        email_match = re.search(r'[\w.-]+@[\w.-]+\.\w+', text)
        if email_match:
            info['email'] = email_match.group()

        # 年龄/出生日期
        age_match = re.search(r'(\d{1,2})\s*[岁]', text)
        if age_match:
            info['age'] = int(age_match.group(1))

        # 性别
        if '男' in text[:200]:
            info['gender'] = '男'
        elif '女' in text[:200]:
            info['gender'] = '女'

        # 现居地
        location_patterns = [
            r'现居[：:]\s*([^\n]+)',
            r'所在地[：:]\s*([^\n]+)',
            r'居住地[：:]\s*([^\n]+)',
        ]
        for pattern in location_patterns:
            match = re.search(pattern, text)
            if match:
                info['location'] = match.group(1).strip()
                break

        return info

    def _extract_education(self, text: str) -> List[Dict[str, Any]]:
        """提取教育背景"""
        education = []

        # 常见教育背景标识
        edu_keywords = ['教育背景', '教育经历', '学历', '院校']
        lines = text.split('\n')

        in_education_section = False
        current_edu = {}

        for i, line in enumerate(lines):
            line = line.strip()

            # 检测教育背景部分
            if any(kw in line for kw in edu_keywords):
                in_education_section = True
                continue

            # 检测下一个部分（停止教育背景提取）
            if in_education_section and any(kw in line for kw in ['工作经历', '项目经历', '专业技能', '自我评价']):
                if current_edu:
                    education.append(current_edu)
                in_education_section = False
                continue

            if in_education_section and line:
                # 尝试提取学校、专业、学历
                # 格式1: 学校 | 专业 | 学历 | 时间
                if '|' in line:
                    parts = [p.strip() for p in line.split('|')]
                    if len(parts) >= 2:
                        current_edu = {
                            'school': parts[0],
                            'major': parts[1] if len(parts) > 1 else '',
                            'degree': parts[2] if len(parts) > 2 else '',
                            'time': parts[3] if len(parts) > 3 else ''
                        }
                        education.append(current_edu)
                        current_edu = {}
                # 格式2: 时间 学校 专业
                elif re.match(r'\d{4}', line):
                    time_match = re.match(r'(\d{4}[.\-]\d{1,2}[.\-]?\d{0,2}[~至至\-]?\d{0,4}[.\-]?\d{0,2})', line)
                    if time_match:
                        current_edu['time'] = time_match.group(1)
                        rest = line[time_match.end():].strip()
                        if rest:
                            # 尝试分离学校和专业
                            parts = re.split(r'[\s,，]+', rest, maxsplit=2)
                            if parts:
                                current_edu['school'] = parts[0]
                                if len(parts) > 1:
                                    current_edu['major'] = parts[1]
                                if len(parts) > 2:
                                    current_edu['degree'] = parts[2]
                            education.append(current_edu)
                            current_edu = {}

        return education

    def _extract_work_experience(self, text: str) -> List[Dict[str, Any]]:
        """提取工作经历"""
        experience = []

        work_keywords = ['工作经历', '工作经验', '职业经历', '工作背景']
        lines = text.split('\n')

        in_work_section = False
        current_work = {}
        content_lines = []

        for i, line in enumerate(lines):
            line = line.strip()

            # 检测工作经历部分
            if any(kw in line for kw in work_keywords):
                in_work_section = True
                continue

            # 检测下一个部分
            if in_work_section and any(kw in line for kw in ['项目经历', '专业技能', '教育背景', '自我评价']):
                if current_work:
                    current_work['content'] = '\n'.join(content_lines)
                    experience.append(current_work)
                in_work_section = False
                current_work = {}
                content_lines = []
                continue

            if in_work_section and line:
                # 尝试提取公司、职位、时间
                # 格式: 时间 公司 职位
                if re.match(r'\d{4}', line) and ('年' in line or '-' in line or '.' in line):
                    if current_work:
                        current_work['content'] = '\n'.join(content_lines)
                        experience.append(current_work)
                        content_lines = []

                    # 解析新条目
                    time_match = re.match(r'(\d{4}[.\-]\d{1,2}[~至至\-]?\d{0,4}[.\-]?\d{0,2})', line)
                    if time_match:
                        current_work = {'time': time_match.group(1)}
                        rest = line[time_match.end():].strip()
                        # 尝试分离公司和职位
                        parts = re.split(r'[\s|｜]+', rest, maxsplit=1)
                        if parts:
                            current_work['company'] = parts[0]
                            if len(parts) > 1:
                                current_work['position'] = parts[1]
                else:
                    # 工作内容描述
                    content_lines.append(line)

        # 添加最后一个条目
        if current_work:
            current_work['content'] = '\n'.join(content_lines)
            experience.append(current_work)

        return experience

    def _extract_projects(self, text: str) -> List[Dict[str, Any]]:
        """提取项目经历"""
        projects = []

        project_keywords = ['项目经历', '项目经验', '项目背景']
        lines = text.split('\n')

        in_project_section = False
        current_project = {}
        content_lines = []

        for line in lines:
            line = line.strip()

            if any(kw in line for kw in project_keywords):
                in_project_section = True
                continue

            if in_project_section and any(kw in line for kw in ['工作经历', '专业技能', '教育背景', '自我评价', '奖项']):
                if current_project:
                    current_project['content'] = '\n'.join(content_lines)
                    projects.append(current_project)
                in_project_section = False
                current_project = {}
                content_lines = []
                continue

            if in_project_section and line:
                # 项目名称和时间
                if re.match(r'\d{4}', line):
                    if current_project:
                        current_project['content'] = '\n'.join(content_lines)
                        projects.append(current_project)
                        content_lines = []

                    time_match = re.match(r'(\d{4}[.\-]\d{1,2}[~至至\-]?\d{0,4}[.\-]?\d{0,2})', line)
                    if time_match:
                        current_project = {'time': time_match.group(1)}
                        rest = line[time_match.end():].strip()
                        if rest:
                            current_project['name'] = rest
                else:
                    content_lines.append(line)

        if current_project:
            current_project['content'] = '\n'.join(content_lines)
            projects.append(current_project)

        return projects

    def _extract_skills(self, text: str) -> List[str]:
        """提取技能"""
        skills = []

        skill_keywords = ['专业技能', '技能特长', '技术栈', '掌握技能']
        lines = text.split('\n')

        in_skill_section = False

        for line in lines:
            line = line.strip()

            if any(kw in line for kw in skill_keywords):
                in_skill_section = True
                continue

            if in_skill_section and any(kw in line for kw in ['工作经历', '项目经历', '教育背景', '自我评价', '奖项']):
                in_skill_section = False
                continue

            if in_skill_section and line:
                # 分割技能标签
                if '、' in line:
                    skills.extend([s.strip() for s in line.split('、') if s.strip()])
                elif '，' in line:
                    skills.extend([s.strip() for s in line.split('，') if s.strip()])
                elif ',' in line:
                    skills.extend([s.strip() for s in line.split(',') if s.strip()])
                else:
                    skills.append(line)

        return list(set(skills))  # 去重

    def _extract_awards(self, text: str) -> List[str]:
        """提取奖项荣誉"""
        awards = []

        award_keywords = ['奖项荣誉', '获奖情况', '荣誉证书', '所获荣誉']
        lines = text.split('\n')

        in_award_section = False

        for line in lines:
            line = line.strip()

            if any(kw in line for kw in award_keywords):
                in_award_section = True
                continue

            if in_award_section and any(kw in line for kw in ['工作经历', '项目经历', '教育背景', '自我评价', '专业技能']):
                in_award_section = False
                continue

            if in_award_section and line:
                awards.append(line)

        return awards

    def _extract_certificates(self, text: str) -> List[str]:
        """提取证书资质"""
        certificates = []

        cert_keywords = ['证书资质', '资格证书', '执业资格', '专业证书']
        lines = text.split('\n')

        in_cert_section = False

        for line in lines:
            line = line.strip()

            if any(kw in line for kw in cert_keywords):
                in_cert_section = True
                continue

            if in_cert_section and any(kw in line for kw in ['工作经历', '项目经历', '教育背景', '自我评价', '专业技能', '奖项']):
                in_cert_section = False
                continue

            if in_cert_section and line:
                certificates.append(line)

        return certificates

    def _extract_self_evaluation(self, text: str) -> str:
        """提取自我评价"""
        eval_keywords = ['自我评价', '个人简介', '个人总结', '自我介绍']
        lines = text.split('\n')

        in_eval_section = False
        eval_lines = []

        for line in lines:
            line = line.strip()

            if any(kw in line for kw in eval_keywords):
                in_eval_section = True
                continue

            if in_eval_section and any(kw in line for kw in ['工作经历', '项目经历', '教育背景', '专业技能', '奖项']):
                break

            if in_eval_section and line:
                eval_lines.append(line)

        return '\n'.join(eval_lines)

    def _calculate_confidence(self, parsed: ParsedResume) -> float:
        """
        计算解析置信度

        适用于所有候选人类型（有经验/应届生/无经验）
        - 有工作经验者：工作经历权重较高
        - 无工作经验者：教育背景和项目经历权重较高
        """
        score = 0.0

        # 基本信息（权重20%）
        if parsed.basic_info.get('name'):
            score += 0.08
        if parsed.basic_info.get('phone') or parsed.basic_info.get('email'):
            score += 0.08
        if parsed.basic_info.get('age') or parsed.basic_info.get('gender'):
            score += 0.04

        # 判断候选人类型
        has_work_experience = len(parsed.work_experience) > 0

        if has_work_experience:
            # 有工作经验者
            # 教育背景（权重15%）
            if parsed.education:
                score += min(0.15, len(parsed.education) * 0.08)

            # 工作经历（权重45%）
            if parsed.work_experience:
                score += min(0.45, len(parsed.work_experience) * 0.2)

            # 项目经历（权重10%）
            if parsed.projects:
                score += min(0.10, len(parsed.projects) * 0.05)
        else:
            # 无工作经验者（应届生/转行者）
            # 教育背景（权重40%）
            if parsed.education:
                score += min(0.40, len(parsed.education) * 0.15)

            # 项目经历（权重30%）
            if parsed.projects:
                score += min(0.30, len(parsed.projects) * 0.12)

        # 技能（权重10%）
        if parsed.skills:
            score += min(0.10, len(parsed.skills) * 0.03)

        # 奖项/证书（额外加分，最多5%）
        if parsed.awards or parsed.certificates:
            bonus = (len(parsed.awards) + len(parsed.certificates)) * 0.02
            score += min(0.05, bonus)

        return min(1.0, score)


class ResumeParser:
    """简历解析器 - 组合各个子模块完成解析"""

    SUPPORTED_FORMATS = ['.pdf', '.docx', '.doc', '.txt', '.md']

    def __init__(self):
        # 使用组合替代继承，将职责委托给专门的类
        self.file_parser = FileParser()
        self.style_extractor = StyleExtractor()
        self.info_extractor = InfoExtractor()

    def parse(self, file_path: str = None, file_content: bytes = None,
              filename: str = None) -> ParsedResume:
        """
        解析简历文件

        Args:
            file_path: 文件路径
            file_content: 文件二进制内容
            filename: 文件名（用于判断格式）

        Returns:
            ParsedResume: 解析后的简历数据
        """
        # 确定文件格式
        if file_path:
            suffix = Path(file_path).suffix.lower()
            with open(file_path, 'rb') as f:
                content = f.read()
        elif file_content and filename:
            suffix = Path(filename).suffix.lower()
            content = file_content
        else:
            raise ValueError("必须提供 file_path 或 (file_content + filename)")

        if suffix not in self.SUPPORTED_FORMATS:
            raise ValueError(f"不支持的文件格式: {suffix}")

        # 根据格式选择解析方法
        style_metadata = StyleMetadata()  # 默认样式
        if suffix == '.pdf':
            text = self.file_parser.parse_pdf(content)
            source_format = 'pdf'
            style_metadata.source = 'pdf'
        elif suffix in ['.docx', '.doc']:
            doc, text = self.file_parser.parse_word(content)
            source_format = 'word'
            # 提取 Word 文档的样式元数据
            style_metadata = self.style_extractor.extract_from_word(doc)
            style_metadata.source = 'word'
        else:
            text = content.decode('utf-8', errors='ignore')
            source_format = 'text'
            style_metadata.source = 'default'

        if not text or len(text.strip()) < 50:
            raise ValueError("解析结果内容过少，可能解析失败")

        # 解析结构化信息
        parsed = self.info_extractor.extract_all(text)
        parsed.source_format = source_format
        parsed.style_metadata = style_metadata

        return parsed

    def get_stats(self) -> Dict[str, int]:
        """获取解析统计信息"""
        return self.file_parser.get_stats()
