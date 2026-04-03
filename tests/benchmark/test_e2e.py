"""TC-E01~E02: 端到端集成测试"""
import pytest
import sys
from pathlib import Path

# 确保能导入 core 模块
sys.path.insert(0, str(Path(__file__).parent.parent.parent))


class TestEndToEnd:
    """端到端测试：从模板预处理到渲染输出"""

    # TC-E01: JinjaInserter 生成 tailored 变量
    def test_jinja_inserter_creates_tailored_variables(self):
        """验证 JinjaInserter 处理动态条目时生成 tailored 变量"""
        try:
            from core.structure_detector import SectionType, EntryInfo
            from core.jinja_inserter import JinjaTagInserter
            from docx import Document
        except ImportError as e:
            pytest.skip(f"无法导入依赖: {e}")

        # 创建一个最小化的测试文档
        doc = Document()

        # 添加标题段落（工作经历）
        doc.add_paragraph('工作经历')

        # 添加一个工作经历条目标题
        doc.add_paragraph('2020.01-2023.12  XX公司  工程师')

        # 添加内容段落
        doc.add_paragraph('负责系统开发和维护')
        doc.add_paragraph('参与项目架构设计')

        entry = EntryInfo(
            entry_type=SectionType.WORK,
            paragraph_index=1,  # entry_para 的索引
            time='2020.01-2023.12',
            organization='XX公司',
            role='工程师',
            content_paragraphs=[2, 3]  # content_para1, content_para2 的索引
        )

        inserter = JinjaTagInserter()
        inserter._insert_entry_simple(doc, entry, 'exp', 'work_experience', 0)

        # 验证内容段落被替换为 tailored 变量
        assert '{{ work_experience_0_tailored }}' in doc.paragraphs[2].text, (
            f"内容段落应包含 tailored 变量，实际: '{doc.paragraphs[2].text}'"
        )

        # 验证第二个内容段落被清除
        assert doc.paragraphs[3].text.strip() == '', (
            f"多余内容段落应被清除，实际: '{doc.paragraphs[3].text}'"
        )

    # TC-E02: _post_process_tailored_content 拆分换行
    def test_post_process_splits_newlines(self):
        """验证渲染后处理能正确拆分 \\n 为独立段落"""
        try:
            from docx import Document
            from core.template_processor import TemplateProcessor
        except ImportError as e:
            pytest.skip(f"无法导入依赖: {e}")

        doc = Document()
        doc.add_paragraph('第一行内容\n第二行内容\n第三行内容')

        processor = TemplateProcessor()
        processor._post_process_tailored_content(doc)

        # 应拆分为 3 个段落
        texts = [p.text for p in doc.paragraphs if p.text.strip()]
        assert len(texts) == 3, (
            f"换行拆分后应有 3 个段落，实际 {len(texts)} 个: {texts}"
        )
        assert texts[0] == '第一行内容'
        assert texts[1] == '第二行内容'
        assert texts[2] == '第三行内容'
