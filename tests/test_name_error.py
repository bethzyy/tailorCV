"""
自测脚本：复现 KeyError('"name"') 错误

用用户实际模板 c3d6868643fdf645.docx + 模拟 tailored_resume 数据，
直接调用 template_processor 的各个方法，捕获完整 traceback。
"""
import sys
import os
import traceback
import io

sys.path.insert(0, os.path.dirname(os.path.dirname(os.path.abspath(__file__))))

from docx import Document


def test_extracted_template():
    """测试 1: 直接用 extracted 模板渲染（会触发动态预处理）"""
    from core.template_processor import TemplateProcessor

    print("=" * 60)
    print("测试 1: render_by_id('c3d6868643fdf645') — 模拟实际调用路径")
    print("=" * 60)

    processor = TemplateProcessor()

    # 模拟 AI 生成的 tailored_resume
    tailored_resume = {
        'basic_info': {
            'name': '张三',
            'phone': '138-0000-0001',
            'email': 'zhangsan@example.com',
            'location': '北京',
            'age': 28
        },
        'summary': '资深软件工程师，5年经验',
        'education': [
            {
                'school': '北京大学',
                'degree': '硕士',
                'major': '计算机',
                'time': '2016-2019',
                'tailored': '北大 计算机硕士 2016-2019'
            }
        ],
        'work_experience': [
            {
                'company': '科技有限公司',
                'position': '高级工程师',
                'time': '2019-至今',
                'tailored': '• 主导架构设计\n• 带领团队'
            }
        ],
        'projects': [
            {
                'name': '推荐系统',
                'role': '负责人',
                'time': '2021-2022',
                'tailored': '• 设计推荐算法'
            }
        ],
        'skills': [
            {'name': 'Python', 'tailored_description': '精通'},
            {'name': 'Java', 'tailored_description': '熟练'}
        ],
        'awards': [],
        'certificates': [],
        'self_evaluation': '良好的团队协作能力'
    }

    try:
        result = processor.render_by_id('c3d6868643fdf645', tailored_resume)
        print(f"✅ 渲染成功! 输出大小: {len(result)} bytes")
        return True
    except Exception as e:
        print(f"❌ 渲染失败!")
        print(f"   异常类型: {type(e).__name__}")
        print(f"   异常消息: {str(e)}")
        print(f"   完整 traceback:")
        traceback.print_exc()
        return False


def test_preprocessed_template():
    """测试 2: 先预处理，再用预处理的模板渲染"""
    from core.template_processor import TemplateProcessor

    print("\n" + "=" * 60)
    print("测试 2: preprocess() + render() — 两步走")
    print("=" * 60)

    processor = TemplateProcessor()

    # 读取 extracted 模板
    template_path = os.path.join(os.path.dirname(os.path.dirname(os.path.abspath(__file__))),
                                  'templates', 'extracted', 'c3d6868643fdf645.docx')
    if not os.path.exists(template_path):
        print(f"❌ 模板文件不存在: {template_path}")
        return False

    with open(template_path, 'rb') as f:
        file_content = f.read()

    doc = Document(io.BytesIO(file_content))

    tailored_resume = {
        'basic_info': {'name': '张三', 'phone': '138-0000-0001',
                       'email': 'zhangsan@example.com', 'location': '北京'},
        'summary': '资深工程师',
        'education': [{'school': '北大', 'degree': '硕士', 'major': 'CS',
                       'time': '2016-2019', 'tailored': '北大 CS 硕士'}],
        'work_experience': [{'company': '科技公司', 'position': '工程师',
                            'time': '2019-至今', 'tailored': '• 主导开发'}],
        'projects': [{'name': '推荐系统', 'role': '负责人',
                      'time': '2021-2022', 'tailored': '• 设计算法'}],
        'skills': [{'name': 'Python', 'tailored_description': '精通'}],
        'awards': [], 'certificates': [], 'self_evaluation': '好'
    }

    try:
        # Step 1: preprocess
        print("   Step 1: preprocess()...")
        preprocess_result = processor.preprocess(doc, 'c3d6868643fdf645.docx',
                                                 original_content=file_content)
        print(f"   preprocess 结果: success={preprocess_result.success}, "
              f"template_id={preprocess_result.template_id}, "
              f"confidence={preprocess_result.metadata.structure_confidence}")

        # Step 2: render
        print("   Step 2: render()...")
        result = processor.render(preprocess_result.template_id, tailored_resume)
        print(f"✅ 渲染成功! 输出大小: {len(result)} bytes")
        return True
    except Exception as e:
        print(f"❌ 失败!")
        print(f"   异常类型: {type(e).__name__}")
        print(f"   异常消息: {str(e)}")
        print(f"   完整 traceback:")
        traceback.print_exc()
        return False


def test_docxtpl_direct():
    """测试 3: 直接用 docxtpl 渲染一个含 {{ basic_info.name }} 的文档"""
    from docxtpl import DocxTemplate

    print("\n" + "=" * 60)
    print("测试 3: docxtpl 直接渲染 — 最小复现")
    print("=" * 60)

    # 创建一个包含 {{ basic_info.name }} 的文档
    doc = Document()
    p = doc.add_paragraph()
    run = p.add_run("{{ basic_info.name }}")

    bio = io.BytesIO()
    doc.save(bio)
    bio.seek(0)

    try:
        tpl = DocxTemplate(bio)
        tpl.render({'basic_info': {'name': '张三'}})
        print("✅ docxtpl 直接渲染成功")

        # 验证输出
        bio.seek(0)
        result_doc = Document(bio)
        for para in result_doc.paragraphs:
            if para.text.strip():
                print(f"   输出文本: {repr(para.text)}")
    except Exception as e:
        print(f"❌ docxtpl 直接渲染失败!")
        print(f"   异常类型: {type(e).__name__}")
        print(f"   异常消息: {str(e)}")
        traceback.print_exc()

    # 测试 2: basic_info 是字符串而不是 dict
    print("\n   --- 测试 basic_info 为字符串 ---")
    bio.seek(0)
    try:
        tpl = DocxTemplate(bio)
        tpl.render({'basic_info': '张三'})  # 错误：应该是 dict
        print("✅ basic_info=字符串 渲染成功")
    except Exception as e:
        print(f"   异常: {type(e).__name__}: {str(e)}")

    # 测试 3: basic_info 是缺少 name 键的 dict
    print("\n   --- 测试 basic_info 缺少 name 键 ---")
    bio.seek(0)
    try:
        tpl = DocxTemplate(bio)
        tpl.render({'basic_info': {'phone': '138'}})
        print("✅ basic_info 缺少 name 渲染成功")
    except Exception as e:
        print(f"   异常: {type(e).__name__}: {str(e)}")

    return True


def test_build_context():
    """测试 4: _build_context 是否安全"""
    from core.template_processor import TemplateProcessor

    print("\n" + "=" * 60)
    print("测试 4: _build_context() 安全性")
    print("=" * 60)

    processor = TemplateProcessor()

    # 正常数据
    normal_data = {
        'basic_info': {'name': '张三'},
        'work_experience': [{'company': 'A', 'tailored': 'B'}],
        'projects': [{'name': 'X', 'tailored': 'Y'}],
    }

    try:
        ctx = processor._build_context(normal_data)
        print(f"✅ 正常数据: context keys = {list(ctx.keys())[:10]}")
    except Exception as e:
        print(f"❌ 正常数据失败: {type(e).__name__}: {e}")
        traceback.print_exc()

    # basic_info 是字符串
    bad_data_1 = {
        'basic_info': '张三',  # 应该是 dict
        'work_experience': [],
    }
    try:
        ctx = processor._build_context(bad_data_1)
        print(f"✅ basic_info=字符串: context['basic_info'] = {repr(ctx.get('basic_info'))}")
    except Exception as e:
        print(f"❌ basic_info=字符串 失败: {type(e).__name__}: {e}")
        traceback.print_exc()

    # basic_info 缺少 name 键
    bad_data_2 = {
        'basic_info': {'phone': '138'},
        'work_experience': [],
    }
    try:
        ctx = processor._build_context(bad_data_2)
        print(f"✅ basic_info 缺 name: context['basic_info'] = {repr(ctx.get('basic_info'))}")
    except Exception as e:
        print(f"❌ basic_info 缺 name 失败: {type(e).__name__}: {e}")
        traceback.print_exc()

    # 空数据
    empty_data = {}
    try:
        ctx = processor._build_context(empty_data)
        print(f"✅ 空数据: context keys = {list(ctx.keys())}")
    except Exception as e:
        print(f"❌ 空数据失败: {type(e).__name__}: {e}")
        traceback.print_exc()

    return True


def test_extracted_template_xml():
    """测试 5: 检查 extracted 模板的 XML 内容"""
    import zipfile

    print("\n" + "=" * 60)
    print("测试 5: 检查模板 XML 中的 Jinja2 标签")
    print("=" * 60)

    template_path = os.path.join(os.path.dirname(os.path.dirname(os.path.abspath(__file__))),
                                  'templates', 'extracted', 'c3d6868643fdf645.docx')
    if not os.path.exists(template_path):
        print(f"❌ 模板不存在: {template_path}")
        return False

    try:
        with zipfile.ZipFile(template_path) as zf:
            with zf.open('word/document.xml') as f:
                content = f.read().decode('utf-8')

        # 查找所有 Jinja2 相关内容
        import re
        # 找 {{ }} 标签
        jinja_tags = re.findall(r'\{\{.*?\}\}', content)
        print(f"   找到 {len(jinja_tags)} 个 Jinja2 标签:")
        for tag in jinja_tags[:20]:
            print(f"   {tag}")

        # 找 {%% %%} 标签
        block_tags = re.findall(r'\{%.*?%\}', content)
        if block_tags:
            print(f"   ⚠️ 找到 {len(block_tags)} 个块标签:")
            for tag in block_tags:
                print(f"   {tag}")

        # 找包含 "name" 的 w:t 元素
        name_elements = re.findall(r'<w:t[^>]*>[^<]*name[^<]*</w:t>', content, re.IGNORECASE)
        print(f"\n   包含 'name' 的 <w:t> 元素: {len(name_elements)}")
        for elem in name_elements[:10]:
            print(f"   {elem[:100]}")

    except Exception as e:
        print(f"❌ XML 检查失败: {e}")
        traceback.print_exc()
        return False

    return True


if __name__ == '__main__':
    results = []
    results.append(("docxtpl 直接渲染", test_docxtpl_direct()))
    results.append(("_build_context 安全性", test_build_context()))
    results.append(("模板 XML 检查", test_extracted_template_xml()))
    results.append(("extracted 模板 render_by_id", test_extracted_template()))
    results.append(("preprocess + render", test_preprocessed_template()))

    print("\n" + "=" * 60)
    print("测试结果汇总")
    print("=" * 60)
    for name, passed in results:
        status = "✅ 通过" if passed else "❌ 失败"
        print(f"   {status}: {name}")
