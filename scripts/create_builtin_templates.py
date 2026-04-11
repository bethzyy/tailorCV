"""
测试 create_builtin_templates 模块
"""

import sys
from pathlib import Path

import pytest
from docx import Document
from docx.shared import Pt, RGBColor

# 添加项目根目录到路径
sys.path.insert(0, str(Path(__file__).parent.parent))

from scripts.create_builtin_templates import (
    set_style,
    create_classic_professional,
    create_modern_minimal,
    create_creative_design,
    create_executive_senior,
    create_academic_research,
    create_tech_engineer,
    main,
)


class TestSetStyle:
    """测试 set_style 函数"""

    def test_default_font_name(self):
        """测试默认字体名称"""
        doc = Document()
        paragraph = doc.add_paragraph()
        run = paragraph.add_run('测试')
        set_style(paragraph)
        assert run.font.name == '微软雅黑'

    def test_default_font_size(self):
        """测试默认字体大小"""
        doc = Document()
        paragraph = doc.add_paragraph()
        run = paragraph.add_run('测试')
        set_style(paragraph)
        assert run.font.size == Pt(11)

    def test_default_bold(self):
        """测试默认不加粗"""
        doc = Document()
        paragraph = doc.add_paragraph()
        run = paragraph.add_run('测试')
        set_style(paragraph)
        assert run.font.bold is False

    def test_custom_font_name(self):
        """测试自定义字体名称"""
        doc = Document()
        paragraph = doc.add_paragraph()
        run = paragraph.add_run('测试')
        set_style(paragraph, font_name='宋体')
        assert run.font.name == '宋体'

    def test_custom_font_size(self):
        """测试自定义字体大小"""
        doc = Document()
        paragraph = doc.add_paragraph()
        run = paragraph.add_run('测试')
        set_style(paragraph, font_size=14)
        assert run.font.size == Pt(14)

    def test_bold_true(self):
        """测试加粗"""
        doc = Document()
        paragraph = doc.add_paragraph()
        run = paragraph.add_run('测试')
        set_style(paragraph, bold=True)
        assert run.font.bold is True

    def test_custom_color(self):
        """测试自定义颜色"""
        doc = Document()
        paragraph = doc.add_paragraph()
        run = paragraph.add_run('测试')
        set_style(paragraph, color=(255, 0, 0))
        assert run.font.color.rgb == RGBColor(255, 0, 0)

    def test_no_color_by_default(self):
        """测试默认无颜色"""
        doc = Document()
        paragraph = doc.add_paragraph()
        run = paragraph.add_run('测试')
        set_style(paragraph)
        assert run.font.color.rgb is None

    def test_multiple_runs(self):
        """测试多个 run 都被设置样式"""
        doc = Document()
        paragraph = doc.add_paragraph()
        run1 = paragraph.add_run('测试1')
        run2 = paragraph.add_run('测试2')
        set_style(paragraph, font_name='宋体', bold=True)
        assert run1.font.name == '宋体'
        assert run1.font.bold is True
        assert run2.font.name == '宋体'
        assert run2.font.bold is True


class TestCreateClassicProfessional:
    """测试经典专业模板"""

    def test_returns_document(self):
        """测试返回 Document 对象"""
        doc = create_classic_professional()
        assert isinstance(doc, Document)

    def test_contains_basic_info_markers(self):
        """测试包含基本信息标记"""
        doc = create_classic_professional()
        text = '\n'.join(p.text for p in doc.paragraphs)
        assert '{{ basic_info.name }}' in text
        assert '{{ basic_info.phone }}' in text
        assert '{{ basic_info.email }}' in text

    def test_contains_section_markers(self):
        """测试包含各板块标记"""
        doc = create_classic_professional()
        text = '\n'.join(p.text for p in doc.paragraphs)
        assert '{{ summary }}' in text
        assert '{% for edu in education %}' in text
        assert '{% for work in work_experience %}' in text
        assert '{% for proj in projects %}' in text
        assert '{% for skill in skills %}' in text


class TestCreateModernMinimal:
    """测试现代简约模板"""

    def test_returns_document(self):
        """测试返回 Document 对象"""
        doc = create_modern_minimal()
        assert isinstance(doc, Document)

    def test_contains_basic_info_markers(self):
        """测试包含基本信息标记"""
        doc = create_modern_minimal()
        text = '\n'.join(p.text for p in doc.paragraphs)
        assert '{{ basic_info.name }}' in text
        assert '{{ basic_info.email }}' in text
        assert '{{ basic_info.phone }}' in text

    def test_contains_section_markers(self):
        """测试包含各板块标记"""
        doc = create_modern_minimal()
        text = '\n'.join(p.text for p in doc.paragraphs)
        assert '{{ summary }}' in text
        assert '{% for work in work_experience %}' in text


class TestCreateCreativeDesign:
    """测试创意设计模板"""

    def test_returns_document(self):
        """测试返回 Document 对象"""
        doc = create_creative_design()
        assert isinstance(doc, Document)

    def test_contains_basic_info_markers(self):
        """测试包含基本信息标记"""
        doc = create_creative_design()
        text = '\n'.join(p.text for p in doc.paragraphs)
        assert '{{ basic_info.name }}' in text
        assert '{{ basic_info.phone }}' in text
        assert '{{ basic_info.email }}' in text

    def test_contains_section_markers(self):
        """测试包含各板块标记"""
        doc = create_creative_design()
        text = '\n'.join(p.text for p in doc.paragraphs)
        assert '{{ summary }}' in text
        assert '{% for work in work_experience %}' in text


class TestCreateExecutiveSenior:
    """测试高管资深模板"""

    def test_returns_document(self):
        """测试返回 Document 对象"""
        doc = create_executive_senior()
        assert isinstance(doc, Document)

    def test_contains_basic_info_markers(self):
        """测试包含基本信息标记"""
        doc = create_executive_senior()
        text = '\n'.join(p.text for p in doc.paragraphs)
        assert '{{ basic_info.name }}' in text
        assert '{{ basic_info.phone }}' in text
        assert '{{ basic_info.email }}' in text

    def test_contains_section_markers(self):
        """测试包含各板块标记"""
        doc = create_executive_senior()
        text = '\n'.join(p.text for p in doc.paragraphs)
        assert '{{ summary }}' in text
        assert '{% for work in work_experience %}' in text


class TestCreateAcademicResearch:
    """测试学术研究模板"""

    def test_returns_document(self):
        """测试返回 Document 对象"""
        doc = create_academic_research()
        assert isinstance(doc, Document)

    def test_contains_basic_info_markers(self):
        """测试包含基本信息标记"""
        doc = create_academic_research()
        text = '\n'.join(p.text for p in doc.paragraphs)
        assert '{{ basic_info.name }}' in text
        assert '{{ basic_info.email }}' in text
        assert '{{ basic_info.phone }}' in text

    def test_contains_section_markers(self):
        """测试包含各板块标记"""
        doc = create_academic_research()
        text = '\n'.join(p.text for p in doc.paragraphs)
        assert '{{ summary }}' in text
        assert '{% for edu in education %}' in text
        assert '{% for award in awards %}' in text


class TestCreateTechEngineer:
    """测试技术工程师模板"""

    def test_returns_document(self):
        """测试返回 Document 对象"""
        doc = create_tech_engineer()
        assert isinstance(doc, Document)

    def test_contains_basic_info_markers(self):
        """测试包含基本信息标记"""
        doc = create_tech_engineer()
        text = '\n'.join(p.text for p in doc.paragraphs)
        assert '{{ basic_info.name }}' in text
        assert '{{ basic_info.phone }}' in text
        assert '{{ basic_info.email }}' in text

    def test_contains_section_markers(self):
        """测试包含各板块标记"""
        doc = create_tech_engineer()
        text = '\n'.join(p.text for p in doc.paragraphs)
        assert '{{ summary }}' in text
        assert '{% for skill in skills %}' in text
        assert '{% for work in work_experience %}' in text


class TestMain:
    """测试 main 函数"""

    def test_main_creates_all_template_files(self, tmp_path):
        """测试 main 函数创建所有模板文件"""
        import scripts.create_builtin_templates as cbt

        original_file = cbt.__file__
        cbt.__file__ = str(tmp_path / 'fake_script.py')
        try:
            cbt.main()
        finally:
            cbt.__file__ = original_file

        builtin_dir = tmp_path / 'templates' / 'builtin'
        expected_files = [
            'classic_professional.docx',
            'modern_minimal.docx',
            'creative_design.docx',
            'executive_senior.docx',
            'academic_research.docx',
            'tech_engineer.docx',
        ]

        for filename in expected_files:
            assert (builtin_dir / filename).exists(), f'模板文件 {filename} 未创建'

    def test_main_created_files_are_valid_docx(self, tmp_path):
        """测试 main 函数创建的文件是有效的 docx 文件"""
        import scripts.create_builtin_templates as cbt

        original_file = cbt.__file__
        cbt.__file__ = str(tmp_path / 'fake_script.py')
        try:
            cbt.main()
        finally:
            cbt.__file__ = original_file

        builtin_dir = tmp_path / 'templates' / 'builtin'
        for docx_file in builtin_dir.glob('*.docx'):
            doc = Document(str(docx_file))
            assert len(doc.paragraphs) > 0, f'{docx_file.name} 没有段落'
